# ocr_agent.py
# Enhanced OCR Agent mit Streaming-Unterstützung 

import asyncio
import base64
import gc
import hashlib
import io
import json
import logging
import os
import re
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any, Union, Literal
from pydantic import BaseModel, Field

# Externe Abhängigkeiten
import asyncpg
import pytesseract
from PIL import Image
import pdf2image

# Optional imports mit besserer Fehlerbehandlung
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logging.warning("python-docx not available - DOCX support disabled")

try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False
    logging.warning("pandas not available - Excel/CSV support limited")

try:
    from bs4 import BeautifulSoup
    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False
    logging.warning("beautifulsoup4 not available - HTML parsing limited")

import email
from email.parser import BytesParser
from email import policy

# Logging setup
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger("EnhancedOCRAgentV2")


# === PYDANTIC MODELS FÜR TYPE SAFETY ===

class DocumentMetadata(BaseModel):
    """Strukturierte Metadaten für Dokumente mit N8N-Kompatibilität"""
    
    # Content sources (mindestens eins muss vorhanden sein)
    file_content: Optional[str] = Field(None, description="Base64-encoded file content")
    file_path: Optional[str] = Field(None, description="Path to file on disk")
    
    # Processing configuration
    processing_mode: Literal["streaming", "in_memory"] = Field("in_memory", description="Processing strategy")
    source: str = Field("unknown", description="Upload source (api, n8n, manual)")
    
    # File metadata
    original_filename: Optional[str] = Field(None, description="Original filename")
    file_size_mb: Optional[float] = Field(None, description="File size in MB")
    estimated_pages: Optional[int] = Field(None, description="Estimated number of pages")
    content_type: str = Field("base64", description="Content encoding type")
    
    # N8N specific fields
    uploadedAt: Optional[str] = Field(None, description="Upload timestamp")
    processed: bool = Field(False, description="Processing status")
    
    # Legacy/alternative field names (for compatibility)
    fileContent: Optional[str] = Field(None, description="Alternative camelCase field")
    filePath: Optional[str] = Field(None, description="Alternative camelCase field")
    
    class Config:
        extra = "allow"  # Allow additional fields for flexibility
        
    @property
    def effective_file_content(self) -> Optional[str]:
        """Get file content with validation and fallback chain"""
        candidates = [self.file_content, self.fileContent]
        for candidate in candidates:
            if candidate and self._basic_validate_base64(candidate):
                return candidate
        return None
    
    @property  
    def effective_file_path(self) -> Optional[str]:
        """Get file path with validation and fallback chain"""
        candidates = [self.file_path, self.filePath]
        for candidate in candidates:
            if candidate and os.path.exists(candidate) and os.path.isfile(candidate):
                return candidate
        return None
    
    def _basic_validate_base64(self, content: str) -> bool:
        """Basic base64 validation for property usage"""
        if not content or not isinstance(content, str):
            return False
        
        # Grundlegende checks ohne imports
        if len(content.strip()) < 4:
            return False
            
        # Bereinige für pattern check
        cleaned = content.strip().replace('\n', '').replace('\r', '').replace(' ', '')
        
        # Entferne data URL prefix falls vorhanden
        if cleaned.startswith('data:') and ',' in cleaned:
            cleaned = cleaned.split(',', 1)[1]
        
        # Basic pattern check (erlaubt base64 chars + padding)
        import re
        if not re.match(r'^[A-Za-z0-9+/]*={0,2}


class OCRAgentV2:
    """
    Enhanced OCR Agent V2 mit vollständiger N8N-Integration
    
    GARANTIEN:
    - Alle extract_* Methoden geben List[Dict] zurück  
    - Einheitliche Chunk-Struktur: {"content": str, "type": str, "metadata": dict, ...}
    - Robuste Error-Behandlung
    - Streaming + In-Memory Modi
    """
    
    def __init__(self, db_url: str):
        self.db_url = db_url
        self.db_pool = None
        self.chunk_size = int(os.getenv("CHUNK_SIZE", "1000"))
        self.running = True

        # Streaming settings
        self.MAX_FILE_SIZE = 100 * 1024 * 1024  # 100 MB
        
        # Tesseract configs
        self.tesseract_config = "--oem 3 --psm 6"
        self.tesseract_config_fast = "--oem 1 --psm 6"
        
        logger.info(f"📊 OCR Agent V2 initialisiert:")
        logger.info(f"   • Chunk-Größe: {self.chunk_size}")
        logger.info(f"   • Max File Size: {self.MAX_FILE_SIZE / (1024*1024):.0f}MB")
        logger.info(f"   • DOCX Support: {HAS_DOCX}")
        logger.info(f"   • Pandas Support: {HAS_PANDAS}")
        logger.info(f"   • BeautifulSoup Support: {HAS_BS4}")

    async def connect(self):
        """Verbinde mit der PostgreSQL-Datenbank"""
        try:
            self.db_pool = await asyncpg.create_pool(self.db_url)
            logger.info("✅ Verbindung zur PostgreSQL-Datenbank hergestellt")
        except Exception as e:
            logger.error(f"❌ Datenbankverbindung fehlgeschlagen: {e}")
            raise

    async def close(self):
        """Schließe die Verbindung zur Datenbank"""
        if self.db_pool:
            await self.db_pool.close()
        logger.info("🛑 OCR Agent V2 beendet")

    async def process_queue(self):
        """Verarbeite die Warteschlange kontinuierlich"""
        while self.running:
            try:
                async with self.db_pool.acquire() as conn:
                    task = await conn.fetchrow(
                        """
                        UPDATE processing_queue 
                        SET status = 'processing', started_at = NOW()
                        WHERE id = (
                            SELECT id FROM processing_queue 
                            WHERE status = 'pending' 
                            AND task_type IN ('ocr', 'extract')
                            AND (NOT deferred OR EXTRACT(hour FROM NOW()) BETWEEN 22 AND 6)
                            ORDER BY priority DESC, created_at ASC
                            LIMIT 1
                            FOR UPDATE SKIP LOCKED
                        )
                        RETURNING *
                        """
                    )
                    if task:
                        await self.process_document(task)
                    else:
                        await asyncio.sleep(5)
            except Exception as e:
                logger.error(f"⚠️ Fehler bei Warteschlange: {e}")
                await asyncio.sleep(10)

    async def process_document(self, task: Dict):
        """
        Verarbeite ein Dokument basierend auf Metadaten
        
        Args:
            task: Database row from processing_queue
        """
        async with self.db_pool.acquire() as conn:
            try:
                doc = await conn.fetchrow(
                    "SELECT * FROM documents WHERE id = $1",
                    task["document_id"]
                )
                if not doc:
                    raise ValueError(f"🚫 Dokument {task['document_id']} nicht gefunden")

                logger.info(f"📄 Verarbeite: {doc['original_name']} (ID: {doc['id']})")

                # Extrahiere Inhalt basierend auf Dateityp - GARANTIERT List[Dict]
                chunks = await self.extract_content(doc)
                
                if not chunks:
                    logger.warning(f"⚠️ Keine Chunks erstellt für {doc['original_name']}")
                    chunks = [self._create_error_chunk("Keine Inhalte extrahiert")]

                # Validiere Chunk-Format
                chunks = self._validate_chunks(chunks)

                # Speichere Chunks in Datenbank
                await self._save_chunks_to_db(conn, doc["id"], chunks)

                # Dokumentstatus aktualisieren
                await conn.execute(
                    """UPDATE documents 
                       SET status = 'completed', processed_at = NOW()
                       WHERE id = $1""",
                    doc["id"]
                )

                # Task abschließen
                await conn.execute(
                    """UPDATE processing_queue 
                       SET status = 'completed', completed_at = NOW()
                       WHERE id = $1""",
                    task["id"]
                )

                # Enhancement-Task für LLM-Pipeline erstellen
                await conn.execute(
                    """INSERT INTO processing_queue (document_id, task_type, priority)
                       VALUES ($1, 'enhance', $2)""",
                    doc["id"], task["priority"]
                )

                logger.info(f"✅ Dokument verarbeitet: {doc['original_name']} ({len(chunks)} Chunks)")

            except Exception as e:
                logger.error(f"❌ Fehler bei Dokumentenverarbeitung: {e}")
                await self._mark_task_failed(conn, task["id"], doc["id"], str(e))

    async def extract_content(self, doc: Dict) -> List[Dict]:
        """
        Hauptmethode für Content-Extraktion mit N8N processing_mode Support
        
        Returns:
            List[Dict]: GARANTIERT List von Chunk-Dictionaries
        """
        try:
            file_type = doc["file_type"].lower()
            raw_metadata = doc.get("metadata", {})
            
            # Parse metadata zu strukturiertem Format
            metadata = self._parse_metadata_robust(raw_metadata)
            
            logger.info(f"🔍 Processing mode: {metadata.processing_mode}, file_type: {file_type}")

            # Automatische Streaming-Entscheidung für große Dateien
            if metadata.effective_file_path:
                file_size = os.path.getsize(metadata.effective_file_path)
                file_size_mb = file_size / (1024 * 1024)
                
                if file_size > self.MAX_FILE_SIZE and metadata.processing_mode != "streaming":
                    metadata.processing_mode = "streaming"
                    logger.warning(f"⚠️ Große Datei ({file_size_mb:.1f}MB) → automatischer Fallback zu 'streaming'")
                else:
                    logger.info(f"📦 Datei: {metadata.effective_file_path} ({file_size_mb:.1f} MB), Modus: {metadata.processing_mode}")

            # === STREAMING MODUS ===
            if metadata.processing_mode == "streaming" and metadata.effective_file_path:
                logger.info(f"⏳ Streaming-Modus aktiviert: {metadata.effective_file_path}")
                return await self._extract_content_streaming(metadata.effective_file_path, file_type, doc["original_name"])

            # === IN-MEMORY MODUS ===
            logger.info(f"💾 In-Memory-Modus für {file_type}")
            
            # Hole file_bytes mit robuster Fallback-Chain
            file_bytes = await self._get_file_bytes_robust(metadata)
            
            # Dispatche zu spezifischer extract-Methode
            return await self._dispatch_extraction(file_type, file_bytes, doc["original_name"])
            
        except Exception as e:
            logger.error(f"❌ Content-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Content-Extraktion fehlgeschlagen: {e}")]

    async def _get_file_bytes_robust(self, metadata: DocumentMetadata) -> bytes:
        """
        Robuste file_bytes Extraktion mit umfassender Fallback-Chain
        
        Args:
            metadata: Strukturierte Metadaten
            
        Returns:
            bytes: File content als bytes
            
        Raises:
            ValueError: Wenn keine gültige Datenquelle gefunden wird
        """
        # === PRIORITÄT 1: EFFECTIVE FILE CONTENT (bereits validiert) ===
        if metadata.effective_file_content:
            try:
                # Bereinige und dekodiere base64 content
                cleaned_content = self._clean_base64_content(metadata.effective_file_content)
                file_bytes = base64.b64decode(cleaned_content)
                
                # Zusätzliche Validierung der decoded bytes
                if not file_bytes:
                    logger.warning(f"⚠️ Base64 dekodiert zu leerem content")
                elif len(file_bytes) < 10:
                    logger.warning(f"⚠️ Base64 content zu klein: {len(file_bytes)} bytes")
                else:
                    logger.info(f"📦 Base64-Inhalt dekodiert: {len(file_bytes)} bytes")
                    return file_bytes
                    
            except Exception as e:
                logger.warning(f"⚠️ Base64-Dekodierung fehlgeschlagen trotz Validation: {e}")
        
        # === PRIORITÄT 2: EFFECTIVE FILE PATH (bereits validiert) ===
        if metadata.effective_file_path:
            try:
                with open(metadata.effective_file_path, "rb") as f:
                    file_bytes = f.read()
                
                # Zusätzliche Validierung der file bytes
                if not file_bytes:
                    logger.warning(f"⚠️ Datei ist leer: {metadata.effective_file_path}")
                else:
                    logger.info(f"📁 Datei gelesen: {len(file_bytes)} bytes von {metadata.effective_file_path}")
                    return file_bytes
                    
            except Exception as e:
                logger.warning(f"⚠️ Datei-Lesen fehlgeschlagen trotz Validation: {e}")
        
        # === FALLBACK CHAIN FÜR NICHT-EFFECTIVE CANDIDATES ===
        # (Falls effective properties None zurückgeben, aber rohe Daten vorhanden sind)
        
        fallback_file_content_candidates = [
            metadata.file_content,
            metadata.fileContent,
        ]
        
        # Suche ersten gültigen base64 content (mit vollständiger Validierung)
        for candidate in fallback_file_content_candidates:
            if candidate and self._validate_base64_content(candidate):
                try:
                    cleaned_content = self._clean_base64_content(candidate)
                    file_bytes = base64.b64decode(cleaned_content)
                    
                    if file_bytes and len(file_bytes) >= 10:
                        logger.info(f"📦 Fallback base64 erfolgreich: {len(file_bytes)} bytes")
                        return file_bytes
                        
                except Exception as e:
                    logger.warning(f"⚠️ Fallback base64 fehlgeschlagen: {e}")
                    continue
        
        fallback_file_path_candidates = [
            metadata.file_path,
            metadata.filePath,
        ]
        
        # Suche ersten gültigen file path (mit vollständiger Validierung)
        for candidate in fallback_file_path_candidates:
            if candidate and self._validate_file_path(candidate):
                try:
                    with open(candidate, "rb") as f:
                        file_bytes = f.read()
                    
                    if file_bytes:
                        logger.info(f"📁 Fallback file path erfolgreich: {len(file_bytes)} bytes")
                        return file_bytes
                        
                except Exception as e:
                    logger.warning(f"⚠️ Fallback file path fehlgeschlagen: {e}")
                    continue
        
        # === ERROR CASE ===
        # Sammle alle candidates für error reporting
        all_file_content_candidates = [
            metadata.effective_file_content,
            metadata.file_content,
            metadata.fileContent,
        ]
        
        all_file_path_candidates = [
            metadata.effective_file_path,
            metadata.file_path,
            metadata.filePath,
        ]
        
        error_details = {
            "effective_file_content": bool(metadata.effective_file_content),
            "effective_file_path": bool(metadata.effective_file_path),
            "all_file_content_candidates": [
                {
                    "exists": bool(c), 
                    "valid": self._validate_base64_content(c) if c else False,
                    "length": len(c) if c else 0,
                    "is_effective": c == metadata.effective_file_content
                } for c in all_file_content_candidates
            ],
            "all_file_path_candidates": [
                {
                    "path": c,
                    "exists": os.path.exists(c) if c else False,
                    "valid": self._validate_file_path(c) if c else False,
                    "is_effective": c == metadata.effective_file_path
                } for c in all_file_path_candidates
            ],
            "metadata_keys": list(metadata.dict(exclude_none=True).keys()),
            "processing_mode": metadata.processing_mode
        }
        
        raise ValueError(f"Keine gültige Datenquelle gefunden. Details: {error_details}")

    def _clean_base64_content(self, content: str) -> str:
        """
        Bereinige base64 content von häufigen Problemen
        
        Args:
            content: Raw base64 string
            
        Returns:
            str: Cleaned base64 string
        """
        if not content:
            return ""
            
        # Entferne whitespace, newlines, carriage returns
        cleaned = content.strip().replace('\n', '').replace('\r', '').replace(' ', '').replace('\t', '')
        
        # Entferne data URL prefix falls vorhanden (data:image/png;base64,...)
        if cleaned.startswith('data:'):
            if ',' in cleaned:
                cleaned = cleaned.split(',', 1)[1]
        
        # Validiere base64 format (grundlegend)
        import re
        if not re.match(r'^[A-Za-z0-9+/]*={0,2}

    async def _dispatch_extraction(self, file_type: str, file_bytes: bytes, filename: str) -> List[Dict]:
        """
        Dispatche zu spezifischer Extraktionsmethode
        
        Returns:
            List[Dict]: GARANTIERT List von Chunk-Dictionaries
        """
        try:
            if file_type == "pdf":
                return await self.extract_pdf(file_bytes, filename)
            elif file_type in ["png", "jpg", "jpeg", "tiff", "tif"]:
                return await self.extract_image(file_bytes)
            elif file_type in ["txt", "md", "log"]:
                content = file_bytes.decode("utf-8", errors="ignore")
                return self.extract_text(content, file_type)
            elif file_type in ["docx", "doc"]:
                return await self.extract_docx(file_bytes)
            elif file_type in ["xlsx", "xls", "csv"]:
                return await self.extract_spreadsheet(file_bytes, file_type)
            elif file_type == "eml":
                return await self.extract_email(file_bytes)
            elif file_type == "html":
                content = file_bytes.decode("utf-8", errors="ignore")
                return self.extract_html(content)
            elif file_type == "zip":
                return await self.extract_archive(file_bytes)
            else:
                # Fallback zu text
                content = file_bytes.decode("utf-8", errors="ignore")
                return self.extract_text(content, file_type)
                
        except Exception as e:
            logger.error(f"❌ Extraktion für {file_type} fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Extraktion für {file_type} fehlgeschlagen: {e}")]

    # === STREAMING METHODEN ===
    
    async def _extract_content_streaming(self, file_path: str, file_type: str, filename: str) -> List[Dict]:
        """Streaming-Extraktion für große Dateien - GARANTIERT List[Dict]"""
        try:
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            logger.info(f"🚀 Streaming für {filename}: {file_size_mb:.1f}MB")
            
            if file_type == "pdf":
                return await self._stream_large_pdf(file_path)
            elif file_type in ["txt", "md", "log"]:
                return await self._stream_large_text(file_path)
            elif file_type in ["docx", "doc"] and HAS_DOCX:
                return await self._stream_large_docx(file_path)
            else:
                # Fallback: normales Laden als Text
                logger.warning(f"⚠️ Kein Streaming für {file_type} - fallback zu normal")
                with open(file_path, "rb") as f:
                    file_bytes = f.read()
                content = file_bytes.decode("utf-8", errors="ignore")
                return self.extract_text(content, file_type)
                
        except Exception as e:
            logger.error(f"❌ Streaming fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Streaming fehlgeschlagen: {e}")]

    async def _stream_large_pdf(self, file_path: str) -> List[Dict]:
        """Verarbeite PDF seitenweise - GARANTIERT List[Dict]"""
        chunks = []
        try:
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            dpi = 150 if file_size_mb > 50 else 300
            tesseract_config = self.tesseract_config_fast if file_size_mb > 100 else self.tesseract_config
            
            logger.info(f"📄 PDF-Streaming: {file_path} ({file_size_mb:.1f} MB, DPI: {dpi})")

            images = pdf2image.convert_from_path(file_path, dpi=dpi)
            total_pages = len(images)

            for page_num, image in enumerate(images, 1):
                if total_pages > 50 and page_num % 10 == 0:
                    logger.info(f"📄 Fortschritt: {page_num}/{total_pages} ({page_num/total_pages*100:.1f}%)")
                
                text = pytesseract.image_to_string(
                    image,
                    lang=os.getenv("TESSERACT_LANG", "deu+eng"),
                    config=tesseract_config
                )

                if text.strip():
                    pdf_chunk_size = int(self.chunk_size * 1.5) if file_size_mb > 100 else self.chunk_size
                    page_text_chunks = self.split_text(text, pdf_chunk_size)
                    
                    for chunk_text in page_text_chunks:
                        chunks.append({
                            "content": chunk_text,
                            "type": "text",
                            "page": page_num,
                            "metadata": {
                                "extraction_method": "pdf_streaming",
                                "dpi": dpi,
                                "chunk_size_used": pdf_chunk_size
                            }
                        })

                del image
                if page_num % 5 == 0:
                    gc.collect()
                    
        except Exception as e:
            logger.error(f"❌ PDF-Streaming fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"PDF-Streaming fehlgeschlagen: {e}")]
            
        logger.info(f"✅ PDF verarbeitet: {len(chunks)} Chunks aus {total_pages} Seiten")
        return chunks

    async def _stream_large_text(self, file_path: str) -> List[Dict]:
        """Verarbeite große Textdateien line-by-line - GARANTIERT List[Dict]"""
        chunks = []
        try:
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            dynamic_chunk_size = min(self.chunk_size * 2, 2000) if file_size_mb > 100 else self.chunk_size
            
            logger.info(f"📦 Text-Streaming: {file_path} ({file_size_mb:.1f} MB), Chunk-Größe: {dynamic_chunk_size}")

            current_chunk = []
            current_size = 0
            line_count = 0
            
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                for line in f:
                    line = line.strip() + "\n"
                    if not line.strip():
                        continue
                        
                    current_chunk.append(line)
                    current_size += len(line)
                    line_count += 1

                    if current_size >= dynamic_chunk_size:
                        chunks.append({
                            "content": "".join(current_chunk),
                            "type": "text",
                            "metadata": {
                                "extraction_method": "text_streaming",
                                "lines_in_chunk": len(current_chunk),
                                "chunk_size_used": dynamic_chunk_size
                            }
                        })
                        current_chunk = []
                        current_size = 0

                        if len(chunks) % 10 == 0:
                            logger.info(f"📊 {len(chunks)} Chunks verarbeitet ({line_count} Zeilen)")
                            
                # Letzter Chunk
                if current_chunk:
                    chunks.append({
                        "content": "".join(current_chunk),
                        "type": "text",
                        "metadata": {
                            "extraction_method": "text_streaming",
                            "lines_in_chunk": len(current_chunk),
                            "chunk_size_used": dynamic_chunk_size,
                            "final_chunk": True
                        }
                    })
                    
        except Exception as e:
            logger.error(f"❌ Text-Streaming fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Text-Streaming fehlgeschlagen: {e}")]
            
        logger.info(f"✅ Textdatei verarbeitet: {len(chunks)} Chunks")
        return chunks

    async def _stream_large_docx(self, file_path: str) -> List[Dict]:
        """Verarbeite große DOCX-Dateien - GARANTIERT List[Dict]"""
        chunks = []
        try:
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            logger.info(f"📄 DOCX-Streaming: {file_path} ({file_size_mb:.1f} MB)")
            
            if not HAS_DOCX:
                return [self._create_error_chunk("DOCX-Support nicht verfügbar")]
            
            doc = DocxDocument(file_path)
            full_text = "\n".join([para.text for para in doc.paragraphs])
            
            docx_chunk_size = int(self.chunk_size * 1.3)
            text_chunks = self.split_text(full_text, docx_chunk_size)
            
            for chunk_text in text_chunks:
                chunks.append({
                    "content": chunk_text,
                    "type": "text",
                    "metadata": {
                        "extraction_method": "docx_streaming",
                        "chunk_size_used": docx_chunk_size
                    }
                })
                
            # Extrahiere Tabellen
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                    
                if table_data:
                    chunks.append({
                        "content": json.dumps(table_data),
                        "type": "table",
                        "metadata": {
                            "extraction_method": "docx_table",
                            "table_index": table_idx
                        }
                    })
                    
        except Exception as e:
            logger.error(f"❌ DOCX-Streaming fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"DOCX-Streaming fehlgeschlagen: {e}")]
            
        logger.info(f"✅ DOCX verarbeitet: {len(chunks)} Chunks")
        return chunks

    # === IN-MEMORY EXTRAKTIONSMETHODEN ===
    
    async def extract_pdf(self, pdf_bytes: bytes, filename: str) -> List[Dict]:
        """Extrahiere Text aus PDF mit OCR (In-Memory) - GARANTIERT List[Dict]"""
        chunks = []
        try:
            images = pdf2image.convert_from_bytes(pdf_bytes, dpi=300)
            logger.info(f"📄 PDF In-Memory: {filename} ({len(images)} Seiten)")
            
            for page_num, image in enumerate(images, 1):
                text = pytesseract.image_to_string(
                    image,
                    lang=os.getenv("TESSERACT_LANG", "deu+eng"),
                    config=self.tesseract_config
                )
                
                if text.strip():
                    page_chunks = self.split_text(text)
                    for chunk_text in page_chunks:
                        chunks.append({
                            "content": chunk_text,
                            "type": "text",
                            "page": page_num,
                            "metadata": {
                                "extraction_method": "pdf_in_memory",
                                "chunk_size_used": self.chunk_size
                            }
                        })
                        
                del image
                if page_num % 5 == 0:
                    gc.collect()
                    
        except Exception as e:
            logger.error(f"❌ PDF-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"PDF-Extraktion fehlgeschlagen: {e}")]
            
        return chunks if chunks else [self._create_error_chunk("Kein Text aus PDF extrahiert")]

    async def extract_image(self, image_bytes: bytes) -> List[Dict]:
        """Extrahiere Text aus einem Bild - GARANTIERT List[Dict]"""
        try:
            image = Image.open(io.BytesIO(image_bytes))
            text = pytesseract.image_to_string(
                image,
                lang=os.getenv("TESSERACT_LANG", "deu+eng"),
                config=self.tesseract_config
            )
            
            if not text.strip():
                return [{"content": "[Bild ohne erkannten Text]", "type": "text", "metadata": {"extraction_method": "image_ocr"}}]
                
            chunks = self.split_text(text)
            return [{
                "content": chunk_text,
                "type": "text",
                "page": 1,
                "metadata": {
                    "extraction_method": "image_ocr",
                    "chunk_size_used": self.chunk_size
                }
            } for chunk_text in chunks]
            
        except Exception as e:
            logger.error(f"❌ Bild-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Bild-Extraktion fehlgeschlagen: {e}")]

    async def extract_docx(self, docx_bytes: bytes) -> List[Dict]:
        """Extrahiere Text aus DOCX (In-Memory) - GARANTIERT List[Dict]"""
        if not HAS_DOCX:
            return [self._create_error_chunk("DOCX-Support nicht verfügbar (python-docx nicht installiert)")]
            
        try:
            doc = DocxDocument(io.BytesIO(docx_bytes))
            full_text = "\n".join([para.text for para in doc.paragraphs])
            
            docx_chunk_size = int(self.chunk_size * 1.3)
            chunks = self.split_text(full_text, docx_chunk_size)
            
            result = [{
                "content": chunk_text,
                "type": "text",
                "metadata": {
                    "extraction_method": "docx_in_memory",
                    "chunk_size_used": docx_chunk_size
                }
            } for chunk_text in chunks]
            
            # Extrahiere Tabellen
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                    
                if table_data:
                    result.append({
                        "content": json.dumps(table_data),
                        "type": "table",
                        "metadata": {
                            "extraction_method": "docx_table",
                            "table_index": table_idx
                        }
                    })
                    
            return result if result else [self._create_error_chunk("Kein Inhalt aus DOCX extrahiert")]
            
        except Exception as e:
            logger.error(f"❌ DOCX-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"DOCX-Extraktion fehlgeschlagen: {e}")]

    def extract_text(self, text: str, file_type: str = "txt") -> List[Dict]:
        """Extrahiere Text mit adaptiver Chunk-Größe - GARANTIERT List[Dict]"""
        if not text.strip():
            return [{"content": "[Leerer Text]", "type": "text", "metadata": {"extraction_method": "text_in_memory"}}]
            
        try:
            # Adaptive Chunk-Größen basierend auf Dateityp
            if file_type in ['csv', 'tsv']:
                chunk_size = self.chunk_size // 2
            elif file_type in ['md', 'txt']:
                chunk_size = int(self.chunk_size * 1.2)
            else:
                chunk_size = self.chunk_size
                
            chunks = self.split_text(text, chunk_size)
            return [{
                "content": chunk_text,
                "type": "text",
                "metadata": {
                    "extraction_method": "text_in_memory",
                    "chunk_size_used": chunk_size,
                    "file_type": file_type
                }
            } for chunk_text in chunks]
            
        except Exception as e:
            logger.error(f"❌ Text-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Text-Extraktion fehlgeschlagen: {e}")]

    async def extract_spreadsheet(self, file_bytes: bytes, file_type: str) -> List[Dict]:
        """Extrahiere Tabellen aus Excel/CSV - GARANTIERT List[Dict]"""
        if not HAS_PANDAS:
            return [self._create_error_chunk("Tabellen-Support nicht verfügbar (pandas nicht installiert)")]
            
        try:
            if file_type == "csv":
                df = pd.read_csv(io.BytesIO(file_bytes))
            else:
                df = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
                
            chunks = []
            if isinstance(df, dict):
                for sheet_name, sheet_df in df.items():
                    chunks.extend(self._process_dataframe(sheet_df, {"sheet": sheet_name}))
            else:
                chunks.extend(self._process_dataframe(df))
                
            return chunks if chunks else [self._create_error_chunk("Keine Tabellendaten extrahiert")]
            
        except Exception as e:
            logger.error(f"❌ Tabelle fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Tabellen-Extraktion fehlgeschlagen: {e}")]

    def _process_dataframe(self, df: "pd.DataFrame", metadata: dict = None) -> List[Dict]:
        """Verarbeite DataFrame in Chunks - GARANTIERT List[Dict]"""
        metadata = metadata or {}
        chunks = []
        
        try:
            for i in range(0, len(df), 50):
                chunk_df = df.iloc[i:i + 50]
                chunks.append({
                    "content": chunk_df.to_json(orient="records"),
                    "type": "table",
                    "metadata": {
                        **metadata, 
                        "row_start": i,
                        "extraction_method": "pandas"
                    }
                })
        except Exception as e:
            logger.error(f"❌ DataFrame-Verarbeitung fehlgeschlagen: {e}")
            chunks.append(self._create_error_chunk(f"DataFrame-Fehler: {e}"))
            
        return chunks

    async def extract_email(self, eml_bytes: bytes) -> List[Dict]:
        """Extrahiere E-Mail-Metadaten und Inhalt - GARANTIERT List[Dict]"""
        try:
            msg = BytesParser(policy=policy.default).parsebytes(eml_bytes)
            chunks = [{
                "content": json.dumps({
                    "from": msg["From"],
                    "to": msg["To"],
                    "subject": msg["Subject"],
                    "date": msg["Date"],
                }),
                "type": "metadata",
                "metadata": {"extraction_method": "email_header"}
            }]
            
            body = msg.get_body(preferencelist=("plain", "html"))
            if body:
                content = body.get_content()
                if content and content.strip():
                    # FIX: Korrekte Verarbeitung der split_text Rückgabe
                    text_chunks = self.split_text(content.strip())
                    chunks.extend([{
                        "content": chunk_text, 
                        "type": "text",
                        "metadata": {"extraction_method": "email_body"}
                    } for chunk_text in text_chunks])
                    
            return chunks
            
        except Exception as e:
            logger.error(f"❌ E-Mail-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"E-Mail-Extraktion fehlgeschlagen: {e}")]

    def extract_html(self, html: str) -> List[Dict]:
        """Extrahiere Text aus HTML - GARANTIERT List[Dict]"""
        try:
            if HAS_BS4:
                soup = BeautifulSoup(html, "html.parser")
                text = soup.get_text(separator="\n", strip=True)
            else:
                import re
                text = re.sub("<[^<]+?>", "", html)
                
            return self.extract_text(text, "html")
            
        except Exception as e:
            logger.error(f"❌ HTML-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"HTML-Extraktion fehlgeschlagen: {e}")]

    async def extract_archive(self, zip_bytes: bytes) -> List[Dict]:
        """Extrahiere Dateiliste aus ZIP - GARANTIERT List[Dict]"""
        try:
            with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
                files = zf.namelist()
                return [{
                    "content": f"Archive enthält: {name}",
                    "type": "metadata",
                    "metadata": {
                        "archived_file": name,
                        "extraction_method": "zip_listing"
                    }
                } for name in files if not name.endswith("/")]
                
        except Exception as e:
            logger.error(f"❌ Archiv-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Archiv-Extraktion fehlgeschlagen: {e}")]

    # === UTILITY METHODEN ===
    
    def split_text(self, text: str, chunk_size: int = None) -> List[str]:
        """
        Teile Text in Chunks mit dynamischer Größe
        
        Returns:
            List[str]: Liste von Text-Chunks (nicht Dict!)
        """
        chunk_size = chunk_size or self.chunk_size
        words = text.split()
        chunks = []
        current_chunk = []
        current_size = 0

        for word in words:
            current_chunk.append(word)
            current_size += len(word) + 1
            
            if current_size >= chunk_size:
                chunks.append(" ".join(current_chunk))
                current_chunk = []
                current_size = 0

        if current_chunk:
            chunks.append(" ".join(current_chunk))

        return chunks or [""]

    def _create_error_chunk(self, error_message: str) -> Dict:
        """Erstelle standardisierten Error-Chunk"""
        return {
            "content": f"[FEHLER: {error_message}]",
            "type": "error",
            "metadata": {
                "extraction_method": "error",
                "error_message": error_message,
                "timestamp": datetime.now().isoformat()
            }
        }

    def _validate_chunks(self, chunks: List[Dict]) -> List[Dict]:
        """Validiere und normalisiere Chunk-Format"""
        validated = []
        for i, chunk in enumerate(chunks):
            if not isinstance(chunk, dict):
                logger.warning(f"⚠️ Chunk {i} ist kein Dict: {type(chunk)}")
                chunk = {"content": str(chunk), "type": "text", "metadata": {}}
                
            # Stelle sicher, dass required fields vorhanden sind
            if "content" not in chunk:
                chunk["content"] = "[Leerer Chunk]"
            if "type" not in chunk:
                chunk["type"] = "text"
            if "metadata" not in chunk:
                chunk["metadata"] = {}
                
            validated.append(chunk)
            
        return validated

    def _parse_metadata(self, metadata: Union[str, dict]) -> dict:
        """DEPRECATED: Legacy method - use _parse_metadata_robust instead"""
        robust_metadata = self._parse_metadata_robust(metadata)
        return robust_metadata.dict(exclude_none=True)

    async def _save_chunks_to_db(self, conn, document_id: str, chunks: List[Dict]):
        """Speichere Chunks in Datenbank"""
        for idx, chunk in enumerate(chunks):
            await conn.execute(
                """
                INSERT INTO chunks (
                    document_id, chunk_index, content, content_type,
                    page_number, position, metadata
                ) VALUES ($1, $2, $3, $4, $5, $6, $7)
                ON CONFLICT (document_id, chunk_index) DO UPDATE
                SET content = EXCLUDED.content,
                    content_type = EXCLUDED.content_type
                """,
                document_id,
                idx,
                chunk["content"],
                chunk.get("type", "text"),
                chunk.get("page"),
                chunk.get("position"),
                json.dumps(chunk.get("metadata", {}))
            )

    async def _mark_task_failed(self, conn, task_id: str, doc_id: str, error_message: str):
        """Markiere Task als fehlgeschlagen"""
        await conn.execute(
            """UPDATE processing_queue 
               SET status = 'failed', error_message = $2, completed_at = NOW()
               WHERE id = $1""",
            task_id, error_message
        )
        await conn.execute(
            """UPDATE documents SET status = 'failed' WHERE id = $1""",
            doc_id
        )

    async def run(self):
        """Starte den OCR-Agent"""
        await self.connect()
        logger.info("🧠 Enhanced OCR Agent V2 gestartet")
        
        try:
            await self.process_queue()
        except KeyboardInterrupt:
            logger.info("🛑 OCR Agent wird beendet (KeyboardInterrupt)")
        except Exception as e:
            logger.error(f"💥 Unerwarteter Fehler: {e}")
            raise
        finally:
            await self.close()


def main():
    """Haupt-Funktion"""
    db_url = os.getenv(
        "DATABASE_URL",
        "postgresql://semanticuser:semantic2024@postgres:5432/semantic_doc_finder"
    )
    
    logger.info(f"🚀 Starte OCR Agent mit DB: {db_url.split('@')[1] if '@' in db_url else 'localhost'}")
    
    agent = OCRAgentV2(db_url)
    asyncio.run(agent.run())


if __name__ == "__main__":
    main(), cleaned):
            logger.warning(f"⚠️ Potentiell ungültiger base64 content (erste 50 chars): {cleaned[:50]}")
        
        return cleaned

    def _validate_file_path(self, file_path: str) -> bool:
        """
        Umfassende Validierung eines file paths
        
        Args:
            file_path: Path to validate
            
        Returns:
            bool: True wenn file path gültig und lesbar ist
        """
        if not file_path or not isinstance(file_path, str):
            return False
            
        try:
            # Prüfe ob Datei existiert
            if not os.path.exists(file_path):
                logger.debug(f"File path existiert nicht: {file_path}")
                return False
            
            # Prüfe ob es eine Datei ist (nicht directory)
            if not os.path.isfile(file_path):
                logger.debug(f"Path ist keine Datei: {file_path}")
                return False
                
            # Prüfe Dateigröße (nicht 0 und nicht zu groß)
            file_size = os.path.getsize(file_path)
            if file_size == 0:
                logger.debug(f"Datei ist leer: {file_path}")
                return False
                
            # Prüfe maximale Dateigröße (2GB limit)
            MAX_FILE_SIZE_ABSOLUTE = 2 * 1024 * 1024 * 1024  # 2GB
            if file_size > MAX_FILE_SIZE_ABSOLUTE:
                logger.warning(f"⚠️ Datei zu groß: {file_size} bytes (Max: {MAX_FILE_SIZE_ABSOLUTE})")
                return False
            
            # Prüfe Lese-Berechtigung
            if not os.access(file_path, os.R_OK):
                logger.debug(f"Keine Lese-Berechtigung für: {file_path}")
                return False
                
            logger.debug(f"✅ File path validiert: {file_path} ({file_size} bytes)")
            return True
            
        except Exception as e:
            logger.debug(f"File path validation fehlgeschlagen: {e}")
            return False

    def _validate_base64_content(self, content: str) -> bool:
        """
        Umfassende Validierung von base64 content
        
        Args:
            content: Base64 string to validate
            
        Returns:
            bool: True wenn content gültiger base64 ist
        """
        if not content or not isinstance(content, str):
            return False
            
        try:
            # Prüfe minimale Länge
            if len(content.strip()) < 4:
                logger.debug(f"Base64 content zu kurz: {len(content)} chars")
                return False
                
            # Prüfe maximale Länge (500MB als base64 ≈ 666MB string)
            MAX_BASE64_LENGTH = 700 * 1024 * 1024  # 700MB string
            if len(content) > MAX_BASE64_LENGTH:
                logger.warning(f"⚠️ Base64 content zu groß: {len(content)} chars")
                return False
            
            # Bereinige content für test decode
            cleaned_content = self._clean_base64_content(content)
            
            # Prüfe base64 format pattern
            if not re.match(r'^[A-Za-z0-9+/]*={0,2}

    def _parse_metadata_robust(self, raw_metadata: Union[str, dict]) -> DocumentMetadata:
        """
        Robustes Parsing von Metadaten zu strukturiertem Format
        
        Args:
            raw_metadata: Raw metadata von database
            
        Returns:
            DocumentMetadata: Strukturierte und validierte Metadaten
        """
        try:
            # Parse zu dict falls string
            if isinstance(raw_metadata, str):
                try:
                    metadata_dict = json.loads(raw_metadata)
                except json.JSONDecodeError as e:
                    logger.warning(f"⚠️ JSON-Parsing fehlgeschlagen: {e}")
                    metadata_dict = {}
            elif isinstance(raw_metadata, dict):
                metadata_dict = raw_metadata
            else:
                logger.warning(f"⚠️ Unerwarteter metadata type: {type(raw_metadata)}")
                metadata_dict = {}
            
            # Erstelle DocumentMetadata mit validation
            metadata = DocumentMetadata(**metadata_dict)
            
            # Additional validation/normalization
            if not metadata.effective_file_content and not metadata.effective_file_path:
                logger.warning("⚠️ Keine file_content oder file_path in metadata gefunden")
            
            return metadata
            
        except Exception as e:
            logger.error(f"❌ Metadata-Parsing fehlgeschlagen: {e}")
            # Fallback zu default metadata
            return DocumentMetadata(source="metadata_parse_error")

    async def _get_file_bytes(self, file_content: Optional[str], file_path: Optional[str]) -> bytes:
        """DEPRECATED: Legacy method - use _get_file_bytes_robust instead"""
        # Erstelle legacy DocumentMetadata für compatibility
        legacy_metadata = DocumentMetadata(
            file_content=file_content,
            file_path=file_path
        )
        return await self._get_file_bytes_robust(legacy_metadata)

    async def _dispatch_extraction(self, file_type: str, file_bytes: bytes, filename: str) -> List[Dict]:
        """
        Dispatche zu spezifischer Extraktionsmethode
        
        Returns:
            List[Dict]: GARANTIERT List von Chunk-Dictionaries
        """
        try:
            if file_type == "pdf":
                return await self.extract_pdf(file_bytes, filename)
            elif file_type in ["png", "jpg", "jpeg", "tiff", "tif"]:
                return await self.extract_image(file_bytes)
            elif file_type in ["txt", "md", "log"]:
                content = file_bytes.decode("utf-8", errors="ignore")
                return self.extract_text(content, file_type)
            elif file_type in ["docx", "doc"]:
                return await self.extract_docx(file_bytes)
            elif file_type in ["xlsx", "xls", "csv"]:
                return await self.extract_spreadsheet(file_bytes, file_type)
            elif file_type == "eml":
                return await self.extract_email(file_bytes)
            elif file_type == "html":
                content = file_bytes.decode("utf-8", errors="ignore")
                return self.extract_html(content)
            elif file_type == "zip":
                return await self.extract_archive(file_bytes)
            else:
                # Fallback zu text
                content = file_bytes.decode("utf-8", errors="ignore")
                return self.extract_text(content, file_type)
                
        except Exception as e:
            logger.error(f"❌ Extraktion für {file_type} fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Extraktion für {file_type} fehlgeschlagen: {e}")]

    # === STREAMING METHODEN ===
    
    async def _extract_content_streaming(self, file_path: str, file_type: str, filename: str) -> List[Dict]:
        """Streaming-Extraktion für große Dateien - GARANTIERT List[Dict]"""
        try:
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            logger.info(f"🚀 Streaming für {filename}: {file_size_mb:.1f}MB")
            
            if file_type == "pdf":
                return await self._stream_large_pdf(file_path)
            elif file_type in ["txt", "md", "log"]:
                return await self._stream_large_text(file_path)
            elif file_type in ["docx", "doc"] and HAS_DOCX:
                return await self._stream_large_docx(file_path)
            else:
                # Fallback: normales Laden als Text
                logger.warning(f"⚠️ Kein Streaming für {file_type} - fallback zu normal")
                with open(file_path, "rb") as f:
                    file_bytes = f.read()
                content = file_bytes.decode("utf-8", errors="ignore")
                return self.extract_text(content, file_type)
                
        except Exception as e:
            logger.error(f"❌ Streaming fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Streaming fehlgeschlagen: {e}")]

    async def _stream_large_pdf(self, file_path: str) -> List[Dict]:
        """Verarbeite PDF seitenweise - GARANTIERT List[Dict]"""
        chunks = []
        try:
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            dpi = 150 if file_size_mb > 50 else 300
            tesseract_config = self.tesseract_config_fast if file_size_mb > 100 else self.tesseract_config
            
            logger.info(f"📄 PDF-Streaming: {file_path} ({file_size_mb:.1f} MB, DPI: {dpi})")

            images = pdf2image.convert_from_path(file_path, dpi=dpi)
            total_pages = len(images)

            for page_num, image in enumerate(images, 1):
                if total_pages > 50 and page_num % 10 == 0:
                    logger.info(f"📄 Fortschritt: {page_num}/{total_pages} ({page_num/total_pages*100:.1f}%)")
                
                text = pytesseract.image_to_string(
                    image,
                    lang=os.getenv("TESSERACT_LANG", "deu+eng"),
                    config=tesseract_config
                )

                if text.strip():
                    pdf_chunk_size = int(self.chunk_size * 1.5) if file_size_mb > 100 else self.chunk_size
                    page_text_chunks = self.split_text(text, pdf_chunk_size)
                    
                    for chunk_text in page_text_chunks:
                        chunks.append({
                            "content": chunk_text,
                            "type": "text",
                            "page": page_num,
                            "metadata": {
                                "extraction_method": "pdf_streaming",
                                "dpi": dpi,
                                "chunk_size_used": pdf_chunk_size
                            }
                        })

                del image
                if page_num % 5 == 0:
                    gc.collect()
                    
        except Exception as e:
            logger.error(f"❌ PDF-Streaming fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"PDF-Streaming fehlgeschlagen: {e}")]
            
        logger.info(f"✅ PDF verarbeitet: {len(chunks)} Chunks aus {total_pages} Seiten")
        return chunks

    async def _stream_large_text(self, file_path: str) -> List[Dict]:
        """Verarbeite große Textdateien line-by-line - GARANTIERT List[Dict]"""
        chunks = []
        try:
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            dynamic_chunk_size = min(self.chunk_size * 2, 2000) if file_size_mb > 100 else self.chunk_size
            
            logger.info(f"📦 Text-Streaming: {file_path} ({file_size_mb:.1f} MB), Chunk-Größe: {dynamic_chunk_size}")

            current_chunk = []
            current_size = 0
            line_count = 0
            
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                for line in f:
                    line = line.strip() + "\n"
                    if not line.strip():
                        continue
                        
                    current_chunk.append(line)
                    current_size += len(line)
                    line_count += 1

                    if current_size >= dynamic_chunk_size:
                        chunks.append({
                            "content": "".join(current_chunk),
                            "type": "text",
                            "metadata": {
                                "extraction_method": "text_streaming",
                                "lines_in_chunk": len(current_chunk),
                                "chunk_size_used": dynamic_chunk_size
                            }
                        })
                        current_chunk = []
                        current_size = 0

                        if len(chunks) % 10 == 0:
                            logger.info(f"📊 {len(chunks)} Chunks verarbeitet ({line_count} Zeilen)")
                            
                # Letzter Chunk
                if current_chunk:
                    chunks.append({
                        "content": "".join(current_chunk),
                        "type": "text",
                        "metadata": {
                            "extraction_method": "text_streaming",
                            "lines_in_chunk": len(current_chunk),
                            "chunk_size_used": dynamic_chunk_size,
                            "final_chunk": True
                        }
                    })
                    
        except Exception as e:
            logger.error(f"❌ Text-Streaming fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Text-Streaming fehlgeschlagen: {e}")]
            
        logger.info(f"✅ Textdatei verarbeitet: {len(chunks)} Chunks")
        return chunks

    async def _stream_large_docx(self, file_path: str) -> List[Dict]:
        """Verarbeite große DOCX-Dateien - GARANTIERT List[Dict]"""
        chunks = []
        try:
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            logger.info(f"📄 DOCX-Streaming: {file_path} ({file_size_mb:.1f} MB)")
            
            if not HAS_DOCX:
                return [self._create_error_chunk("DOCX-Support nicht verfügbar")]
            
            doc = DocxDocument(file_path)
            full_text = "\n".join([para.text for para in doc.paragraphs])
            
            docx_chunk_size = int(self.chunk_size * 1.3)
            text_chunks = self.split_text(full_text, docx_chunk_size)
            
            for chunk_text in text_chunks:
                chunks.append({
                    "content": chunk_text,
                    "type": "text",
                    "metadata": {
                        "extraction_method": "docx_streaming",
                        "chunk_size_used": docx_chunk_size
                    }
                })
                
            # Extrahiere Tabellen
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                    
                if table_data:
                    chunks.append({
                        "content": json.dumps(table_data),
                        "type": "table",
                        "metadata": {
                            "extraction_method": "docx_table",
                            "table_index": table_idx
                        }
                    })
                    
        except Exception as e:
            logger.error(f"❌ DOCX-Streaming fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"DOCX-Streaming fehlgeschlagen: {e}")]
            
        logger.info(f"✅ DOCX verarbeitet: {len(chunks)} Chunks")
        return chunks

    # === IN-MEMORY EXTRAKTIONSMETHODEN ===
    
    async def extract_pdf(self, pdf_bytes: bytes, filename: str) -> List[Dict]:
        """Extrahiere Text aus PDF mit OCR (In-Memory) - GARANTIERT List[Dict]"""
        chunks = []
        try:
            images = pdf2image.convert_from_bytes(pdf_bytes, dpi=300)
            logger.info(f"📄 PDF In-Memory: {filename} ({len(images)} Seiten)")
            
            for page_num, image in enumerate(images, 1):
                text = pytesseract.image_to_string(
                    image,
                    lang=os.getenv("TESSERACT_LANG", "deu+eng"),
                    config=self.tesseract_config
                )
                
                if text.strip():
                    page_chunks = self.split_text(text)
                    for chunk_text in page_chunks:
                        chunks.append({
                            "content": chunk_text,
                            "type": "text",
                            "page": page_num,
                            "metadata": {
                                "extraction_method": "pdf_in_memory",
                                "chunk_size_used": self.chunk_size
                            }
                        })
                        
                del image
                if page_num % 5 == 0:
                    gc.collect()
                    
        except Exception as e:
            logger.error(f"❌ PDF-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"PDF-Extraktion fehlgeschlagen: {e}")]
            
        return chunks if chunks else [self._create_error_chunk("Kein Text aus PDF extrahiert")]

    async def extract_image(self, image_bytes: bytes) -> List[Dict]:
        """Extrahiere Text aus einem Bild - GARANTIERT List[Dict]"""
        try:
            image = Image.open(io.BytesIO(image_bytes))
            text = pytesseract.image_to_string(
                image,
                lang=os.getenv("TESSERACT_LANG", "deu+eng"),
                config=self.tesseract_config
            )
            
            if not text.strip():
                return [{"content": "[Bild ohne erkannten Text]", "type": "text", "metadata": {"extraction_method": "image_ocr"}}]
                
            chunks = self.split_text(text)
            return [{
                "content": chunk_text,
                "type": "text",
                "page": 1,
                "metadata": {
                    "extraction_method": "image_ocr",
                    "chunk_size_used": self.chunk_size
                }
            } for chunk_text in chunks]
            
        except Exception as e:
            logger.error(f"❌ Bild-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Bild-Extraktion fehlgeschlagen: {e}")]

    async def extract_docx(self, docx_bytes: bytes) -> List[Dict]:
        """Extrahiere Text aus DOCX (In-Memory) - GARANTIERT List[Dict]"""
        if not HAS_DOCX:
            return [self._create_error_chunk("DOCX-Support nicht verfügbar (python-docx nicht installiert)")]
            
        try:
            doc = DocxDocument(io.BytesIO(docx_bytes))
            full_text = "\n".join([para.text for para in doc.paragraphs])
            
            docx_chunk_size = int(self.chunk_size * 1.3)
            chunks = self.split_text(full_text, docx_chunk_size)
            
            result = [{
                "content": chunk_text,
                "type": "text",
                "metadata": {
                    "extraction_method": "docx_in_memory",
                    "chunk_size_used": docx_chunk_size
                }
            } for chunk_text in chunks]
            
            # Extrahiere Tabellen
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                    
                if table_data:
                    result.append({
                        "content": json.dumps(table_data),
                        "type": "table",
                        "metadata": {
                            "extraction_method": "docx_table",
                            "table_index": table_idx
                        }
                    })
                    
            return result if result else [self._create_error_chunk("Kein Inhalt aus DOCX extrahiert")]
            
        except Exception as e:
            logger.error(f"❌ DOCX-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"DOCX-Extraktion fehlgeschlagen: {e}")]

    def extract_text(self, text: str, file_type: str = "txt") -> List[Dict]:
        """Extrahiere Text mit adaptiver Chunk-Größe - GARANTIERT List[Dict]"""
        if not text.strip():
            return [{"content": "[Leerer Text]", "type": "text", "metadata": {"extraction_method": "text_in_memory"}}]
            
        try:
            # Adaptive Chunk-Größen basierend auf Dateityp
            if file_type in ['csv', 'tsv']:
                chunk_size = self.chunk_size // 2
            elif file_type in ['md', 'txt']:
                chunk_size = int(self.chunk_size * 1.2)
            else:
                chunk_size = self.chunk_size
                
            chunks = self.split_text(text, chunk_size)
            return [{
                "content": chunk_text,
                "type": "text",
                "metadata": {
                    "extraction_method": "text_in_memory",
                    "chunk_size_used": chunk_size,
                    "file_type": file_type
                }
            } for chunk_text in chunks]
            
        except Exception as e:
            logger.error(f"❌ Text-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Text-Extraktion fehlgeschlagen: {e}")]

    async def extract_spreadsheet(self, file_bytes: bytes, file_type: str) -> List[Dict]:
        """Extrahiere Tabellen aus Excel/CSV - GARANTIERT List[Dict]"""
        if not HAS_PANDAS:
            return [self._create_error_chunk("Tabellen-Support nicht verfügbar (pandas nicht installiert)")]
            
        try:
            if file_type == "csv":
                df = pd.read_csv(io.BytesIO(file_bytes))
            else:
                df = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
                
            chunks = []
            if isinstance(df, dict):
                for sheet_name, sheet_df in df.items():
                    chunks.extend(self._process_dataframe(sheet_df, {"sheet": sheet_name}))
            else:
                chunks.extend(self._process_dataframe(df))
                
            return chunks if chunks else [self._create_error_chunk("Keine Tabellendaten extrahiert")]
            
        except Exception as e:
            logger.error(f"❌ Tabelle fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Tabellen-Extraktion fehlgeschlagen: {e}")]

    def _process_dataframe(self, df: "pd.DataFrame", metadata: dict = None) -> List[Dict]:
        """Verarbeite DataFrame in Chunks - GARANTIERT List[Dict]"""
        metadata = metadata or {}
        chunks = []
        
        try:
            for i in range(0, len(df), 50):
                chunk_df = df.iloc[i:i + 50]
                chunks.append({
                    "content": chunk_df.to_json(orient="records"),
                    "type": "table",
                    "metadata": {
                        **metadata, 
                        "row_start": i,
                        "extraction_method": "pandas"
                    }
                })
        except Exception as e:
            logger.error(f"❌ DataFrame-Verarbeitung fehlgeschlagen: {e}")
            chunks.append(self._create_error_chunk(f"DataFrame-Fehler: {e}"))
            
        return chunks

    async def extract_email(self, eml_bytes: bytes) -> List[Dict]:
        """Extrahiere E-Mail-Metadaten und Inhalt - GARANTIERT List[Dict]"""
        try:
            msg = BytesParser(policy=policy.default).parsebytes(eml_bytes)
            chunks = [{
                "content": json.dumps({
                    "from": msg["From"],
                    "to": msg["To"],
                    "subject": msg["Subject"],
                    "date": msg["Date"],
                }),
                "type": "metadata",
                "metadata": {"extraction_method": "email_header"}
            }]
            
            body = msg.get_body(preferencelist=("plain", "html"))
            if body:
                content = body.get_content()
                if content and content.strip():
                    # FIX: Korrekte Verarbeitung der split_text Rückgabe
                    text_chunks = self.split_text(content.strip())
                    chunks.extend([{
                        "content": chunk_text, 
                        "type": "text",
                        "metadata": {"extraction_method": "email_body"}
                    } for chunk_text in text_chunks])
                    
            return chunks
            
        except Exception as e:
            logger.error(f"❌ E-Mail-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"E-Mail-Extraktion fehlgeschlagen: {e}")]

    def extract_html(self, html: str) -> List[Dict]:
        """Extrahiere Text aus HTML - GARANTIERT List[Dict]"""
        try:
            if HAS_BS4:
                soup = BeautifulSoup(html, "html.parser")
                text = soup.get_text(separator="\n", strip=True)
            else:
                import re
                text = re.sub("<[^<]+?>", "", html)
                
            return self.extract_text(text, "html")
            
        except Exception as e:
            logger.error(f"❌ HTML-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"HTML-Extraktion fehlgeschlagen: {e}")]

    async def extract_archive(self, zip_bytes: bytes) -> List[Dict]:
        """Extrahiere Dateiliste aus ZIP - GARANTIERT List[Dict]"""
        try:
            with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
                files = zf.namelist()
                return [{
                    "content": f"Archive enthält: {name}",
                    "type": "metadata",
                    "metadata": {
                        "archived_file": name,
                        "extraction_method": "zip_listing"
                    }
                } for name in files if not name.endswith("/")]
                
        except Exception as e:
            logger.error(f"❌ Archiv-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Archiv-Extraktion fehlgeschlagen: {e}")]

    # === UTILITY METHODEN ===
    
    def split_text(self, text: str, chunk_size: int = None) -> List[str]:
        """
        Teile Text in Chunks mit dynamischer Größe
        
        Returns:
            List[str]: Liste von Text-Chunks (nicht Dict!)
        """
        chunk_size = chunk_size or self.chunk_size
        words = text.split()
        chunks = []
        current_chunk = []
        current_size = 0

        for word in words:
            current_chunk.append(word)
            current_size += len(word) + 1
            
            if current_size >= chunk_size:
                chunks.append(" ".join(current_chunk))
                current_chunk = []
                current_size = 0

        if current_chunk:
            chunks.append(" ".join(current_chunk))

        return chunks or [""]

    def _create_error_chunk(self, error_message: str) -> Dict:
        """Erstelle standardisierten Error-Chunk"""
        return {
            "content": f"[FEHLER: {error_message}]",
            "type": "error",
            "metadata": {
                "extraction_method": "error",
                "error_message": error_message,
                "timestamp": datetime.now().isoformat()
            }
        }

    def _validate_chunks(self, chunks: List[Dict]) -> List[Dict]:
        """Validiere und normalisiere Chunk-Format"""
        validated = []
        for i, chunk in enumerate(chunks):
            if not isinstance(chunk, dict):
                logger.warning(f"⚠️ Chunk {i} ist kein Dict: {type(chunk)}")
                chunk = {"content": str(chunk), "type": "text", "metadata": {}}
                
            # Stelle sicher, dass required fields vorhanden sind
            if "content" not in chunk:
                chunk["content"] = "[Leerer Chunk]"
            if "type" not in chunk:
                chunk["type"] = "text"
            if "metadata" not in chunk:
                chunk["metadata"] = {}
                
            validated.append(chunk)
            
        return validated

    def _parse_metadata(self, metadata: Union[str, dict]) -> dict:
        """Parse metadata sicher zu dict"""
        if isinstance(metadata, dict):
            return metadata
        if isinstance(metadata, str):
            try:
                return json.loads(metadata)
            except:
                return {}
        return {}

    async def _save_chunks_to_db(self, conn, document_id: str, chunks: List[Dict]):
        """Speichere Chunks in Datenbank"""
        for idx, chunk in enumerate(chunks):
            await conn.execute(
                """
                INSERT INTO chunks (
                    document_id, chunk_index, content, content_type,
                    page_number, position, metadata
                ) VALUES ($1, $2, $3, $4, $5, $6, $7)
                ON CONFLICT (document_id, chunk_index) DO UPDATE
                SET content = EXCLUDED.content,
                    content_type = EXCLUDED.content_type
                """,
                document_id,
                idx,
                chunk["content"],
                chunk.get("type", "text"),
                chunk.get("page"),
                chunk.get("position"),
                json.dumps(chunk.get("metadata", {}))
            )

    async def _mark_task_failed(self, conn, task_id: str, doc_id: str, error_message: str):
        """Markiere Task als fehlgeschlagen"""
        await conn.execute(
            """UPDATE processing_queue 
               SET status = 'failed', error_message = $2, completed_at = NOW()
               WHERE id = $1""",
            task_id, error_message
        )
        await conn.execute(
            """UPDATE documents SET status = 'failed' WHERE id = $1""",
            doc_id
        )

    async def run(self):
        """Starte den OCR-Agent"""
        await self.connect()
        logger.info("🧠 Enhanced OCR Agent V2 gestartet")
        
        try:
            await self.process_queue()
        except KeyboardInterrupt:
            logger.info("🛑 OCR Agent wird beendet (KeyboardInterrupt)")
        except Exception as e:
            logger.error(f"💥 Unerwarteter Fehler: {e}")
            raise
        finally:
            await self.close()


def main():
    """Haupt-Funktion"""
    db_url = os.getenv(
        "DATABASE_URL",
        "postgresql://semanticuser:semantic2024@postgres:5432/semantic_doc_finder"
    )
    
    logger.info(f"🚀 Starte OCR Agent mit DB: {db_url.split('@')[1] if '@' in db_url else 'localhost'}")
    
    agent = OCRAgentV2(db_url)
    asyncio.run(agent.run())


if __name__ == "__main__":
    main(), cleaned_content):
                logger.debug(f"Base64 format ungültig (erste 50 chars): {cleaned_content[:50]}")
                return False
            
            # Test decode (nur erste 1000 chars für performance)
            test_content = cleaned_content[:1000] if len(cleaned_content) > 1000 else cleaned_content
            
            # Padding für test hinzufügen falls nötig
            if len(test_content) % 4 != 0:
                test_content += '=' * (4 - len(test_content) % 4)
                
            decoded = base64.b64decode(test_content)
            
            # Prüfe dass decode nicht leer ist
            if not decoded:
                logger.debug("Base64 dekodiert zu leerem content")
                return False
                
            logger.debug(f"✅ Base64 content validiert: {len(content)} chars → {len(decoded)} test bytes")
            return True
            
        except Exception as e:
            logger.debug(f"Base64 validation fehlgeschlagen: {e}")
            return False

    def _parse_metadata_robust(self, raw_metadata: Union[str, dict]) -> DocumentMetadata:
        """
        Robustes Parsing von Metadaten zu strukturiertem Format
        
        Args:
            raw_metadata: Raw metadata von database
            
        Returns:
            DocumentMetadata: Strukturierte und validierte Metadaten
        """
        try:
            # Parse zu dict falls string
            if isinstance(raw_metadata, str):
                try:
                    metadata_dict = json.loads(raw_metadata)
                except json.JSONDecodeError as e:
                    logger.warning(f"⚠️ JSON-Parsing fehlgeschlagen: {e}")
                    metadata_dict = {}
            elif isinstance(raw_metadata, dict):
                metadata_dict = raw_metadata
            else:
                logger.warning(f"⚠️ Unerwarteter metadata type: {type(raw_metadata)}")
                metadata_dict = {}
            
            # Erstelle DocumentMetadata mit validation
            metadata = DocumentMetadata(**metadata_dict)
            
            # Additional validation/normalization
            if not metadata.effective_file_content and not metadata.effective_file_path:
                logger.warning("⚠️ Keine file_content oder file_path in metadata gefunden")
            
            return metadata
            
        except Exception as e:
            logger.error(f"❌ Metadata-Parsing fehlgeschlagen: {e}")
            # Fallback zu default metadata
            return DocumentMetadata(source="metadata_parse_error")

    async def _get_file_bytes(self, file_content: Optional[str], file_path: Optional[str]) -> bytes:
        """DEPRECATED: Legacy method - use _get_file_bytes_robust instead"""
        # Erstelle legacy DocumentMetadata für compatibility
        legacy_metadata = DocumentMetadata(
            file_content=file_content,
            file_path=file_path
        )
        return await self._get_file_bytes_robust(legacy_metadata)

    async def _dispatch_extraction(self, file_type: str, file_bytes: bytes, filename: str) -> List[Dict]:
        """
        Dispatche zu spezifischer Extraktionsmethode
        
        Returns:
            List[Dict]: GARANTIERT List von Chunk-Dictionaries
        """
        try:
            if file_type == "pdf":
                return await self.extract_pdf(file_bytes, filename)
            elif file_type in ["png", "jpg", "jpeg", "tiff", "tif"]:
                return await self.extract_image(file_bytes)
            elif file_type in ["txt", "md", "log"]:
                content = file_bytes.decode("utf-8", errors="ignore")
                return self.extract_text(content, file_type)
            elif file_type in ["docx", "doc"]:
                return await self.extract_docx(file_bytes)
            elif file_type in ["xlsx", "xls", "csv"]:
                return await self.extract_spreadsheet(file_bytes, file_type)
            elif file_type == "eml":
                return await self.extract_email(file_bytes)
            elif file_type == "html":
                content = file_bytes.decode("utf-8", errors="ignore")
                return self.extract_html(content)
            elif file_type == "zip":
                return await self.extract_archive(file_bytes)
            else:
                # Fallback zu text
                content = file_bytes.decode("utf-8", errors="ignore")
                return self.extract_text(content, file_type)
                
        except Exception as e:
            logger.error(f"❌ Extraktion für {file_type} fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Extraktion für {file_type} fehlgeschlagen: {e}")]

    # === STREAMING METHODEN ===
    
    async def _extract_content_streaming(self, file_path: str, file_type: str, filename: str) -> List[Dict]:
        """Streaming-Extraktion für große Dateien - GARANTIERT List[Dict]"""
        try:
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            logger.info(f"🚀 Streaming für {filename}: {file_size_mb:.1f}MB")
            
            if file_type == "pdf":
                return await self._stream_large_pdf(file_path)
            elif file_type in ["txt", "md", "log"]:
                return await self._stream_large_text(file_path)
            elif file_type in ["docx", "doc"] and HAS_DOCX:
                return await self._stream_large_docx(file_path)
            else:
                # Fallback: normales Laden als Text
                logger.warning(f"⚠️ Kein Streaming für {file_type} - fallback zu normal")
                with open(file_path, "rb") as f:
                    file_bytes = f.read()
                content = file_bytes.decode("utf-8", errors="ignore")
                return self.extract_text(content, file_type)
                
        except Exception as e:
            logger.error(f"❌ Streaming fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Streaming fehlgeschlagen: {e}")]

    async def _stream_large_pdf(self, file_path: str) -> List[Dict]:
        """Verarbeite PDF seitenweise - GARANTIERT List[Dict]"""
        chunks = []
        try:
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            dpi = 150 if file_size_mb > 50 else 300
            tesseract_config = self.tesseract_config_fast if file_size_mb > 100 else self.tesseract_config
            
            logger.info(f"📄 PDF-Streaming: {file_path} ({file_size_mb:.1f} MB, DPI: {dpi})")

            images = pdf2image.convert_from_path(file_path, dpi=dpi)
            total_pages = len(images)

            for page_num, image in enumerate(images, 1):
                if total_pages > 50 and page_num % 10 == 0:
                    logger.info(f"📄 Fortschritt: {page_num}/{total_pages} ({page_num/total_pages*100:.1f}%)")
                
                text = pytesseract.image_to_string(
                    image,
                    lang=os.getenv("TESSERACT_LANG", "deu+eng"),
                    config=tesseract_config
                )

                if text.strip():
                    pdf_chunk_size = int(self.chunk_size * 1.5) if file_size_mb > 100 else self.chunk_size
                    page_text_chunks = self.split_text(text, pdf_chunk_size)
                    
                    for chunk_text in page_text_chunks:
                        chunks.append({
                            "content": chunk_text,
                            "type": "text",
                            "page": page_num,
                            "metadata": {
                                "extraction_method": "pdf_streaming",
                                "dpi": dpi,
                                "chunk_size_used": pdf_chunk_size
                            }
                        })

                del image
                if page_num % 5 == 0:
                    gc.collect()
                    
        except Exception as e:
            logger.error(f"❌ PDF-Streaming fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"PDF-Streaming fehlgeschlagen: {e}")]
            
        logger.info(f"✅ PDF verarbeitet: {len(chunks)} Chunks aus {total_pages} Seiten")
        return chunks

    async def _stream_large_text(self, file_path: str) -> List[Dict]:
        """Verarbeite große Textdateien line-by-line - GARANTIERT List[Dict]"""
        chunks = []
        try:
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            dynamic_chunk_size = min(self.chunk_size * 2, 2000) if file_size_mb > 100 else self.chunk_size
            
            logger.info(f"📦 Text-Streaming: {file_path} ({file_size_mb:.1f} MB), Chunk-Größe: {dynamic_chunk_size}")

            current_chunk = []
            current_size = 0
            line_count = 0
            
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                for line in f:
                    line = line.strip() + "\n"
                    if not line.strip():
                        continue
                        
                    current_chunk.append(line)
                    current_size += len(line)
                    line_count += 1

                    if current_size >= dynamic_chunk_size:
                        chunks.append({
                            "content": "".join(current_chunk),
                            "type": "text",
                            "metadata": {
                                "extraction_method": "text_streaming",
                                "lines_in_chunk": len(current_chunk),
                                "chunk_size_used": dynamic_chunk_size
                            }
                        })
                        current_chunk = []
                        current_size = 0

                        if len(chunks) % 10 == 0:
                            logger.info(f"📊 {len(chunks)} Chunks verarbeitet ({line_count} Zeilen)")
                            
                # Letzter Chunk
                if current_chunk:
                    chunks.append({
                        "content": "".join(current_chunk),
                        "type": "text",
                        "metadata": {
                            "extraction_method": "text_streaming",
                            "lines_in_chunk": len(current_chunk),
                            "chunk_size_used": dynamic_chunk_size,
                            "final_chunk": True
                        }
                    })
                    
        except Exception as e:
            logger.error(f"❌ Text-Streaming fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Text-Streaming fehlgeschlagen: {e}")]
            
        logger.info(f"✅ Textdatei verarbeitet: {len(chunks)} Chunks")
        return chunks

    async def _stream_large_docx(self, file_path: str) -> List[Dict]:
        """Verarbeite große DOCX-Dateien - GARANTIERT List[Dict]"""
        chunks = []
        try:
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            logger.info(f"📄 DOCX-Streaming: {file_path} ({file_size_mb:.1f} MB)")
            
            if not HAS_DOCX:
                return [self._create_error_chunk("DOCX-Support nicht verfügbar")]
            
            doc = DocxDocument(file_path)
            full_text = "\n".join([para.text for para in doc.paragraphs])
            
            docx_chunk_size = int(self.chunk_size * 1.3)
            text_chunks = self.split_text(full_text, docx_chunk_size)
            
            for chunk_text in text_chunks:
                chunks.append({
                    "content": chunk_text,
                    "type": "text",
                    "metadata": {
                        "extraction_method": "docx_streaming",
                        "chunk_size_used": docx_chunk_size
                    }
                })
                
            # Extrahiere Tabellen
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                    
                if table_data:
                    chunks.append({
                        "content": json.dumps(table_data),
                        "type": "table",
                        "metadata": {
                            "extraction_method": "docx_table",
                            "table_index": table_idx
                        }
                    })
                    
        except Exception as e:
            logger.error(f"❌ DOCX-Streaming fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"DOCX-Streaming fehlgeschlagen: {e}")]
            
        logger.info(f"✅ DOCX verarbeitet: {len(chunks)} Chunks")
        return chunks

    # === IN-MEMORY EXTRAKTIONSMETHODEN ===
    
    async def extract_pdf(self, pdf_bytes: bytes, filename: str) -> List[Dict]:
        """Extrahiere Text aus PDF mit OCR (In-Memory) - GARANTIERT List[Dict]"""
        chunks = []
        try:
            images = pdf2image.convert_from_bytes(pdf_bytes, dpi=300)
            logger.info(f"📄 PDF In-Memory: {filename} ({len(images)} Seiten)")
            
            for page_num, image in enumerate(images, 1):
                text = pytesseract.image_to_string(
                    image,
                    lang=os.getenv("TESSERACT_LANG", "deu+eng"),
                    config=self.tesseract_config
                )
                
                if text.strip():
                    page_chunks = self.split_text(text)
                    for chunk_text in page_chunks:
                        chunks.append({
                            "content": chunk_text,
                            "type": "text",
                            "page": page_num,
                            "metadata": {
                                "extraction_method": "pdf_in_memory",
                                "chunk_size_used": self.chunk_size
                            }
                        })
                        
                del image
                if page_num % 5 == 0:
                    gc.collect()
                    
        except Exception as e:
            logger.error(f"❌ PDF-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"PDF-Extraktion fehlgeschlagen: {e}")]
            
        return chunks if chunks else [self._create_error_chunk("Kein Text aus PDF extrahiert")]

    async def extract_image(self, image_bytes: bytes) -> List[Dict]:
        """Extrahiere Text aus einem Bild - GARANTIERT List[Dict]"""
        try:
            image = Image.open(io.BytesIO(image_bytes))
            text = pytesseract.image_to_string(
                image,
                lang=os.getenv("TESSERACT_LANG", "deu+eng"),
                config=self.tesseract_config
            )
            
            if not text.strip():
                return [{"content": "[Bild ohne erkannten Text]", "type": "text", "metadata": {"extraction_method": "image_ocr"}}]
                
            chunks = self.split_text(text)
            return [{
                "content": chunk_text,
                "type": "text",
                "page": 1,
                "metadata": {
                    "extraction_method": "image_ocr",
                    "chunk_size_used": self.chunk_size
                }
            } for chunk_text in chunks]
            
        except Exception as e:
            logger.error(f"❌ Bild-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Bild-Extraktion fehlgeschlagen: {e}")]

    async def extract_docx(self, docx_bytes: bytes) -> List[Dict]:
        """Extrahiere Text aus DOCX (In-Memory) - GARANTIERT List[Dict]"""
        if not HAS_DOCX:
            return [self._create_error_chunk("DOCX-Support nicht verfügbar (python-docx nicht installiert)")]
            
        try:
            doc = DocxDocument(io.BytesIO(docx_bytes))
            full_text = "\n".join([para.text for para in doc.paragraphs])
            
            docx_chunk_size = int(self.chunk_size * 1.3)
            chunks = self.split_text(full_text, docx_chunk_size)
            
            result = [{
                "content": chunk_text,
                "type": "text",
                "metadata": {
                    "extraction_method": "docx_in_memory",
                    "chunk_size_used": docx_chunk_size
                }
            } for chunk_text in chunks]
            
            # Extrahiere Tabellen
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                    
                if table_data:
                    result.append({
                        "content": json.dumps(table_data),
                        "type": "table",
                        "metadata": {
                            "extraction_method": "docx_table",
                            "table_index": table_idx
                        }
                    })
                    
            return result if result else [self._create_error_chunk("Kein Inhalt aus DOCX extrahiert")]
            
        except Exception as e:
            logger.error(f"❌ DOCX-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"DOCX-Extraktion fehlgeschlagen: {e}")]

    def extract_text(self, text: str, file_type: str = "txt") -> List[Dict]:
        """Extrahiere Text mit adaptiver Chunk-Größe - GARANTIERT List[Dict]"""
        if not text.strip():
            return [{"content": "[Leerer Text]", "type": "text", "metadata": {"extraction_method": "text_in_memory"}}]
            
        try:
            # Adaptive Chunk-Größen basierend auf Dateityp
            if file_type in ['csv', 'tsv']:
                chunk_size = self.chunk_size // 2
            elif file_type in ['md', 'txt']:
                chunk_size = int(self.chunk_size * 1.2)
            else:
                chunk_size = self.chunk_size
                
            chunks = self.split_text(text, chunk_size)
            return [{
                "content": chunk_text,
                "type": "text",
                "metadata": {
                    "extraction_method": "text_in_memory",
                    "chunk_size_used": chunk_size,
                    "file_type": file_type
                }
            } for chunk_text in chunks]
            
        except Exception as e:
            logger.error(f"❌ Text-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Text-Extraktion fehlgeschlagen: {e}")]

    async def extract_spreadsheet(self, file_bytes: bytes, file_type: str) -> List[Dict]:
        """Extrahiere Tabellen aus Excel/CSV - GARANTIERT List[Dict]"""
        if not HAS_PANDAS:
            return [self._create_error_chunk("Tabellen-Support nicht verfügbar (pandas nicht installiert)")]
            
        try:
            if file_type == "csv":
                df = pd.read_csv(io.BytesIO(file_bytes))
            else:
                df = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
                
            chunks = []
            if isinstance(df, dict):
                for sheet_name, sheet_df in df.items():
                    chunks.extend(self._process_dataframe(sheet_df, {"sheet": sheet_name}))
            else:
                chunks.extend(self._process_dataframe(df))
                
            return chunks if chunks else [self._create_error_chunk("Keine Tabellendaten extrahiert")]
            
        except Exception as e:
            logger.error(f"❌ Tabelle fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Tabellen-Extraktion fehlgeschlagen: {e}")]

    def _process_dataframe(self, df: "pd.DataFrame", metadata: dict = None) -> List[Dict]:
        """Verarbeite DataFrame in Chunks - GARANTIERT List[Dict]"""
        metadata = metadata or {}
        chunks = []
        
        try:
            for i in range(0, len(df), 50):
                chunk_df = df.iloc[i:i + 50]
                chunks.append({
                    "content": chunk_df.to_json(orient="records"),
                    "type": "table",
                    "metadata": {
                        **metadata, 
                        "row_start": i,
                        "extraction_method": "pandas"
                    }
                })
        except Exception as e:
            logger.error(f"❌ DataFrame-Verarbeitung fehlgeschlagen: {e}")
            chunks.append(self._create_error_chunk(f"DataFrame-Fehler: {e}"))
            
        return chunks

    async def extract_email(self, eml_bytes: bytes) -> List[Dict]:
        """Extrahiere E-Mail-Metadaten und Inhalt - GARANTIERT List[Dict]"""
        try:
            msg = BytesParser(policy=policy.default).parsebytes(eml_bytes)
            chunks = [{
                "content": json.dumps({
                    "from": msg["From"],
                    "to": msg["To"],
                    "subject": msg["Subject"],
                    "date": msg["Date"],
                }),
                "type": "metadata",
                "metadata": {"extraction_method": "email_header"}
            }]
            
            body = msg.get_body(preferencelist=("plain", "html"))
            if body:
                content = body.get_content()
                if content and content.strip():
                    # FIX: Korrekte Verarbeitung der split_text Rückgabe
                    text_chunks = self.split_text(content.strip())
                    chunks.extend([{
                        "content": chunk_text, 
                        "type": "text",
                        "metadata": {"extraction_method": "email_body"}
                    } for chunk_text in text_chunks])
                    
            return chunks
            
        except Exception as e:
            logger.error(f"❌ E-Mail-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"E-Mail-Extraktion fehlgeschlagen: {e}")]

    def extract_html(self, html: str) -> List[Dict]:
        """Extrahiere Text aus HTML - GARANTIERT List[Dict]"""
        try:
            if HAS_BS4:
                soup = BeautifulSoup(html, "html.parser")
                text = soup.get_text(separator="\n", strip=True)
            else:
                import re
                text = re.sub("<[^<]+?>", "", html)
                
            return self.extract_text(text, "html")
            
        except Exception as e:
            logger.error(f"❌ HTML-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"HTML-Extraktion fehlgeschlagen: {e}")]

    async def extract_archive(self, zip_bytes: bytes) -> List[Dict]:
        """Extrahiere Dateiliste aus ZIP - GARANTIERT List[Dict]"""
        try:
            with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
                files = zf.namelist()
                return [{
                    "content": f"Archive enthält: {name}",
                    "type": "metadata",
                    "metadata": {
                        "archived_file": name,
                        "extraction_method": "zip_listing"
                    }
                } for name in files if not name.endswith("/")]
                
        except Exception as e:
            logger.error(f"❌ Archiv-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Archiv-Extraktion fehlgeschlagen: {e}")]

    # === UTILITY METHODEN ===
    
    def split_text(self, text: str, chunk_size: int = None) -> List[str]:
        """
        Teile Text in Chunks mit dynamischer Größe
        
        Returns:
            List[str]: Liste von Text-Chunks (nicht Dict!)
        """
        chunk_size = chunk_size or self.chunk_size
        words = text.split()
        chunks = []
        current_chunk = []
        current_size = 0

        for word in words:
            current_chunk.append(word)
            current_size += len(word) + 1
            
            if current_size >= chunk_size:
                chunks.append(" ".join(current_chunk))
                current_chunk = []
                current_size = 0

        if current_chunk:
            chunks.append(" ".join(current_chunk))

        return chunks or [""]

    def _create_error_chunk(self, error_message: str) -> Dict:
        """Erstelle standardisierten Error-Chunk"""
        return {
            "content": f"[FEHLER: {error_message}]",
            "type": "error",
            "metadata": {
                "extraction_method": "error",
                "error_message": error_message,
                "timestamp": datetime.now().isoformat()
            }
        }

    def _validate_chunks(self, chunks: List[Dict]) -> List[Dict]:
        """Validiere und normalisiere Chunk-Format"""
        validated = []
        for i, chunk in enumerate(chunks):
            if not isinstance(chunk, dict):
                logger.warning(f"⚠️ Chunk {i} ist kein Dict: {type(chunk)}")
                chunk = {"content": str(chunk), "type": "text", "metadata": {}}
                
            # Stelle sicher, dass required fields vorhanden sind
            if "content" not in chunk:
                chunk["content"] = "[Leerer Chunk]"
            if "type" not in chunk:
                chunk["type"] = "text"
            if "metadata" not in chunk:
                chunk["metadata"] = {}
                
            validated.append(chunk)
            
        return validated

    def _parse_metadata(self, metadata: Union[str, dict]) -> dict:
        """Parse metadata sicher zu dict"""
        if isinstance(metadata, dict):
            return metadata
        if isinstance(metadata, str):
            try:
                return json.loads(metadata)
            except:
                return {}
        return {}

    async def _save_chunks_to_db(self, conn, document_id: str, chunks: List[Dict]):
        """Speichere Chunks in Datenbank"""
        for idx, chunk in enumerate(chunks):
            await conn.execute(
                """
                INSERT INTO chunks (
                    document_id, chunk_index, content, content_type,
                    page_number, position, metadata
                ) VALUES ($1, $2, $3, $4, $5, $6, $7)
                ON CONFLICT (document_id, chunk_index) DO UPDATE
                SET content = EXCLUDED.content,
                    content_type = EXCLUDED.content_type
                """,
                document_id,
                idx,
                chunk["content"],
                chunk.get("type", "text"),
                chunk.get("page"),
                chunk.get("position"),
                json.dumps(chunk.get("metadata", {}))
            )

    async def _mark_task_failed(self, conn, task_id: str, doc_id: str, error_message: str):
        """Markiere Task als fehlgeschlagen"""
        await conn.execute(
            """UPDATE processing_queue 
               SET status = 'failed', error_message = $2, completed_at = NOW()
               WHERE id = $1""",
            task_id, error_message
        )
        await conn.execute(
            """UPDATE documents SET status = 'failed' WHERE id = $1""",
            doc_id
        )

    async def run(self):
        """Starte den OCR-Agent"""
        await self.connect()
        logger.info("🧠 Enhanced OCR Agent V2 gestartet")
        
        try:
            await self.process_queue()
        except KeyboardInterrupt:
            logger.info("🛑 OCR Agent wird beendet (KeyboardInterrupt)")
        except Exception as e:
            logger.error(f"💥 Unerwarteter Fehler: {e}")
            raise
        finally:
            await self.close()


def main():
    """Haupt-Funktion"""
    db_url = os.getenv(
        "DATABASE_URL",
        "postgresql://semanticuser:semantic2024@postgres:5432/semantic_doc_finder"
    )
    
    logger.info(f"🚀 Starte OCR Agent mit DB: {db_url.split('@')[1] if '@' in db_url else 'localhost'}")
    
    agent = OCRAgentV2(db_url)
    asyncio.run(agent.run())


if __name__ == "__main__":
    main(), cleaned):
            return False
            
        return True


class OCRAgentV2:
    """
    Enhanced OCR Agent V2 mit vollständiger N8N-Integration
    
    GARANTIEN:
    - Alle extract_* Methoden geben List[Dict] zurück  
    - Einheitliche Chunk-Struktur: {"content": str, "type": str, "metadata": dict, ...}
    - Robuste Error-Behandlung
    - Streaming + In-Memory Modi
    """
    
    def __init__(self, db_url: str):
        self.db_url = db_url
        self.db_pool = None
        self.chunk_size = int(os.getenv("CHUNK_SIZE", "1000"))
        self.running = True

        # Streaming settings
        self.MAX_FILE_SIZE = 100 * 1024 * 1024  # 100 MB
        
        # Tesseract configs
        self.tesseract_config = "--oem 3 --psm 6"
        self.tesseract_config_fast = "--oem 1 --psm 6"
        
        logger.info(f"📊 OCR Agent V2 initialisiert:")
        logger.info(f"   • Chunk-Größe: {self.chunk_size}")
        logger.info(f"   • Max File Size: {self.MAX_FILE_SIZE / (1024*1024):.0f}MB")
        logger.info(f"   • DOCX Support: {HAS_DOCX}")
        logger.info(f"   • Pandas Support: {HAS_PANDAS}")
        logger.info(f"   • BeautifulSoup Support: {HAS_BS4}")

    async def connect(self):
        """Verbinde mit der PostgreSQL-Datenbank"""
        try:
            self.db_pool = await asyncpg.create_pool(self.db_url)
            logger.info("✅ Verbindung zur PostgreSQL-Datenbank hergestellt")
        except Exception as e:
            logger.error(f"❌ Datenbankverbindung fehlgeschlagen: {e}")
            raise

    async def close(self):
        """Schließe die Verbindung zur Datenbank"""
        if self.db_pool:
            await self.db_pool.close()
        logger.info("🛑 OCR Agent V2 beendet")

    async def process_queue(self):
        """Verarbeite die Warteschlange kontinuierlich"""
        while self.running:
            try:
                async with self.db_pool.acquire() as conn:
                    task = await conn.fetchrow(
                        """
                        UPDATE processing_queue 
                        SET status = 'processing', started_at = NOW()
                        WHERE id = (
                            SELECT id FROM processing_queue 
                            WHERE status = 'pending' 
                            AND task_type IN ('ocr', 'extract')
                            AND (NOT deferred OR EXTRACT(hour FROM NOW()) BETWEEN 22 AND 6)
                            ORDER BY priority DESC, created_at ASC
                            LIMIT 1
                            FOR UPDATE SKIP LOCKED
                        )
                        RETURNING *
                        """
                    )
                    if task:
                        await self.process_document(task)
                    else:
                        await asyncio.sleep(5)
            except Exception as e:
                logger.error(f"⚠️ Fehler bei Warteschlange: {e}")
                await asyncio.sleep(10)

    async def process_document(self, task: Dict):
        """
        Verarbeite ein Dokument basierend auf Metadaten
        
        Args:
            task: Database row from processing_queue
        """
        async with self.db_pool.acquire() as conn:
            try:
                doc = await conn.fetchrow(
                    "SELECT * FROM documents WHERE id = $1",
                    task["document_id"]
                )
                if not doc:
                    raise ValueError(f"🚫 Dokument {task['document_id']} nicht gefunden")

                logger.info(f"📄 Verarbeite: {doc['original_name']} (ID: {doc['id']})")

                # Extrahiere Inhalt basierend auf Dateityp - GARANTIERT List[Dict]
                chunks = await self.extract_content(doc)
                
                if not chunks:
                    logger.warning(f"⚠️ Keine Chunks erstellt für {doc['original_name']}")
                    chunks = [self._create_error_chunk("Keine Inhalte extrahiert")]

                # Validiere Chunk-Format
                chunks = self._validate_chunks(chunks)

                # Speichere Chunks in Datenbank
                await self._save_chunks_to_db(conn, doc["id"], chunks)

                # Dokumentstatus aktualisieren
                await conn.execute(
                    """UPDATE documents 
                       SET status = 'completed', processed_at = NOW()
                       WHERE id = $1""",
                    doc["id"]
                )

                # Task abschließen
                await conn.execute(
                    """UPDATE processing_queue 
                       SET status = 'completed', completed_at = NOW()
                       WHERE id = $1""",
                    task["id"]
                )

                # Enhancement-Task für LLM-Pipeline erstellen
                await conn.execute(
                    """INSERT INTO processing_queue (document_id, task_type, priority)
                       VALUES ($1, 'enhance', $2)""",
                    doc["id"], task["priority"]
                )

                logger.info(f"✅ Dokument verarbeitet: {doc['original_name']} ({len(chunks)} Chunks)")

            except Exception as e:
                logger.error(f"❌ Fehler bei Dokumentenverarbeitung: {e}")
                await self._mark_task_failed(conn, task["id"], doc["id"], str(e))

    async def extract_content(self, doc: Dict) -> List[Dict]:
        """
        Hauptmethode für Content-Extraktion mit N8N processing_mode Support
        
        Returns:
            List[Dict]: GARANTIERT List von Chunk-Dictionaries
        """
        try:
            file_type = doc["file_type"].lower()
            raw_metadata = doc.get("metadata", {})
            
            # Parse metadata zu strukturiertem Format
            metadata = self._parse_metadata_robust(raw_metadata)
            
            logger.info(f"🔍 Processing mode: {metadata.processing_mode}, file_type: {file_type}")

            # Automatische Streaming-Entscheidung für große Dateien
            if metadata.effective_file_path and os.path.exists(metadata.effective_file_path):
                file_size = os.path.getsize(metadata.effective_file_path)
                file_size_mb = file_size / (1024 * 1024)
                
                if file_size > self.MAX_FILE_SIZE and metadata.processing_mode != "streaming":
                    metadata.processing_mode = "streaming"
                    logger.warning(f"⚠️ Große Datei ({file_size_mb:.1f}MB) → automatischer Fallback zu 'streaming'")
                else:
                    logger.info(f"📦 Datei: {metadata.effective_file_path} ({file_size_mb:.1f} MB), Modus: {metadata.processing_mode}")

            # === STREAMING MODUS ===
            if metadata.processing_mode == "streaming" and metadata.effective_file_path and os.path.exists(metadata.effective_file_path):
                logger.info(f"⏳ Streaming-Modus aktiviert: {metadata.effective_file_path}")
                return await self._extract_content_streaming(metadata.effective_file_path, file_type, doc["original_name"])

            # === IN-MEMORY MODUS ===
            logger.info(f"💾 In-Memory-Modus für {file_type}")
            
            # Hole file_bytes mit robuster Fallback-Chain
            file_bytes = await self._get_file_bytes_robust(metadata)
            
            # Dispatche zu spezifischer extract-Methode
            return await self._dispatch_extraction(file_type, file_bytes, doc["original_name"])
            
        except Exception as e:
            logger.error(f"❌ Content-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Content-Extraktion fehlgeschlagen: {e}")]

    async def _get_file_bytes_robust(self, metadata: DocumentMetadata) -> bytes:
        """
        Robuste file_bytes Extraktion mit umfassender Fallback-Chain
        
        Args:
            metadata: Strukturierte Metadaten
            
        Returns:
            bytes: File content als bytes
            
        Raises:
            ValueError: Wenn keine gültige Datenquelle gefunden wird
        """
        # === FALLBACK CHAIN FÜR FILE CONTENT ===
        file_content_candidates = [
            metadata.effective_file_content,  # Nutzt property mit fallback
            metadata.file_content,
            metadata.fileContent,  # camelCase variant
        ]
        
        # Suche ersten gültigen base64 content
        for candidate in file_content_candidates:
            if candidate and self._validate_base64_content(candidate):
                try:
                    # Bereinige base64 content (entferne whitespace, newlines)
                    cleaned_content = self._clean_base64_content(candidate)
                    file_bytes = base64.b64decode(cleaned_content)
                    
                    # Validiere dass decoded content nicht leer ist
                    if not file_bytes:
                        logger.warning(f"⚠️ Base64 dekodiert zu leerem content")
                        continue
                        
                    # Validiere dass decoded content sinnvoll aussieht (mindestens 10 bytes)
                    if len(file_bytes) < 10:
                        logger.warning(f"⚠️ Base64 content zu klein: {len(file_bytes)} bytes")
                        continue
                        
                    logger.info(f"📦 Base64-Inhalt dekodiert: {len(file_bytes)} bytes (aus {len(candidate)} chars)")
                    return file_bytes
                except Exception as e:
                    logger.warning(f"⚠️ Base64-Dekodierung fehlgeschlagen für candidate: {e}")
                    continue
        
        # === FALLBACK CHAIN FÜR FILE PATH ===
        file_path_candidates = [
            metadata.effective_file_path,  # Nutzt property mit fallback  
            metadata.file_path,
            metadata.filePath,  # camelCase variant
        ]
        
        # Suche ersten gültigen file path
        for candidate in file_path_candidates:
            if candidate and self._validate_file_path(candidate):
                try:
                    with open(candidate, "rb") as f:
                        file_bytes = f.read()
                    
                    # Validiere dass file content nicht leer ist
                    if not file_bytes:
                        logger.warning(f"⚠️ Datei ist leer: {candidate}")
                        continue
                        
                    logger.info(f"📁 Datei gelesen: {len(file_bytes)} bytes von {candidate}")
                    return file_bytes
                except Exception as e:
                    logger.warning(f"⚠️ Datei-Lesen fehlgeschlagen für {candidate}: {e}")
                    continue
        
        # === ERROR CASE ===
        error_details = {
            "file_content_candidates": [
                {
                    "exists": bool(c), 
                    "valid": self._validate_base64_content(c) if c else False,
                    "length": len(c) if c else 0
                } for c in file_content_candidates
            ],
            "file_path_candidates": [
                {
                    "path": c,
                    "exists": os.path.exists(c) if c else False,
                    "valid": self._validate_file_path(c) if c else False
                } for c in file_path_candidates
            ],
            "metadata_keys": list(metadata.dict(exclude_none=True).keys()),
            "processing_mode": metadata.processing_mode
        }
        
        raise ValueError(f"Keine gültige Datenquelle gefunden. Details: {error_details}")

    def _clean_base64_content(self, content: str) -> str:
        """
        Bereinige base64 content von häufigen Problemen
        
        Args:
            content: Raw base64 string
            
        Returns:
            str: Cleaned base64 string
        """
        if not content:
            return ""
            
        # Entferne whitespace, newlines, carriage returns
        cleaned = content.strip().replace('\n', '').replace('\r', '').replace(' ', '').replace('\t', '')
        
        # Entferne data URL prefix falls vorhanden (data:image/png;base64,...)
        if cleaned.startswith('data:'):
            if ',' in cleaned:
                cleaned = cleaned.split(',', 1)[1]
        
        # Validiere base64 format (grundlegend)
        import re
        if not re.match(r'^[A-Za-z0-9+/]*={0,2}

    async def _dispatch_extraction(self, file_type: str, file_bytes: bytes, filename: str) -> List[Dict]:
        """
        Dispatche zu spezifischer Extraktionsmethode
        
        Returns:
            List[Dict]: GARANTIERT List von Chunk-Dictionaries
        """
        try:
            if file_type == "pdf":
                return await self.extract_pdf(file_bytes, filename)
            elif file_type in ["png", "jpg", "jpeg", "tiff", "tif"]:
                return await self.extract_image(file_bytes)
            elif file_type in ["txt", "md", "log"]:
                content = file_bytes.decode("utf-8", errors="ignore")
                return self.extract_text(content, file_type)
            elif file_type in ["docx", "doc"]:
                return await self.extract_docx(file_bytes)
            elif file_type in ["xlsx", "xls", "csv"]:
                return await self.extract_spreadsheet(file_bytes, file_type)
            elif file_type == "eml":
                return await self.extract_email(file_bytes)
            elif file_type == "html":
                content = file_bytes.decode("utf-8", errors="ignore")
                return self.extract_html(content)
            elif file_type == "zip":
                return await self.extract_archive(file_bytes)
            else:
                # Fallback zu text
                content = file_bytes.decode("utf-8", errors="ignore")
                return self.extract_text(content, file_type)
                
        except Exception as e:
            logger.error(f"❌ Extraktion für {file_type} fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Extraktion für {file_type} fehlgeschlagen: {e}")]

    # === STREAMING METHODEN ===
    
    async def _extract_content_streaming(self, file_path: str, file_type: str, filename: str) -> List[Dict]:
        """Streaming-Extraktion für große Dateien - GARANTIERT List[Dict]"""
        try:
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            logger.info(f"🚀 Streaming für {filename}: {file_size_mb:.1f}MB")
            
            if file_type == "pdf":
                return await self._stream_large_pdf(file_path)
            elif file_type in ["txt", "md", "log"]:
                return await self._stream_large_text(file_path)
            elif file_type in ["docx", "doc"] and HAS_DOCX:
                return await self._stream_large_docx(file_path)
            else:
                # Fallback: normales Laden als Text
                logger.warning(f"⚠️ Kein Streaming für {file_type} - fallback zu normal")
                with open(file_path, "rb") as f:
                    file_bytes = f.read()
                content = file_bytes.decode("utf-8", errors="ignore")
                return self.extract_text(content, file_type)
                
        except Exception as e:
            logger.error(f"❌ Streaming fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Streaming fehlgeschlagen: {e}")]

    async def _stream_large_pdf(self, file_path: str) -> List[Dict]:
        """Verarbeite PDF seitenweise - GARANTIERT List[Dict]"""
        chunks = []
        try:
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            dpi = 150 if file_size_mb > 50 else 300
            tesseract_config = self.tesseract_config_fast if file_size_mb > 100 else self.tesseract_config
            
            logger.info(f"📄 PDF-Streaming: {file_path} ({file_size_mb:.1f} MB, DPI: {dpi})")

            images = pdf2image.convert_from_path(file_path, dpi=dpi)
            total_pages = len(images)

            for page_num, image in enumerate(images, 1):
                if total_pages > 50 and page_num % 10 == 0:
                    logger.info(f"📄 Fortschritt: {page_num}/{total_pages} ({page_num/total_pages*100:.1f}%)")
                
                text = pytesseract.image_to_string(
                    image,
                    lang=os.getenv("TESSERACT_LANG", "deu+eng"),
                    config=tesseract_config
                )

                if text.strip():
                    pdf_chunk_size = int(self.chunk_size * 1.5) if file_size_mb > 100 else self.chunk_size
                    page_text_chunks = self.split_text(text, pdf_chunk_size)
                    
                    for chunk_text in page_text_chunks:
                        chunks.append({
                            "content": chunk_text,
                            "type": "text",
                            "page": page_num,
                            "metadata": {
                                "extraction_method": "pdf_streaming",
                                "dpi": dpi,
                                "chunk_size_used": pdf_chunk_size
                            }
                        })

                del image
                if page_num % 5 == 0:
                    gc.collect()
                    
        except Exception as e:
            logger.error(f"❌ PDF-Streaming fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"PDF-Streaming fehlgeschlagen: {e}")]
            
        logger.info(f"✅ PDF verarbeitet: {len(chunks)} Chunks aus {total_pages} Seiten")
        return chunks

    async def _stream_large_text(self, file_path: str) -> List[Dict]:
        """Verarbeite große Textdateien line-by-line - GARANTIERT List[Dict]"""
        chunks = []
        try:
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            dynamic_chunk_size = min(self.chunk_size * 2, 2000) if file_size_mb > 100 else self.chunk_size
            
            logger.info(f"📦 Text-Streaming: {file_path} ({file_size_mb:.1f} MB), Chunk-Größe: {dynamic_chunk_size}")

            current_chunk = []
            current_size = 0
            line_count = 0
            
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                for line in f:
                    line = line.strip() + "\n"
                    if not line.strip():
                        continue
                        
                    current_chunk.append(line)
                    current_size += len(line)
                    line_count += 1

                    if current_size >= dynamic_chunk_size:
                        chunks.append({
                            "content": "".join(current_chunk),
                            "type": "text",
                            "metadata": {
                                "extraction_method": "text_streaming",
                                "lines_in_chunk": len(current_chunk),
                                "chunk_size_used": dynamic_chunk_size
                            }
                        })
                        current_chunk = []
                        current_size = 0

                        if len(chunks) % 10 == 0:
                            logger.info(f"📊 {len(chunks)} Chunks verarbeitet ({line_count} Zeilen)")
                            
                # Letzter Chunk
                if current_chunk:
                    chunks.append({
                        "content": "".join(current_chunk),
                        "type": "text",
                        "metadata": {
                            "extraction_method": "text_streaming",
                            "lines_in_chunk": len(current_chunk),
                            "chunk_size_used": dynamic_chunk_size,
                            "final_chunk": True
                        }
                    })
                    
        except Exception as e:
            logger.error(f"❌ Text-Streaming fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Text-Streaming fehlgeschlagen: {e}")]
            
        logger.info(f"✅ Textdatei verarbeitet: {len(chunks)} Chunks")
        return chunks

    async def _stream_large_docx(self, file_path: str) -> List[Dict]:
        """Verarbeite große DOCX-Dateien - GARANTIERT List[Dict]"""
        chunks = []
        try:
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            logger.info(f"📄 DOCX-Streaming: {file_path} ({file_size_mb:.1f} MB)")
            
            if not HAS_DOCX:
                return [self._create_error_chunk("DOCX-Support nicht verfügbar")]
            
            doc = DocxDocument(file_path)
            full_text = "\n".join([para.text for para in doc.paragraphs])
            
            docx_chunk_size = int(self.chunk_size * 1.3)
            text_chunks = self.split_text(full_text, docx_chunk_size)
            
            for chunk_text in text_chunks:
                chunks.append({
                    "content": chunk_text,
                    "type": "text",
                    "metadata": {
                        "extraction_method": "docx_streaming",
                        "chunk_size_used": docx_chunk_size
                    }
                })
                
            # Extrahiere Tabellen
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                    
                if table_data:
                    chunks.append({
                        "content": json.dumps(table_data),
                        "type": "table",
                        "metadata": {
                            "extraction_method": "docx_table",
                            "table_index": table_idx
                        }
                    })
                    
        except Exception as e:
            logger.error(f"❌ DOCX-Streaming fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"DOCX-Streaming fehlgeschlagen: {e}")]
            
        logger.info(f"✅ DOCX verarbeitet: {len(chunks)} Chunks")
        return chunks

    # === IN-MEMORY EXTRAKTIONSMETHODEN ===
    
    async def extract_pdf(self, pdf_bytes: bytes, filename: str) -> List[Dict]:
        """Extrahiere Text aus PDF mit OCR (In-Memory) - GARANTIERT List[Dict]"""
        chunks = []
        try:
            images = pdf2image.convert_from_bytes(pdf_bytes, dpi=300)
            logger.info(f"📄 PDF In-Memory: {filename} ({len(images)} Seiten)")
            
            for page_num, image in enumerate(images, 1):
                text = pytesseract.image_to_string(
                    image,
                    lang=os.getenv("TESSERACT_LANG", "deu+eng"),
                    config=self.tesseract_config
                )
                
                if text.strip():
                    page_chunks = self.split_text(text)
                    for chunk_text in page_chunks:
                        chunks.append({
                            "content": chunk_text,
                            "type": "text",
                            "page": page_num,
                            "metadata": {
                                "extraction_method": "pdf_in_memory",
                                "chunk_size_used": self.chunk_size
                            }
                        })
                        
                del image
                if page_num % 5 == 0:
                    gc.collect()
                    
        except Exception as e:
            logger.error(f"❌ PDF-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"PDF-Extraktion fehlgeschlagen: {e}")]
            
        return chunks if chunks else [self._create_error_chunk("Kein Text aus PDF extrahiert")]

    async def extract_image(self, image_bytes: bytes) -> List[Dict]:
        """Extrahiere Text aus einem Bild - GARANTIERT List[Dict]"""
        try:
            image = Image.open(io.BytesIO(image_bytes))
            text = pytesseract.image_to_string(
                image,
                lang=os.getenv("TESSERACT_LANG", "deu+eng"),
                config=self.tesseract_config
            )
            
            if not text.strip():
                return [{"content": "[Bild ohne erkannten Text]", "type": "text", "metadata": {"extraction_method": "image_ocr"}}]
                
            chunks = self.split_text(text)
            return [{
                "content": chunk_text,
                "type": "text",
                "page": 1,
                "metadata": {
                    "extraction_method": "image_ocr",
                    "chunk_size_used": self.chunk_size
                }
            } for chunk_text in chunks]
            
        except Exception as e:
            logger.error(f"❌ Bild-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Bild-Extraktion fehlgeschlagen: {e}")]

    async def extract_docx(self, docx_bytes: bytes) -> List[Dict]:
        """Extrahiere Text aus DOCX (In-Memory) - GARANTIERT List[Dict]"""
        if not HAS_DOCX:
            return [self._create_error_chunk("DOCX-Support nicht verfügbar (python-docx nicht installiert)")]
            
        try:
            doc = DocxDocument(io.BytesIO(docx_bytes))
            full_text = "\n".join([para.text for para in doc.paragraphs])
            
            docx_chunk_size = int(self.chunk_size * 1.3)
            chunks = self.split_text(full_text, docx_chunk_size)
            
            result = [{
                "content": chunk_text,
                "type": "text",
                "metadata": {
                    "extraction_method": "docx_in_memory",
                    "chunk_size_used": docx_chunk_size
                }
            } for chunk_text in chunks]
            
            # Extrahiere Tabellen
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                    
                if table_data:
                    result.append({
                        "content": json.dumps(table_data),
                        "type": "table",
                        "metadata": {
                            "extraction_method": "docx_table",
                            "table_index": table_idx
                        }
                    })
                    
            return result if result else [self._create_error_chunk("Kein Inhalt aus DOCX extrahiert")]
            
        except Exception as e:
            logger.error(f"❌ DOCX-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"DOCX-Extraktion fehlgeschlagen: {e}")]

    def extract_text(self, text: str, file_type: str = "txt") -> List[Dict]:
        """Extrahiere Text mit adaptiver Chunk-Größe - GARANTIERT List[Dict]"""
        if not text.strip():
            return [{"content": "[Leerer Text]", "type": "text", "metadata": {"extraction_method": "text_in_memory"}}]
            
        try:
            # Adaptive Chunk-Größen basierend auf Dateityp
            if file_type in ['csv', 'tsv']:
                chunk_size = self.chunk_size // 2
            elif file_type in ['md', 'txt']:
                chunk_size = int(self.chunk_size * 1.2)
            else:
                chunk_size = self.chunk_size
                
            chunks = self.split_text(text, chunk_size)
            return [{
                "content": chunk_text,
                "type": "text",
                "metadata": {
                    "extraction_method": "text_in_memory",
                    "chunk_size_used": chunk_size,
                    "file_type": file_type
                }
            } for chunk_text in chunks]
            
        except Exception as e:
            logger.error(f"❌ Text-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Text-Extraktion fehlgeschlagen: {e}")]

    async def extract_spreadsheet(self, file_bytes: bytes, file_type: str) -> List[Dict]:
        """Extrahiere Tabellen aus Excel/CSV - GARANTIERT List[Dict]"""
        if not HAS_PANDAS:
            return [self._create_error_chunk("Tabellen-Support nicht verfügbar (pandas nicht installiert)")]
            
        try:
            if file_type == "csv":
                df = pd.read_csv(io.BytesIO(file_bytes))
            else:
                df = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
                
            chunks = []
            if isinstance(df, dict):
                for sheet_name, sheet_df in df.items():
                    chunks.extend(self._process_dataframe(sheet_df, {"sheet": sheet_name}))
            else:
                chunks.extend(self._process_dataframe(df))
                
            return chunks if chunks else [self._create_error_chunk("Keine Tabellendaten extrahiert")]
            
        except Exception as e:
            logger.error(f"❌ Tabelle fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Tabellen-Extraktion fehlgeschlagen: {e}")]

    def _process_dataframe(self, df: "pd.DataFrame", metadata: dict = None) -> List[Dict]:
        """Verarbeite DataFrame in Chunks - GARANTIERT List[Dict]"""
        metadata = metadata or {}
        chunks = []
        
        try:
            for i in range(0, len(df), 50):
                chunk_df = df.iloc[i:i + 50]
                chunks.append({
                    "content": chunk_df.to_json(orient="records"),
                    "type": "table",
                    "metadata": {
                        **metadata, 
                        "row_start": i,
                        "extraction_method": "pandas"
                    }
                })
        except Exception as e:
            logger.error(f"❌ DataFrame-Verarbeitung fehlgeschlagen: {e}")
            chunks.append(self._create_error_chunk(f"DataFrame-Fehler: {e}"))
            
        return chunks

    async def extract_email(self, eml_bytes: bytes) -> List[Dict]:
        """Extrahiere E-Mail-Metadaten und Inhalt - GARANTIERT List[Dict]"""
        try:
            msg = BytesParser(policy=policy.default).parsebytes(eml_bytes)
            chunks = [{
                "content": json.dumps({
                    "from": msg["From"],
                    "to": msg["To"],
                    "subject": msg["Subject"],
                    "date": msg["Date"],
                }),
                "type": "metadata",
                "metadata": {"extraction_method": "email_header"}
            }]
            
            body = msg.get_body(preferencelist=("plain", "html"))
            if body:
                content = body.get_content()
                if content and content.strip():
                    # FIX: Korrekte Verarbeitung der split_text Rückgabe
                    text_chunks = self.split_text(content.strip())
                    chunks.extend([{
                        "content": chunk_text, 
                        "type": "text",
                        "metadata": {"extraction_method": "email_body"}
                    } for chunk_text in text_chunks])
                    
            return chunks
            
        except Exception as e:
            logger.error(f"❌ E-Mail-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"E-Mail-Extraktion fehlgeschlagen: {e}")]

    def extract_html(self, html: str) -> List[Dict]:
        """Extrahiere Text aus HTML - GARANTIERT List[Dict]"""
        try:
            if HAS_BS4:
                soup = BeautifulSoup(html, "html.parser")
                text = soup.get_text(separator="\n", strip=True)
            else:
                import re
                text = re.sub("<[^<]+?>", "", html)
                
            return self.extract_text(text, "html")
            
        except Exception as e:
            logger.error(f"❌ HTML-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"HTML-Extraktion fehlgeschlagen: {e}")]

    async def extract_archive(self, zip_bytes: bytes) -> List[Dict]:
        """Extrahiere Dateiliste aus ZIP - GARANTIERT List[Dict]"""
        try:
            with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
                files = zf.namelist()
                return [{
                    "content": f"Archive enthält: {name}",
                    "type": "metadata",
                    "metadata": {
                        "archived_file": name,
                        "extraction_method": "zip_listing"
                    }
                } for name in files if not name.endswith("/")]
                
        except Exception as e:
            logger.error(f"❌ Archiv-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Archiv-Extraktion fehlgeschlagen: {e}")]

    # === UTILITY METHODEN ===
    
    def split_text(self, text: str, chunk_size: int = None) -> List[str]:
        """
        Teile Text in Chunks mit dynamischer Größe
        
        Returns:
            List[str]: Liste von Text-Chunks (nicht Dict!)
        """
        chunk_size = chunk_size or self.chunk_size
        words = text.split()
        chunks = []
        current_chunk = []
        current_size = 0

        for word in words:
            current_chunk.append(word)
            current_size += len(word) + 1
            
            if current_size >= chunk_size:
                chunks.append(" ".join(current_chunk))
                current_chunk = []
                current_size = 0

        if current_chunk:
            chunks.append(" ".join(current_chunk))

        return chunks or [""]

    def _create_error_chunk(self, error_message: str) -> Dict:
        """Erstelle standardisierten Error-Chunk"""
        return {
            "content": f"[FEHLER: {error_message}]",
            "type": "error",
            "metadata": {
                "extraction_method": "error",
                "error_message": error_message,
                "timestamp": datetime.now().isoformat()
            }
        }

    def _validate_chunks(self, chunks: List[Dict]) -> List[Dict]:
        """Validiere und normalisiere Chunk-Format"""
        validated = []
        for i, chunk in enumerate(chunks):
            if not isinstance(chunk, dict):
                logger.warning(f"⚠️ Chunk {i} ist kein Dict: {type(chunk)}")
                chunk = {"content": str(chunk), "type": "text", "metadata": {}}
                
            # Stelle sicher, dass required fields vorhanden sind
            if "content" not in chunk:
                chunk["content"] = "[Leerer Chunk]"
            if "type" not in chunk:
                chunk["type"] = "text"
            if "metadata" not in chunk:
                chunk["metadata"] = {}
                
            validated.append(chunk)
            
        return validated

    def _parse_metadata(self, metadata: Union[str, dict]) -> dict:
        """DEPRECATED: Legacy method - use _parse_metadata_robust instead"""
        robust_metadata = self._parse_metadata_robust(metadata)
        return robust_metadata.dict(exclude_none=True)

    async def _save_chunks_to_db(self, conn, document_id: str, chunks: List[Dict]):
        """Speichere Chunks in Datenbank"""
        for idx, chunk in enumerate(chunks):
            await conn.execute(
                """
                INSERT INTO chunks (
                    document_id, chunk_index, content, content_type,
                    page_number, position, metadata
                ) VALUES ($1, $2, $3, $4, $5, $6, $7)
                ON CONFLICT (document_id, chunk_index) DO UPDATE
                SET content = EXCLUDED.content,
                    content_type = EXCLUDED.content_type
                """,
                document_id,
                idx,
                chunk["content"],
                chunk.get("type", "text"),
                chunk.get("page"),
                chunk.get("position"),
                json.dumps(chunk.get("metadata", {}))
            )

    async def _mark_task_failed(self, conn, task_id: str, doc_id: str, error_message: str):
        """Markiere Task als fehlgeschlagen"""
        await conn.execute(
            """UPDATE processing_queue 
               SET status = 'failed', error_message = $2, completed_at = NOW()
               WHERE id = $1""",
            task_id, error_message
        )
        await conn.execute(
            """UPDATE documents SET status = 'failed' WHERE id = $1""",
            doc_id
        )

    async def run(self):
        """Starte den OCR-Agent"""
        await self.connect()
        logger.info("🧠 Enhanced OCR Agent V2 gestartet")
        
        try:
            await self.process_queue()
        except KeyboardInterrupt:
            logger.info("🛑 OCR Agent wird beendet (KeyboardInterrupt)")
        except Exception as e:
            logger.error(f"💥 Unerwarteter Fehler: {e}")
            raise
        finally:
            await self.close()


def main():
    """Haupt-Funktion"""
    db_url = os.getenv(
        "DATABASE_URL",
        "postgresql://semanticuser:semantic2024@postgres:5432/semantic_doc_finder"
    )
    
    logger.info(f"🚀 Starte OCR Agent mit DB: {db_url.split('@')[1] if '@' in db_url else 'localhost'}")
    
    agent = OCRAgentV2(db_url)
    asyncio.run(agent.run())


if __name__ == "__main__":
    main(), cleaned):
            logger.warning(f"⚠️ Potentiell ungültiger base64 content (erste 50 chars): {cleaned[:50]}")
        
        return cleaned

    def _validate_file_path(self, file_path: str) -> bool:
        """
        Umfassende Validierung eines file paths
        
        Args:
            file_path: Path to validate
            
        Returns:
            bool: True wenn file path gültig und lesbar ist
        """
        if not file_path or not isinstance(file_path, str):
            return False
            
        try:
            # Prüfe ob Datei existiert
            if not os.path.exists(file_path):
                logger.debug(f"File path existiert nicht: {file_path}")
                return False
            
            # Prüfe ob es eine Datei ist (nicht directory)
            if not os.path.isfile(file_path):
                logger.debug(f"Path ist keine Datei: {file_path}")
                return False
                
            # Prüfe Dateigröße (nicht 0 und nicht zu groß)
            file_size = os.path.getsize(file_path)
            if file_size == 0:
                logger.debug(f"Datei ist leer: {file_path}")
                return False
                
            # Prüfe maximale Dateigröße (2GB limit)
            MAX_FILE_SIZE_ABSOLUTE = 2 * 1024 * 1024 * 1024  # 2GB
            if file_size > MAX_FILE_SIZE_ABSOLUTE:
                logger.warning(f"⚠️ Datei zu groß: {file_size} bytes (Max: {MAX_FILE_SIZE_ABSOLUTE})")
                return False
            
            # Prüfe Lese-Berechtigung
            if not os.access(file_path, os.R_OK):
                logger.debug(f"Keine Lese-Berechtigung für: {file_path}")
                return False
                
            logger.debug(f"✅ File path validiert: {file_path} ({file_size} bytes)")
            return True
            
        except Exception as e:
            logger.debug(f"File path validation fehlgeschlagen: {e}")
            return False

    def _validate_base64_content(self, content: str) -> bool:
        """
        Umfassende Validierung von base64 content
        
        Args:
            content: Base64 string to validate
            
        Returns:
            bool: True wenn content gültiger base64 ist
        """
        if not content or not isinstance(content, str):
            return False
            
        try:
            # Prüfe minimale Länge
            if len(content.strip()) < 4:
                logger.debug(f"Base64 content zu kurz: {len(content)} chars")
                return False
                
            # Prüfe maximale Länge (500MB als base64 ≈ 666MB string)
            MAX_BASE64_LENGTH = 700 * 1024 * 1024  # 700MB string
            if len(content) > MAX_BASE64_LENGTH:
                logger.warning(f"⚠️ Base64 content zu groß: {len(content)} chars")
                return False
            
            # Bereinige content für test decode
            cleaned_content = self._clean_base64_content(content)
            
            # Prüfe base64 format pattern
            if not re.match(r'^[A-Za-z0-9+/]*={0,2}

    def _parse_metadata_robust(self, raw_metadata: Union[str, dict]) -> DocumentMetadata:
        """
        Robustes Parsing von Metadaten zu strukturiertem Format
        
        Args:
            raw_metadata: Raw metadata von database
            
        Returns:
            DocumentMetadata: Strukturierte und validierte Metadaten
        """
        try:
            # Parse zu dict falls string
            if isinstance(raw_metadata, str):
                try:
                    metadata_dict = json.loads(raw_metadata)
                except json.JSONDecodeError as e:
                    logger.warning(f"⚠️ JSON-Parsing fehlgeschlagen: {e}")
                    metadata_dict = {}
            elif isinstance(raw_metadata, dict):
                metadata_dict = raw_metadata
            else:
                logger.warning(f"⚠️ Unerwarteter metadata type: {type(raw_metadata)}")
                metadata_dict = {}
            
            # Erstelle DocumentMetadata mit validation
            metadata = DocumentMetadata(**metadata_dict)
            
            # Additional validation/normalization
            if not metadata.effective_file_content and not metadata.effective_file_path:
                logger.warning("⚠️ Keine file_content oder file_path in metadata gefunden")
            
            return metadata
            
        except Exception as e:
            logger.error(f"❌ Metadata-Parsing fehlgeschlagen: {e}")
            # Fallback zu default metadata
            return DocumentMetadata(source="metadata_parse_error")

    async def _get_file_bytes(self, file_content: Optional[str], file_path: Optional[str]) -> bytes:
        """DEPRECATED: Legacy method - use _get_file_bytes_robust instead"""
        # Erstelle legacy DocumentMetadata für compatibility
        legacy_metadata = DocumentMetadata(
            file_content=file_content,
            file_path=file_path
        )
        return await self._get_file_bytes_robust(legacy_metadata)

    async def _dispatch_extraction(self, file_type: str, file_bytes: bytes, filename: str) -> List[Dict]:
        """
        Dispatche zu spezifischer Extraktionsmethode
        
        Returns:
            List[Dict]: GARANTIERT List von Chunk-Dictionaries
        """
        try:
            if file_type == "pdf":
                return await self.extract_pdf(file_bytes, filename)
            elif file_type in ["png", "jpg", "jpeg", "tiff", "tif"]:
                return await self.extract_image(file_bytes)
            elif file_type in ["txt", "md", "log"]:
                content = file_bytes.decode("utf-8", errors="ignore")
                return self.extract_text(content, file_type)
            elif file_type in ["docx", "doc"]:
                return await self.extract_docx(file_bytes)
            elif file_type in ["xlsx", "xls", "csv"]:
                return await self.extract_spreadsheet(file_bytes, file_type)
            elif file_type == "eml":
                return await self.extract_email(file_bytes)
            elif file_type == "html":
                content = file_bytes.decode("utf-8", errors="ignore")
                return self.extract_html(content)
            elif file_type == "zip":
                return await self.extract_archive(file_bytes)
            else:
                # Fallback zu text
                content = file_bytes.decode("utf-8", errors="ignore")
                return self.extract_text(content, file_type)
                
        except Exception as e:
            logger.error(f"❌ Extraktion für {file_type} fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Extraktion für {file_type} fehlgeschlagen: {e}")]

    # === STREAMING METHODEN ===
    
    async def _extract_content_streaming(self, file_path: str, file_type: str, filename: str) -> List[Dict]:
        """Streaming-Extraktion für große Dateien - GARANTIERT List[Dict]"""
        try:
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            logger.info(f"🚀 Streaming für {filename}: {file_size_mb:.1f}MB")
            
            if file_type == "pdf":
                return await self._stream_large_pdf(file_path)
            elif file_type in ["txt", "md", "log"]:
                return await self._stream_large_text(file_path)
            elif file_type in ["docx", "doc"] and HAS_DOCX:
                return await self._stream_large_docx(file_path)
            else:
                # Fallback: normales Laden als Text
                logger.warning(f"⚠️ Kein Streaming für {file_type} - fallback zu normal")
                with open(file_path, "rb") as f:
                    file_bytes = f.read()
                content = file_bytes.decode("utf-8", errors="ignore")
                return self.extract_text(content, file_type)
                
        except Exception as e:
            logger.error(f"❌ Streaming fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Streaming fehlgeschlagen: {e}")]

    async def _stream_large_pdf(self, file_path: str) -> List[Dict]:
        """Verarbeite PDF seitenweise - GARANTIERT List[Dict]"""
        chunks = []
        try:
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            dpi = 150 if file_size_mb > 50 else 300
            tesseract_config = self.tesseract_config_fast if file_size_mb > 100 else self.tesseract_config
            
            logger.info(f"📄 PDF-Streaming: {file_path} ({file_size_mb:.1f} MB, DPI: {dpi})")

            images = pdf2image.convert_from_path(file_path, dpi=dpi)
            total_pages = len(images)

            for page_num, image in enumerate(images, 1):
                if total_pages > 50 and page_num % 10 == 0:
                    logger.info(f"📄 Fortschritt: {page_num}/{total_pages} ({page_num/total_pages*100:.1f}%)")
                
                text = pytesseract.image_to_string(
                    image,
                    lang=os.getenv("TESSERACT_LANG", "deu+eng"),
                    config=tesseract_config
                )

                if text.strip():
                    pdf_chunk_size = int(self.chunk_size * 1.5) if file_size_mb > 100 else self.chunk_size
                    page_text_chunks = self.split_text(text, pdf_chunk_size)
                    
                    for chunk_text in page_text_chunks:
                        chunks.append({
                            "content": chunk_text,
                            "type": "text",
                            "page": page_num,
                            "metadata": {
                                "extraction_method": "pdf_streaming",
                                "dpi": dpi,
                                "chunk_size_used": pdf_chunk_size
                            }
                        })

                del image
                if page_num % 5 == 0:
                    gc.collect()
                    
        except Exception as e:
            logger.error(f"❌ PDF-Streaming fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"PDF-Streaming fehlgeschlagen: {e}")]
            
        logger.info(f"✅ PDF verarbeitet: {len(chunks)} Chunks aus {total_pages} Seiten")
        return chunks

    async def _stream_large_text(self, file_path: str) -> List[Dict]:
        """Verarbeite große Textdateien line-by-line - GARANTIERT List[Dict]"""
        chunks = []
        try:
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            dynamic_chunk_size = min(self.chunk_size * 2, 2000) if file_size_mb > 100 else self.chunk_size
            
            logger.info(f"📦 Text-Streaming: {file_path} ({file_size_mb:.1f} MB), Chunk-Größe: {dynamic_chunk_size}")

            current_chunk = []
            current_size = 0
            line_count = 0
            
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                for line in f:
                    line = line.strip() + "\n"
                    if not line.strip():
                        continue
                        
                    current_chunk.append(line)
                    current_size += len(line)
                    line_count += 1

                    if current_size >= dynamic_chunk_size:
                        chunks.append({
                            "content": "".join(current_chunk),
                            "type": "text",
                            "metadata": {
                                "extraction_method": "text_streaming",
                                "lines_in_chunk": len(current_chunk),
                                "chunk_size_used": dynamic_chunk_size
                            }
                        })
                        current_chunk = []
                        current_size = 0

                        if len(chunks) % 10 == 0:
                            logger.info(f"📊 {len(chunks)} Chunks verarbeitet ({line_count} Zeilen)")
                            
                # Letzter Chunk
                if current_chunk:
                    chunks.append({
                        "content": "".join(current_chunk),
                        "type": "text",
                        "metadata": {
                            "extraction_method": "text_streaming",
                            "lines_in_chunk": len(current_chunk),
                            "chunk_size_used": dynamic_chunk_size,
                            "final_chunk": True
                        }
                    })
                    
        except Exception as e:
            logger.error(f"❌ Text-Streaming fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Text-Streaming fehlgeschlagen: {e}")]
            
        logger.info(f"✅ Textdatei verarbeitet: {len(chunks)} Chunks")
        return chunks

    async def _stream_large_docx(self, file_path: str) -> List[Dict]:
        """Verarbeite große DOCX-Dateien - GARANTIERT List[Dict]"""
        chunks = []
        try:
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            logger.info(f"📄 DOCX-Streaming: {file_path} ({file_size_mb:.1f} MB)")
            
            if not HAS_DOCX:
                return [self._create_error_chunk("DOCX-Support nicht verfügbar")]
            
            doc = DocxDocument(file_path)
            full_text = "\n".join([para.text for para in doc.paragraphs])
            
            docx_chunk_size = int(self.chunk_size * 1.3)
            text_chunks = self.split_text(full_text, docx_chunk_size)
            
            for chunk_text in text_chunks:
                chunks.append({
                    "content": chunk_text,
                    "type": "text",
                    "metadata": {
                        "extraction_method": "docx_streaming",
                        "chunk_size_used": docx_chunk_size
                    }
                })
                
            # Extrahiere Tabellen
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                    
                if table_data:
                    chunks.append({
                        "content": json.dumps(table_data),
                        "type": "table",
                        "metadata": {
                            "extraction_method": "docx_table",
                            "table_index": table_idx
                        }
                    })
                    
        except Exception as e:
            logger.error(f"❌ DOCX-Streaming fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"DOCX-Streaming fehlgeschlagen: {e}")]
            
        logger.info(f"✅ DOCX verarbeitet: {len(chunks)} Chunks")
        return chunks

    # === IN-MEMORY EXTRAKTIONSMETHODEN ===
    
    async def extract_pdf(self, pdf_bytes: bytes, filename: str) -> List[Dict]:
        """Extrahiere Text aus PDF mit OCR (In-Memory) - GARANTIERT List[Dict]"""
        chunks = []
        try:
            images = pdf2image.convert_from_bytes(pdf_bytes, dpi=300)
            logger.info(f"📄 PDF In-Memory: {filename} ({len(images)} Seiten)")
            
            for page_num, image in enumerate(images, 1):
                text = pytesseract.image_to_string(
                    image,
                    lang=os.getenv("TESSERACT_LANG", "deu+eng"),
                    config=self.tesseract_config
                )
                
                if text.strip():
                    page_chunks = self.split_text(text)
                    for chunk_text in page_chunks:
                        chunks.append({
                            "content": chunk_text,
                            "type": "text",
                            "page": page_num,
                            "metadata": {
                                "extraction_method": "pdf_in_memory",
                                "chunk_size_used": self.chunk_size
                            }
                        })
                        
                del image
                if page_num % 5 == 0:
                    gc.collect()
                    
        except Exception as e:
            logger.error(f"❌ PDF-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"PDF-Extraktion fehlgeschlagen: {e}")]
            
        return chunks if chunks else [self._create_error_chunk("Kein Text aus PDF extrahiert")]

    async def extract_image(self, image_bytes: bytes) -> List[Dict]:
        """Extrahiere Text aus einem Bild - GARANTIERT List[Dict]"""
        try:
            image = Image.open(io.BytesIO(image_bytes))
            text = pytesseract.image_to_string(
                image,
                lang=os.getenv("TESSERACT_LANG", "deu+eng"),
                config=self.tesseract_config
            )
            
            if not text.strip():
                return [{"content": "[Bild ohne erkannten Text]", "type": "text", "metadata": {"extraction_method": "image_ocr"}}]
                
            chunks = self.split_text(text)
            return [{
                "content": chunk_text,
                "type": "text",
                "page": 1,
                "metadata": {
                    "extraction_method": "image_ocr",
                    "chunk_size_used": self.chunk_size
                }
            } for chunk_text in chunks]
            
        except Exception as e:
            logger.error(f"❌ Bild-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Bild-Extraktion fehlgeschlagen: {e}")]

    async def extract_docx(self, docx_bytes: bytes) -> List[Dict]:
        """Extrahiere Text aus DOCX (In-Memory) - GARANTIERT List[Dict]"""
        if not HAS_DOCX:
            return [self._create_error_chunk("DOCX-Support nicht verfügbar (python-docx nicht installiert)")]
            
        try:
            doc = DocxDocument(io.BytesIO(docx_bytes))
            full_text = "\n".join([para.text for para in doc.paragraphs])
            
            docx_chunk_size = int(self.chunk_size * 1.3)
            chunks = self.split_text(full_text, docx_chunk_size)
            
            result = [{
                "content": chunk_text,
                "type": "text",
                "metadata": {
                    "extraction_method": "docx_in_memory",
                    "chunk_size_used": docx_chunk_size
                }
            } for chunk_text in chunks]
            
            # Extrahiere Tabellen
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                    
                if table_data:
                    result.append({
                        "content": json.dumps(table_data),
                        "type": "table",
                        "metadata": {
                            "extraction_method": "docx_table",
                            "table_index": table_idx
                        }
                    })
                    
            return result if result else [self._create_error_chunk("Kein Inhalt aus DOCX extrahiert")]
            
        except Exception as e:
            logger.error(f"❌ DOCX-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"DOCX-Extraktion fehlgeschlagen: {e}")]

    def extract_text(self, text: str, file_type: str = "txt") -> List[Dict]:
        """Extrahiere Text mit adaptiver Chunk-Größe - GARANTIERT List[Dict]"""
        if not text.strip():
            return [{"content": "[Leerer Text]", "type": "text", "metadata": {"extraction_method": "text_in_memory"}}]
            
        try:
            # Adaptive Chunk-Größen basierend auf Dateityp
            if file_type in ['csv', 'tsv']:
                chunk_size = self.chunk_size // 2
            elif file_type in ['md', 'txt']:
                chunk_size = int(self.chunk_size * 1.2)
            else:
                chunk_size = self.chunk_size
                
            chunks = self.split_text(text, chunk_size)
            return [{
                "content": chunk_text,
                "type": "text",
                "metadata": {
                    "extraction_method": "text_in_memory",
                    "chunk_size_used": chunk_size,
                    "file_type": file_type
                }
            } for chunk_text in chunks]
            
        except Exception as e:
            logger.error(f"❌ Text-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Text-Extraktion fehlgeschlagen: {e}")]

    async def extract_spreadsheet(self, file_bytes: bytes, file_type: str) -> List[Dict]:
        """Extrahiere Tabellen aus Excel/CSV - GARANTIERT List[Dict]"""
        if not HAS_PANDAS:
            return [self._create_error_chunk("Tabellen-Support nicht verfügbar (pandas nicht installiert)")]
            
        try:
            if file_type == "csv":
                df = pd.read_csv(io.BytesIO(file_bytes))
            else:
                df = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
                
            chunks = []
            if isinstance(df, dict):
                for sheet_name, sheet_df in df.items():
                    chunks.extend(self._process_dataframe(sheet_df, {"sheet": sheet_name}))
            else:
                chunks.extend(self._process_dataframe(df))
                
            return chunks if chunks else [self._create_error_chunk("Keine Tabellendaten extrahiert")]
            
        except Exception as e:
            logger.error(f"❌ Tabelle fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Tabellen-Extraktion fehlgeschlagen: {e}")]

    def _process_dataframe(self, df: "pd.DataFrame", metadata: dict = None) -> List[Dict]:
        """Verarbeite DataFrame in Chunks - GARANTIERT List[Dict]"""
        metadata = metadata or {}
        chunks = []
        
        try:
            for i in range(0, len(df), 50):
                chunk_df = df.iloc[i:i + 50]
                chunks.append({
                    "content": chunk_df.to_json(orient="records"),
                    "type": "table",
                    "metadata": {
                        **metadata, 
                        "row_start": i,
                        "extraction_method": "pandas"
                    }
                })
        except Exception as e:
            logger.error(f"❌ DataFrame-Verarbeitung fehlgeschlagen: {e}")
            chunks.append(self._create_error_chunk(f"DataFrame-Fehler: {e}"))
            
        return chunks

    async def extract_email(self, eml_bytes: bytes) -> List[Dict]:
        """Extrahiere E-Mail-Metadaten und Inhalt - GARANTIERT List[Dict]"""
        try:
            msg = BytesParser(policy=policy.default).parsebytes(eml_bytes)
            chunks = [{
                "content": json.dumps({
                    "from": msg["From"],
                    "to": msg["To"],
                    "subject": msg["Subject"],
                    "date": msg["Date"],
                }),
                "type": "metadata",
                "metadata": {"extraction_method": "email_header"}
            }]
            
            body = msg.get_body(preferencelist=("plain", "html"))
            if body:
                content = body.get_content()
                if content and content.strip():
                    # FIX: Korrekte Verarbeitung der split_text Rückgabe
                    text_chunks = self.split_text(content.strip())
                    chunks.extend([{
                        "content": chunk_text, 
                        "type": "text",
                        "metadata": {"extraction_method": "email_body"}
                    } for chunk_text in text_chunks])
                    
            return chunks
            
        except Exception as e:
            logger.error(f"❌ E-Mail-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"E-Mail-Extraktion fehlgeschlagen: {e}")]

    def extract_html(self, html: str) -> List[Dict]:
        """Extrahiere Text aus HTML - GARANTIERT List[Dict]"""
        try:
            if HAS_BS4:
                soup = BeautifulSoup(html, "html.parser")
                text = soup.get_text(separator="\n", strip=True)
            else:
                import re
                text = re.sub("<[^<]+?>", "", html)
                
            return self.extract_text(text, "html")
            
        except Exception as e:
            logger.error(f"❌ HTML-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"HTML-Extraktion fehlgeschlagen: {e}")]

    async def extract_archive(self, zip_bytes: bytes) -> List[Dict]:
        """Extrahiere Dateiliste aus ZIP - GARANTIERT List[Dict]"""
        try:
            with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
                files = zf.namelist()
                return [{
                    "content": f"Archive enthält: {name}",
                    "type": "metadata",
                    "metadata": {
                        "archived_file": name,
                        "extraction_method": "zip_listing"
                    }
                } for name in files if not name.endswith("/")]
                
        except Exception as e:
            logger.error(f"❌ Archiv-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Archiv-Extraktion fehlgeschlagen: {e}")]

    # === UTILITY METHODEN ===
    
    def split_text(self, text: str, chunk_size: int = None) -> List[str]:
        """
        Teile Text in Chunks mit dynamischer Größe
        
        Returns:
            List[str]: Liste von Text-Chunks (nicht Dict!)
        """
        chunk_size = chunk_size or self.chunk_size
        words = text.split()
        chunks = []
        current_chunk = []
        current_size = 0

        for word in words:
            current_chunk.append(word)
            current_size += len(word) + 1
            
            if current_size >= chunk_size:
                chunks.append(" ".join(current_chunk))
                current_chunk = []
                current_size = 0

        if current_chunk:
            chunks.append(" ".join(current_chunk))

        return chunks or [""]

    def _create_error_chunk(self, error_message: str) -> Dict:
        """Erstelle standardisierten Error-Chunk"""
        return {
            "content": f"[FEHLER: {error_message}]",
            "type": "error",
            "metadata": {
                "extraction_method": "error",
                "error_message": error_message,
                "timestamp": datetime.now().isoformat()
            }
        }

    def _validate_chunks(self, chunks: List[Dict]) -> List[Dict]:
        """Validiere und normalisiere Chunk-Format"""
        validated = []
        for i, chunk in enumerate(chunks):
            if not isinstance(chunk, dict):
                logger.warning(f"⚠️ Chunk {i} ist kein Dict: {type(chunk)}")
                chunk = {"content": str(chunk), "type": "text", "metadata": {}}
                
            # Stelle sicher, dass required fields vorhanden sind
            if "content" not in chunk:
                chunk["content"] = "[Leerer Chunk]"
            if "type" not in chunk:
                chunk["type"] = "text"
            if "metadata" not in chunk:
                chunk["metadata"] = {}
                
            validated.append(chunk)
            
        return validated

    def _parse_metadata(self, metadata: Union[str, dict]) -> dict:
        """Parse metadata sicher zu dict"""
        if isinstance(metadata, dict):
            return metadata
        if isinstance(metadata, str):
            try:
                return json.loads(metadata)
            except:
                return {}
        return {}

    async def _save_chunks_to_db(self, conn, document_id: str, chunks: List[Dict]):
        """Speichere Chunks in Datenbank"""
        for idx, chunk in enumerate(chunks):
            await conn.execute(
                """
                INSERT INTO chunks (
                    document_id, chunk_index, content, content_type,
                    page_number, position, metadata
                ) VALUES ($1, $2, $3, $4, $5, $6, $7)
                ON CONFLICT (document_id, chunk_index) DO UPDATE
                SET content = EXCLUDED.content,
                    content_type = EXCLUDED.content_type
                """,
                document_id,
                idx,
                chunk["content"],
                chunk.get("type", "text"),
                chunk.get("page"),
                chunk.get("position"),
                json.dumps(chunk.get("metadata", {}))
            )

    async def _mark_task_failed(self, conn, task_id: str, doc_id: str, error_message: str):
        """Markiere Task als fehlgeschlagen"""
        await conn.execute(
            """UPDATE processing_queue 
               SET status = 'failed', error_message = $2, completed_at = NOW()
               WHERE id = $1""",
            task_id, error_message
        )
        await conn.execute(
            """UPDATE documents SET status = 'failed' WHERE id = $1""",
            doc_id
        )

    async def run(self):
        """Starte den OCR-Agent"""
        await self.connect()
        logger.info("🧠 Enhanced OCR Agent V2 gestartet")
        
        try:
            await self.process_queue()
        except KeyboardInterrupt:
            logger.info("🛑 OCR Agent wird beendet (KeyboardInterrupt)")
        except Exception as e:
            logger.error(f"💥 Unerwarteter Fehler: {e}")
            raise
        finally:
            await self.close()


def main():
    """Haupt-Funktion"""
    db_url = os.getenv(
        "DATABASE_URL",
        "postgresql://semanticuser:semantic2024@postgres:5432/semantic_doc_finder"
    )
    
    logger.info(f"🚀 Starte OCR Agent mit DB: {db_url.split('@')[1] if '@' in db_url else 'localhost'}")
    
    agent = OCRAgentV2(db_url)
    asyncio.run(agent.run())


if __name__ == "__main__":
    main(), cleaned_content):
                logger.debug(f"Base64 format ungültig (erste 50 chars): {cleaned_content[:50]}")
                return False
            
            # Test decode (nur erste 1000 chars für performance)
            test_content = cleaned_content[:1000] if len(cleaned_content) > 1000 else cleaned_content
            
            # Padding für test hinzufügen falls nötig
            if len(test_content) % 4 != 0:
                test_content += '=' * (4 - len(test_content) % 4)
                
            decoded = base64.b64decode(test_content)
            
            # Prüfe dass decode nicht leer ist
            if not decoded:
                logger.debug("Base64 dekodiert zu leerem content")
                return False
                
            logger.debug(f"✅ Base64 content validiert: {len(content)} chars → {len(decoded)} test bytes")
            return True
            
        except Exception as e:
            logger.debug(f"Base64 validation fehlgeschlagen: {e}")
            return False

    def _parse_metadata_robust(self, raw_metadata: Union[str, dict]) -> DocumentMetadata:
        """
        Robustes Parsing von Metadaten zu strukturiertem Format
        
        Args:
            raw_metadata: Raw metadata von database
            
        Returns:
            DocumentMetadata: Strukturierte und validierte Metadaten
        """
        try:
            # Parse zu dict falls string
            if isinstance(raw_metadata, str):
                try:
                    metadata_dict = json.loads(raw_metadata)
                except json.JSONDecodeError as e:
                    logger.warning(f"⚠️ JSON-Parsing fehlgeschlagen: {e}")
                    metadata_dict = {}
            elif isinstance(raw_metadata, dict):
                metadata_dict = raw_metadata
            else:
                logger.warning(f"⚠️ Unerwarteter metadata type: {type(raw_metadata)}")
                metadata_dict = {}
            
            # Erstelle DocumentMetadata mit validation
            metadata = DocumentMetadata(**metadata_dict)
            
            # Additional validation/normalization
            if not metadata.effective_file_content and not metadata.effective_file_path:
                logger.warning("⚠️ Keine file_content oder file_path in metadata gefunden")
            
            return metadata
            
        except Exception as e:
            logger.error(f"❌ Metadata-Parsing fehlgeschlagen: {e}")
            # Fallback zu default metadata
            return DocumentMetadata(source="metadata_parse_error")

    async def _get_file_bytes(self, file_content: Optional[str], file_path: Optional[str]) -> bytes:
        """DEPRECATED: Legacy method - use _get_file_bytes_robust instead"""
        # Erstelle legacy DocumentMetadata für compatibility
        legacy_metadata = DocumentMetadata(
            file_content=file_content,
            file_path=file_path
        )
        return await self._get_file_bytes_robust(legacy_metadata)

    async def _dispatch_extraction(self, file_type: str, file_bytes: bytes, filename: str) -> List[Dict]:
        """
        Dispatche zu spezifischer Extraktionsmethode
        
        Returns:
            List[Dict]: GARANTIERT List von Chunk-Dictionaries
        """
        try:
            if file_type == "pdf":
                return await self.extract_pdf(file_bytes, filename)
            elif file_type in ["png", "jpg", "jpeg", "tiff", "tif"]:
                return await self.extract_image(file_bytes)
            elif file_type in ["txt", "md", "log"]:
                content = file_bytes.decode("utf-8", errors="ignore")
                return self.extract_text(content, file_type)
            elif file_type in ["docx", "doc"]:
                return await self.extract_docx(file_bytes)
            elif file_type in ["xlsx", "xls", "csv"]:
                return await self.extract_spreadsheet(file_bytes, file_type)
            elif file_type == "eml":
                return await self.extract_email(file_bytes)
            elif file_type == "html":
                content = file_bytes.decode("utf-8", errors="ignore")
                return self.extract_html(content)
            elif file_type == "zip":
                return await self.extract_archive(file_bytes)
            else:
                # Fallback zu text
                content = file_bytes.decode("utf-8", errors="ignore")
                return self.extract_text(content, file_type)
                
        except Exception as e:
            logger.error(f"❌ Extraktion für {file_type} fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Extraktion für {file_type} fehlgeschlagen: {e}")]

    # === STREAMING METHODEN ===
    
    async def _extract_content_streaming(self, file_path: str, file_type: str, filename: str) -> List[Dict]:
        """Streaming-Extraktion für große Dateien - GARANTIERT List[Dict]"""
        try:
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            logger.info(f"🚀 Streaming für {filename}: {file_size_mb:.1f}MB")
            
            if file_type == "pdf":
                return await self._stream_large_pdf(file_path)
            elif file_type in ["txt", "md", "log"]:
                return await self._stream_large_text(file_path)
            elif file_type in ["docx", "doc"] and HAS_DOCX:
                return await self._stream_large_docx(file_path)
            else:
                # Fallback: normales Laden als Text
                logger.warning(f"⚠️ Kein Streaming für {file_type} - fallback zu normal")
                with open(file_path, "rb") as f:
                    file_bytes = f.read()
                content = file_bytes.decode("utf-8", errors="ignore")
                return self.extract_text(content, file_type)
                
        except Exception as e:
            logger.error(f"❌ Streaming fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Streaming fehlgeschlagen: {e}")]

    async def _stream_large_pdf(self, file_path: str) -> List[Dict]:
        """Verarbeite PDF seitenweise - GARANTIERT List[Dict]"""
        chunks = []
        try:
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            dpi = 150 if file_size_mb > 50 else 300
            tesseract_config = self.tesseract_config_fast if file_size_mb > 100 else self.tesseract_config
            
            logger.info(f"📄 PDF-Streaming: {file_path} ({file_size_mb:.1f} MB, DPI: {dpi})")

            images = pdf2image.convert_from_path(file_path, dpi=dpi)
            total_pages = len(images)

            for page_num, image in enumerate(images, 1):
                if total_pages > 50 and page_num % 10 == 0:
                    logger.info(f"📄 Fortschritt: {page_num}/{total_pages} ({page_num/total_pages*100:.1f}%)")
                
                text = pytesseract.image_to_string(
                    image,
                    lang=os.getenv("TESSERACT_LANG", "deu+eng"),
                    config=tesseract_config
                )

                if text.strip():
                    pdf_chunk_size = int(self.chunk_size * 1.5) if file_size_mb > 100 else self.chunk_size
                    page_text_chunks = self.split_text(text, pdf_chunk_size)
                    
                    for chunk_text in page_text_chunks:
                        chunks.append({
                            "content": chunk_text,
                            "type": "text",
                            "page": page_num,
                            "metadata": {
                                "extraction_method": "pdf_streaming",
                                "dpi": dpi,
                                "chunk_size_used": pdf_chunk_size
                            }
                        })

                del image
                if page_num % 5 == 0:
                    gc.collect()
                    
        except Exception as e:
            logger.error(f"❌ PDF-Streaming fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"PDF-Streaming fehlgeschlagen: {e}")]
            
        logger.info(f"✅ PDF verarbeitet: {len(chunks)} Chunks aus {total_pages} Seiten")
        return chunks

    async def _stream_large_text(self, file_path: str) -> List[Dict]:
        """Verarbeite große Textdateien line-by-line - GARANTIERT List[Dict]"""
        chunks = []
        try:
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            dynamic_chunk_size = min(self.chunk_size * 2, 2000) if file_size_mb > 100 else self.chunk_size
            
            logger.info(f"📦 Text-Streaming: {file_path} ({file_size_mb:.1f} MB), Chunk-Größe: {dynamic_chunk_size}")

            current_chunk = []
            current_size = 0
            line_count = 0
            
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                for line in f:
                    line = line.strip() + "\n"
                    if not line.strip():
                        continue
                        
                    current_chunk.append(line)
                    current_size += len(line)
                    line_count += 1

                    if current_size >= dynamic_chunk_size:
                        chunks.append({
                            "content": "".join(current_chunk),
                            "type": "text",
                            "metadata": {
                                "extraction_method": "text_streaming",
                                "lines_in_chunk": len(current_chunk),
                                "chunk_size_used": dynamic_chunk_size
                            }
                        })
                        current_chunk = []
                        current_size = 0

                        if len(chunks) % 10 == 0:
                            logger.info(f"📊 {len(chunks)} Chunks verarbeitet ({line_count} Zeilen)")
                            
                # Letzter Chunk
                if current_chunk:
                    chunks.append({
                        "content": "".join(current_chunk),
                        "type": "text",
                        "metadata": {
                            "extraction_method": "text_streaming",
                            "lines_in_chunk": len(current_chunk),
                            "chunk_size_used": dynamic_chunk_size,
                            "final_chunk": True
                        }
                    })
                    
        except Exception as e:
            logger.error(f"❌ Text-Streaming fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Text-Streaming fehlgeschlagen: {e}")]
            
        logger.info(f"✅ Textdatei verarbeitet: {len(chunks)} Chunks")
        return chunks

    async def _stream_large_docx(self, file_path: str) -> List[Dict]:
        """Verarbeite große DOCX-Dateien - GARANTIERT List[Dict]"""
        chunks = []
        try:
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            logger.info(f"📄 DOCX-Streaming: {file_path} ({file_size_mb:.1f} MB)")
            
            if not HAS_DOCX:
                return [self._create_error_chunk("DOCX-Support nicht verfügbar")]
            
            doc = DocxDocument(file_path)
            full_text = "\n".join([para.text for para in doc.paragraphs])
            
            docx_chunk_size = int(self.chunk_size * 1.3)
            text_chunks = self.split_text(full_text, docx_chunk_size)
            
            for chunk_text in text_chunks:
                chunks.append({
                    "content": chunk_text,
                    "type": "text",
                    "metadata": {
                        "extraction_method": "docx_streaming",
                        "chunk_size_used": docx_chunk_size
                    }
                })
                
            # Extrahiere Tabellen
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                    
                if table_data:
                    chunks.append({
                        "content": json.dumps(table_data),
                        "type": "table",
                        "metadata": {
                            "extraction_method": "docx_table",
                            "table_index": table_idx
                        }
                    })
                    
        except Exception as e:
            logger.error(f"❌ DOCX-Streaming fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"DOCX-Streaming fehlgeschlagen: {e}")]
            
        logger.info(f"✅ DOCX verarbeitet: {len(chunks)} Chunks")
        return chunks

    # === IN-MEMORY EXTRAKTIONSMETHODEN ===
    
    async def extract_pdf(self, pdf_bytes: bytes, filename: str) -> List[Dict]:
        """Extrahiere Text aus PDF mit OCR (In-Memory) - GARANTIERT List[Dict]"""
        chunks = []
        try:
            images = pdf2image.convert_from_bytes(pdf_bytes, dpi=300)
            logger.info(f"📄 PDF In-Memory: {filename} ({len(images)} Seiten)")
            
            for page_num, image in enumerate(images, 1):
                text = pytesseract.image_to_string(
                    image,
                    lang=os.getenv("TESSERACT_LANG", "deu+eng"),
                    config=self.tesseract_config
                )
                
                if text.strip():
                    page_chunks = self.split_text(text)
                    for chunk_text in page_chunks:
                        chunks.append({
                            "content": chunk_text,
                            "type": "text",
                            "page": page_num,
                            "metadata": {
                                "extraction_method": "pdf_in_memory",
                                "chunk_size_used": self.chunk_size
                            }
                        })
                        
                del image
                if page_num % 5 == 0:
                    gc.collect()
                    
        except Exception as e:
            logger.error(f"❌ PDF-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"PDF-Extraktion fehlgeschlagen: {e}")]
            
        return chunks if chunks else [self._create_error_chunk("Kein Text aus PDF extrahiert")]

    async def extract_image(self, image_bytes: bytes) -> List[Dict]:
        """Extrahiere Text aus einem Bild - GARANTIERT List[Dict]"""
        try:
            image = Image.open(io.BytesIO(image_bytes))
            text = pytesseract.image_to_string(
                image,
                lang=os.getenv("TESSERACT_LANG", "deu+eng"),
                config=self.tesseract_config
            )
            
            if not text.strip():
                return [{"content": "[Bild ohne erkannten Text]", "type": "text", "metadata": {"extraction_method": "image_ocr"}}]
                
            chunks = self.split_text(text)
            return [{
                "content": chunk_text,
                "type": "text",
                "page": 1,
                "metadata": {
                    "extraction_method": "image_ocr",
                    "chunk_size_used": self.chunk_size
                }
            } for chunk_text in chunks]
            
        except Exception as e:
            logger.error(f"❌ Bild-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Bild-Extraktion fehlgeschlagen: {e}")]

    async def extract_docx(self, docx_bytes: bytes) -> List[Dict]:
        """Extrahiere Text aus DOCX (In-Memory) - GARANTIERT List[Dict]"""
        if not HAS_DOCX:
            return [self._create_error_chunk("DOCX-Support nicht verfügbar (python-docx nicht installiert)")]
            
        try:
            doc = DocxDocument(io.BytesIO(docx_bytes))
            full_text = "\n".join([para.text for para in doc.paragraphs])
            
            docx_chunk_size = int(self.chunk_size * 1.3)
            chunks = self.split_text(full_text, docx_chunk_size)
            
            result = [{
                "content": chunk_text,
                "type": "text",
                "metadata": {
                    "extraction_method": "docx_in_memory",
                    "chunk_size_used": docx_chunk_size
                }
            } for chunk_text in chunks]
            
            # Extrahiere Tabellen
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                    
                if table_data:
                    result.append({
                        "content": json.dumps(table_data),
                        "type": "table",
                        "metadata": {
                            "extraction_method": "docx_table",
                            "table_index": table_idx
                        }
                    })
                    
            return result if result else [self._create_error_chunk("Kein Inhalt aus DOCX extrahiert")]
            
        except Exception as e:
            logger.error(f"❌ DOCX-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"DOCX-Extraktion fehlgeschlagen: {e}")]

    def extract_text(self, text: str, file_type: str = "txt") -> List[Dict]:
        """Extrahiere Text mit adaptiver Chunk-Größe - GARANTIERT List[Dict]"""
        if not text.strip():
            return [{"content": "[Leerer Text]", "type": "text", "metadata": {"extraction_method": "text_in_memory"}}]
            
        try:
            # Adaptive Chunk-Größen basierend auf Dateityp
            if file_type in ['csv', 'tsv']:
                chunk_size = self.chunk_size // 2
            elif file_type in ['md', 'txt']:
                chunk_size = int(self.chunk_size * 1.2)
            else:
                chunk_size = self.chunk_size
                
            chunks = self.split_text(text, chunk_size)
            return [{
                "content": chunk_text,
                "type": "text",
                "metadata": {
                    "extraction_method": "text_in_memory",
                    "chunk_size_used": chunk_size,
                    "file_type": file_type
                }
            } for chunk_text in chunks]
            
        except Exception as e:
            logger.error(f"❌ Text-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Text-Extraktion fehlgeschlagen: {e}")]

    async def extract_spreadsheet(self, file_bytes: bytes, file_type: str) -> List[Dict]:
        """Extrahiere Tabellen aus Excel/CSV - GARANTIERT List[Dict]"""
        if not HAS_PANDAS:
            return [self._create_error_chunk("Tabellen-Support nicht verfügbar (pandas nicht installiert)")]
            
        try:
            if file_type == "csv":
                df = pd.read_csv(io.BytesIO(file_bytes))
            else:
                df = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
                
            chunks = []
            if isinstance(df, dict):
                for sheet_name, sheet_df in df.items():
                    chunks.extend(self._process_dataframe(sheet_df, {"sheet": sheet_name}))
            else:
                chunks.extend(self._process_dataframe(df))
                
            return chunks if chunks else [self._create_error_chunk("Keine Tabellendaten extrahiert")]
            
        except Exception as e:
            logger.error(f"❌ Tabelle fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Tabellen-Extraktion fehlgeschlagen: {e}")]

    def _process_dataframe(self, df: "pd.DataFrame", metadata: dict = None) -> List[Dict]:
        """Verarbeite DataFrame in Chunks - GARANTIERT List[Dict]"""
        metadata = metadata or {}
        chunks = []
        
        try:
            for i in range(0, len(df), 50):
                chunk_df = df.iloc[i:i + 50]
                chunks.append({
                    "content": chunk_df.to_json(orient="records"),
                    "type": "table",
                    "metadata": {
                        **metadata, 
                        "row_start": i,
                        "extraction_method": "pandas"
                    }
                })
        except Exception as e:
            logger.error(f"❌ DataFrame-Verarbeitung fehlgeschlagen: {e}")
            chunks.append(self._create_error_chunk(f"DataFrame-Fehler: {e}"))
            
        return chunks

    async def extract_email(self, eml_bytes: bytes) -> List[Dict]:
        """Extrahiere E-Mail-Metadaten und Inhalt - GARANTIERT List[Dict]"""
        try:
            msg = BytesParser(policy=policy.default).parsebytes(eml_bytes)
            chunks = [{
                "content": json.dumps({
                    "from": msg["From"],
                    "to": msg["To"],
                    "subject": msg["Subject"],
                    "date": msg["Date"],
                }),
                "type": "metadata",
                "metadata": {"extraction_method": "email_header"}
            }]
            
            body = msg.get_body(preferencelist=("plain", "html"))
            if body:
                content = body.get_content()
                if content and content.strip():
                    # FIX: Korrekte Verarbeitung der split_text Rückgabe
                    text_chunks = self.split_text(content.strip())
                    chunks.extend([{
                        "content": chunk_text, 
                        "type": "text",
                        "metadata": {"extraction_method": "email_body"}
                    } for chunk_text in text_chunks])
                    
            return chunks
            
        except Exception as e:
            logger.error(f"❌ E-Mail-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"E-Mail-Extraktion fehlgeschlagen: {e}")]

    def extract_html(self, html: str) -> List[Dict]:
        """Extrahiere Text aus HTML - GARANTIERT List[Dict]"""
        try:
            if HAS_BS4:
                soup = BeautifulSoup(html, "html.parser")
                text = soup.get_text(separator="\n", strip=True)
            else:
                import re
                text = re.sub("<[^<]+?>", "", html)
                
            return self.extract_text(text, "html")
            
        except Exception as e:
            logger.error(f"❌ HTML-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"HTML-Extraktion fehlgeschlagen: {e}")]

    async def extract_archive(self, zip_bytes: bytes) -> List[Dict]:
        """Extrahiere Dateiliste aus ZIP - GARANTIERT List[Dict]"""
        try:
            with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
                files = zf.namelist()
                return [{
                    "content": f"Archive enthält: {name}",
                    "type": "metadata",
                    "metadata": {
                        "archived_file": name,
                        "extraction_method": "zip_listing"
                    }
                } for name in files if not name.endswith("/")]
                
        except Exception as e:
            logger.error(f"❌ Archiv-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Archiv-Extraktion fehlgeschlagen: {e}")]

    # === UTILITY METHODEN ===
    
    def split_text(self, text: str, chunk_size: int = None) -> List[str]:
        """
        Teile Text in Chunks mit dynamischer Größe
        
        Returns:
            List[str]: Liste von Text-Chunks (nicht Dict!)
        """
        chunk_size = chunk_size or self.chunk_size
        words = text.split()
        chunks = []
        current_chunk = []
        current_size = 0

        for word in words:
            current_chunk.append(word)
            current_size += len(word) + 1
            
            if current_size >= chunk_size:
                chunks.append(" ".join(current_chunk))
                current_chunk = []
                current_size = 0

        if current_chunk:
            chunks.append(" ".join(current_chunk))

        return chunks or [""]

    def _create_error_chunk(self, error_message: str) -> Dict:
        """Erstelle standardisierten Error-Chunk"""
        return {
            "content": f"[FEHLER: {error_message}]",
            "type": "error",
            "metadata": {
                "extraction_method": "error",
                "error_message": error_message,
                "timestamp": datetime.now().isoformat()
            }
        }

    def _validate_chunks(self, chunks: List[Dict]) -> List[Dict]:
        """Validiere und normalisiere Chunk-Format"""
        validated = []
        for i, chunk in enumerate(chunks):
            if not isinstance(chunk, dict):
                logger.warning(f"⚠️ Chunk {i} ist kein Dict: {type(chunk)}")
                chunk = {"content": str(chunk), "type": "text", "metadata": {}}
                
            # Stelle sicher, dass required fields vorhanden sind
            if "content" not in chunk:
                chunk["content"] = "[Leerer Chunk]"
            if "type" not in chunk:
                chunk["type"] = "text"
            if "metadata" not in chunk:
                chunk["metadata"] = {}
                
            validated.append(chunk)
            
        return validated

    def _parse_metadata(self, metadata: Union[str, dict]) -> dict:
        """Parse metadata sicher zu dict"""
        if isinstance(metadata, dict):
            return metadata
        if isinstance(metadata, str):
            try:
                return json.loads(metadata)
            except:
                return {}
        return {}

    async def _save_chunks_to_db(self, conn, document_id: str, chunks: List[Dict]):
        """Speichere Chunks in Datenbank"""
        for idx, chunk in enumerate(chunks):
            await conn.execute(
                """
                INSERT INTO chunks (
                    document_id, chunk_index, content, content_type,
                    page_number, position, metadata
                ) VALUES ($1, $2, $3, $4, $5, $6, $7)
                ON CONFLICT (document_id, chunk_index) DO UPDATE
                SET content = EXCLUDED.content,
                    content_type = EXCLUDED.content_type
                """,
                document_id,
                idx,
                chunk["content"],
                chunk.get("type", "text"),
                chunk.get("page"),
                chunk.get("position"),
                json.dumps(chunk.get("metadata", {}))
            )

    async def _mark_task_failed(self, conn, task_id: str, doc_id: str, error_message: str):
        """Markiere Task als fehlgeschlagen"""
        await conn.execute(
            """UPDATE processing_queue 
               SET status = 'failed', error_message = $2, completed_at = NOW()
               WHERE id = $1""",
            task_id, error_message
        )
        await conn.execute(
            """UPDATE documents SET status = 'failed' WHERE id = $1""",
            doc_id
        )

    async def run(self):
        """Starte den OCR-Agent"""
        await self.connect()
        logger.info("🧠 Enhanced OCR Agent V2 gestartet")
        
        try:
            await self.process_queue()
        except KeyboardInterrupt:
            logger.info("🛑 OCR Agent wird beendet (KeyboardInterrupt)")
        except Exception as e:
            logger.error(f"💥 Unerwarteter Fehler: {e}")
            raise
        finally:
            await self.close()


def main():
    """Haupt-Funktion"""
    db_url = os.getenv(
        "DATABASE_URL",
        "postgresql://semanticuser:semantic2024@postgres:5432/semantic_doc_finder"
    )
    
    logger.info(f"🚀 Starte OCR Agent mit DB: {db_url.split('@')[1] if '@' in db_url else 'localhost'}")
    
    agent = OCRAgentV2(db_url)
    asyncio.run(agent.run())


if __name__ == "__main__":
    main()
