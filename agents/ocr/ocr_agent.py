# ocr_agent_optimized_v2.py
# Erweiterte Version mit intelligentem PDF-Handling und Multi-Engine Support

import asyncio
import base64
import os
import gc
import json
import logging
import re
import hashlib
import io
import zipfile
import torch
import psutil
import fitz  # PyMuPDF
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Union, Any, Literal, Tuple
from enum import Enum

# Externe Abhängigkeiten
import asyncpg
import pytesseract
from PIL import Image
import pdf2image
from pydantic import BaseModel, Field, field_validator

# Erweiterte OCR Engines
try:
    from paddleocr import PaddleOCR
    HAS_PADDLE = True
except ImportError:
    HAS_PADDLE = False
    logging.warning("PaddleOCR not available - falling back to Tesseract")

try:
    import easyocr
    HAS_EASYOCR = True
except ImportError:
    HAS_EASYOCR = False
    logging.warning("EasyOCR not available - falling back to Tesseract")

# Bedingte Imports (bleibt wie es war)
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logging.warning("python-docx not available - DOCX support disabled")

try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False
    logging.warning("pandas not available - Excel/CSV support limited")

try:
    from bs4 import BeautifulSoup
    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False
    logging.warning("beautifulsoup4 not available - HTML parsing limited")

# Email parsing
import email
from email.parser import BytesParser
from email import policy

# Logging setup
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger("OCRAgentV2")


# === ENUMS & CONSTANTS ===

class OCREngine(str, Enum):
    """Verfügbare OCR Engines"""
    TESSERACT = "tesseract"
    PADDLEOCR = "paddleocr"
    EASYOCR = "easyocr"
    AUTO = "auto"


class ExtractionMode(str, Enum):
    """Extraktionsmodi"""
    NATIVE = "native"           # Nur native Textextraktion
    OCR = "ocr"                # Nur OCR
    HYBRID = "hybrid"          # Native + OCR für beste Ergebnisse
    AUTO = "auto"              # Automatische Auswahl


# === PYDANTIC MODELS ===

class OCRConfig(BaseModel):
    """OCR Konfiguration"""
    engine: OCREngine = OCREngine.AUTO
    languages: List[str] = ["deu", "eng"]
    confidence_threshold: float = 0.6
    enable_coordinates: bool = False
    enable_confidence: bool = True
    extraction_mode: ExtractionMode = ExtractionMode.AUTO
    dpi: int = 300
    
    class Config:
        use_enum_values = True


class DocumentMetadata(BaseModel):
    """Strukturierte Metadaten mit N8N-Kompatibilität"""
    
    # Content sources
    file_content: Optional[str] = None
    file_path: Optional[str] = None
    
    # Processing configuration
    processing_mode: Literal["auto", "streaming", "in_memory"] = "auto"
    source: str = "unknown"
    
    # OCR Configuration
    ocr_config: Optional[OCRConfig] = None
    
    # File metadata
    original_filename: Optional[str] = None
    file_size_mb: Optional[Union[float, str]] = None
    
    # N8N compatibility fields (camelCase)
    fileContent: Optional[str] = None
    filePath: Optional[str] = None
    
    # Additional N8N fields
    fileName: Optional[str] = None
    fileType: Optional[str] = None
    
    # Additional metadata
    uploadedAt: Optional[str] = None
    processed: bool = False
    
    class Config:
        extra = "allow"
    
    from pydantic import field_validator  
    @field_validator('file_size_mb', mode='before')  
    @classmethod  
    def convert_file_size(cls, v):
        if v is None:
            return None
        if isinstance(v, (int, float)):
            return float(v)
        if isinstance(v, str):
            try:
                return float(v)
            except ValueError:
                return None
        return None
    
    def get_file_content(self) -> Optional[str]:
        """Get file content with fallback"""
        return self.file_content or self.fileContent
    
    def get_file_path(self) -> Optional[str]:
        """Get file path with validation"""
        path = self.file_path or self.filePath
        if path and os.path.exists(path) and os.path.isfile(path):
            return path
        return None
    
    def get_original_filename(self) -> Optional[str]:
        """Get filename with fallback"""
        return self.original_filename or self.fileName
    
    def get_ocr_config(self) -> OCRConfig:
        """Get OCR config with defaults from environment"""
        if self.ocr_config:
            return self.ocr_config
            
        # Build from environment variables
        return OCRConfig(
            engine=os.getenv("OCR_ENGINE", "auto").lower(),
            languages=os.getenv("OCR_LANGUAGES", "deu,eng").split(","),
            confidence_threshold=float(os.getenv("OCR_CONFIDENCE_THRESHOLD", "0.6")),
            enable_coordinates=os.getenv("ENABLE_COORDINATES", "false").lower() == "true",
            enable_confidence=os.getenv("ENABLE_CONFIDENCE", "true").lower() == "true",
            extraction_mode=os.getenv("EXTRACTION_MODE", "auto").lower(),
            dpi=int(os.getenv("OCR_DPI", "300"))
        )
    
    def determine_processing_mode(self) -> str:
        """Auto-determine processing mode based on file size"""
        if self.processing_mode != "auto":
            return self.processing_mode
        
        file_path = self.get_file_path()
        file_content = self.get_file_content()
    
        size_mb = 0
        if file_path and os.path.exists(file_path):
            size_mb = os.path.getsize(file_path) / (1024 * 1024)
            logger.info(f"📏 File size from path: {size_mb:.1f}MB")
        elif file_content:
            size_mb = (len(file_content) * 0.75) / (1024 * 1024)
            logger.info(f"📏 File size from base64: {size_mb:.1f}MB")
    
        mode = "streaming" if size_mb > 10 else "in_memory" 
        logger.info(f"🔧 Processing mode: {mode} (size: {size_mb:.1f}MB)")
        return mode


class OCRResult(BaseModel):
    """OCR Ergebnis mit optionalen Koordinaten"""
    text: str
    confidence: Optional[float] = None
    bbox: Optional[List[float]] = None  # [x1, y1, x2, y2]
    language: Optional[str] = None


# === OCR ENGINE IMPLEMENTATIONS ===

class BaseOCREngine:
    """Basis-Klasse für OCR Engines"""
    
    def __init__(self, config: OCRConfig):
        self.config = config
        
    async def process(self, image_bytes: bytes) -> List[OCRResult]:
        """Process image and return OCR results"""
        raise NotImplementedError


class TesseractEngine(BaseOCREngine):
    """Tesseract OCR Engine"""
    
    def __init__(self, config: OCRConfig):
        super().__init__(config)
        self.tesseract_config = "--oem 3 --psm 6"
        self.tesseract_config_fast = "--oem 1 --psm 6"
        
    async def process(self, image_bytes: bytes) -> List[OCRResult]:
        """Process with Tesseract"""
        try:
            image = Image.open(io.BytesIO(image_bytes))
            
            if self.config.enable_coordinates or self.config.enable_confidence:
                # Use image_to_data for detailed output
                data = pytesseract.image_to_data(
                    image,
                    lang="+".join(self.config.languages),
                    config=self.tesseract_config,
                    output_type=pytesseract.Output.DICT
                )
                
                results = []
                n_boxes = len(data['text'])
                
                for i in range(n_boxes):
                    if data['text'][i].strip():
                        result = OCRResult(
                            text=data['text'][i],
                            confidence=data['conf'][i] / 100.0 if self.config.enable_confidence else None,
                            bbox=[
                                data['left'][i], 
                                data['top'][i], 
                                data['left'][i] + data['width'][i], 
                                data['top'][i] + data['height'][i]
                            ] if self.config.enable_coordinates else None
                        )
                        
                        if result.confidence is None or result.confidence >= self.config.confidence_threshold:
                            results.append(result)
                
                return results
            else:
                # Simple text extraction
                text = pytesseract.image_to_string(
                    image,
                    lang="+".join(self.config.languages),
                    config=self.tesseract_config
                )
                return [OCRResult(text=text)]
                
        except Exception as e:
            logger.error(f"Tesseract error: {e}")
            return []


class PaddleOCREngine(BaseOCREngine):
    def __init__(self, config: OCRConfig):
        super().__init__(config)
        if not HAS_PADDLE:
            raise ImportError("PaddleOCR not available")
            
        # Map language codes
        lang_map = {
            "deu": "german", 
            "eng": "en",
            "fra": "french",
            "spa": "spanish",
            "ita": "italian"
        }
        
        paddle_lang = lang_map.get(self.config.languages[0], "en")
        
        # KORRIGIERT: Moderne PaddleOCR Parameter
        self.ocr = PaddleOCR(
            lang=paddle_lang,
            use_textline_orientation=True,
            text_recognition_model_dir=None,
            text_detection_model_dir=None,
            device="gpu" if torch.cuda.is_available() else "cpu"
        )
        
    async def process(self, image_bytes: bytes) -> List[OCRResult]:
        """Process with PaddleOCR - ULTRA-ROBUSTE VERSION"""
        try:
            # Convert to numpy array
            image = Image.open(io.BytesIO(image_bytes))
            img_array = np.array(image)
            
            result = self.ocr.predict(img_array)
            
            # ✅ ULTRA-ROBUSTE PRÜFUNG - verhindert ALLE "string index out of range" Fehler
            if result is None:
                logger.warning("PaddleOCR returned None")
                return []
                
            if not isinstance(result, (list, tuple)):
                logger.warning(f"PaddleOCR returned unexpected type: {type(result)}")
                return []
                
            if len(result) == 0:
                logger.warning("PaddleOCR returned empty result")
                return []
                
            # Prüfe erste Seite
            first_page = result[0]
            if first_page is None:
                logger.warning("PaddleOCR first page is None")
                return []
                
            if not isinstance(first_page, (list, tuple)):
                logger.warning(f"PaddleOCR first page unexpected type: {type(first_page)}")
                return []
                
            if len(first_page) == 0:
                logger.info("PaddleOCR found no text on page")
                return []
            
            ocr_results = []
            
            # Iteriere über alle erkannten Textzeilen mit extra Sicherheit
            for line_idx, line in enumerate(first_page):
                try:
                    # Umfassende Line-Validierung
                    if line is None:
                        continue
                        
                    if not isinstance(line, (list, tuple)):
                        logger.debug(f"Line {line_idx} wrong type: {type(line)}")
                        continue
                        
                    if len(line) < 2:
                        logger.debug(f"Line {line_idx} too short: {len(line)}")
                        continue
                        
                    bbox, text_info = line[0], line[1]
                    
                    # Text-Info Validierung
                    if text_info is None:
                        continue
                        
                    if not isinstance(text_info, (list, tuple)):
                        logger.debug(f"Line {line_idx} text_info wrong type: {type(text_info)}")
                        continue
                        
                    if len(text_info) < 2:
                        logger.debug(f"Line {line_idx} text_info too short: {len(text_info)}")
                        continue
                        
                    text, confidence = text_info[0], text_info[1]
                    
                    # Text-Validierung
                    if text is None or text == "":
                        continue
                        
                    if not isinstance(text, str):
                        text = str(text)
                        
                    text = text.strip()
                    if not text:
                        continue
                    
                    # Confidence-Validierung
                    if not isinstance(confidence, (int, float)):
                        try:
                            confidence = float(confidence)
                        except (ValueError, TypeError):
                            confidence = 0.0
                    
                    # Confidence-Schwelle prüfen
                    if confidence < self.config.confidence_threshold:
                        continue
                    
                    # BBox-Validierung (optional)
                    processed_bbox = None
                    if self.config.enable_coordinates and bbox is not None:
                        try:
                            if isinstance(bbox, (list, tuple)) and len(bbox) >= 4:
                                processed_bbox = [
                                    float(bbox[0][0]), float(bbox[0][1]),  # top-left
                                    float(bbox[2][0]), float(bbox[2][1])   # bottom-right
                                ]
                        except (IndexError, TypeError, ValueError):
                            logger.debug(f"Invalid bbox for line {line_idx}")
                    
                    # OCR-Result erstellen
                    ocr_result = OCRResult(
                        text=text,
                        confidence=confidence if self.config.enable_confidence else None,
                        bbox=processed_bbox
                    )
                    
                    ocr_results.append(ocr_result)
                        
                except Exception as line_error:
                    logger.debug(f"Error processing line {line_idx}: {line_error}")
                    continue
                    
            if not ocr_results:
                logger.info("No valid OCR results found after processing")
                    
            return ocr_results
            
        except Exception as e:
            logger.error(f"PaddleOCR error: {e}")
            logger.debug(f"Image shape: {img_array.shape if 'img_array' in locals() else 'No array'}")
            return []


class EasyOCREngine(BaseOCREngine):
    """EasyOCR Engine - Easy to use, good accuracy"""
    
    def __init__(self, config: OCRConfig):
        super().__init__(config)
        if not HAS_EASYOCR:
            raise ImportError("EasyOCR not available")
            
        # Map language codes
        lang_map = {
            "deu": "de",
            "eng": "en",
            "fra": "fr",
            "spa": "es",
            "ita": "it"
        }
        
        easy_langs = [lang_map.get(lang, lang) for lang in self.config.languages]
        
        self.reader = easyocr.Reader(
            easy_langs,
            gpu=torch.cuda.is_available()
        )
        
    async def process(self, image_bytes: bytes) -> List[OCRResult]:
        """Process with EasyOCR"""
        try:
            # EasyOCR can work directly with bytes
            result = self.reader.readtext(image_bytes)
            
            ocr_results = []
            for (bbox, text, confidence) in result:
                ocr_result = OCRResult(
                    text=text,
                    confidence=confidence if self.config.enable_confidence else None,
                    bbox=[
                        bbox[0][0], bbox[0][1],  # top-left
                        bbox[2][0], bbox[2][1]   # bottom-right
                    ] if self.config.enable_coordinates else None
                )
                
                if confidence >= self.config.confidence_threshold:
                    ocr_results.append(ocr_result)
                    
            return ocr_results
            
        except Exception as e:
            logger.error(f"EasyOCR error: {e}")
            return []


class OCREngineManager:
    """Manager für OCR Engines"""
    
    def __init__(self):
        self.engines: Dict[OCREngine, BaseOCREngine] = {}
        
    def get_engine(self, engine_type: OCREngine, config: OCRConfig) -> BaseOCREngine:
        """Get or create OCR engine"""
        if engine_type == OCREngine.AUTO:
            engine_type = self._select_best_engine()
            
        if engine_type not in self.engines:
            if engine_type == OCREngine.TESSERACT:
                self.engines[engine_type] = TesseractEngine(config)
            elif engine_type == OCREngine.PADDLEOCR and HAS_PADDLE:
                self.engines[engine_type] = PaddleOCREngine(config)
            elif engine_type == OCREngine.EASYOCR and HAS_EASYOCR:
                self.engines[engine_type] = EasyOCREngine(config)
            else:
                # Fallback to Tesseract
                logger.warning(f"Engine {engine_type} not available, falling back to Tesseract")
                self.engines[engine_type] = TesseractEngine(config)
                
        return self.engines[engine_type]
    
    def _select_best_engine(self) -> OCREngine:
        """Select best available engine"""
        try:
            if HAS_PADDLE:
                # Test mit einfachem Bild ob PaddleOCR funktioniert
                test_ocr = PaddleOCR(
                    lang='en', 
                    use_textline_orientation=False,
                    device="gpu" if torch.cuda.is_available() else "cpu"
                )
            return OCREngine.PADDLEOCR
        except Exception as e:
            logger.warning(f"PaddleOCR test failed: {e}")
    
    # Fallback zu anderen Engines...
        return OCREngine.TESSERACT


# === MAIN OCR AGENT CLASS ===

class OCRAgentOptimizedV2:
    """
    Erweiterte OCR Agent Version mit intelligentem PDF-Handling und Multi-Engine Support
    """
    
    def __init__(self, db_url: str):
        self.db_url = db_url
        self.db_pool = None
        self.chunk_size = int(os.getenv("CHUNK_SIZE", "1000"))
        self.running = True
        
        # OCR Engine Manager
        self.ocr_manager = OCREngineManager()
        
        # Size limits
        self.MAX_FILE_SIZE = 100 * 1024 * 1024  # 100 MB
        self.MAX_BASE64_LENGTH = 700 * 1024 * 1024  # 700MB string
        
        logger.info(f"📊 OCR Agent V2 initialisiert:")
        logger.info(f"   • Chunk-Größe: {self.chunk_size}")
        logger.info(f"   • Max File Size: {self.MAX_FILE_SIZE / (1024*1024):.0f}MB")
        logger.info(f"   • Verfügbare Engines: {self._get_available_engines()}")
        
        try:
            if torch.cuda.is_available():
                gpu_count = torch.cuda.device_count()
                gpu_name = torch.cuda.get_device_name(0) if gpu_count > 0 else "Unknown"
                logger.info(f"🚀 GPU available: {gpu_count}x {gpu_name}")
            else:
                logger.warning("❌ No GPU available - running CPU only")
        except Exception as e:
            logger.warning(f"❌ GPU check failed: {e}")
            
    def _get_available_engines(self) -> List[str]:
        """Get list of available OCR engines"""
        engines = ["tesseract"]
        if HAS_PADDLE:
            engines.append("paddleocr")
        if HAS_EASYOCR:
            engines.append("easyocr")
        return engines
        
    def check_memory_before_processing(self, file_size_mb: float) -> str:
        """Check system memory and decide processing mode"""
        try:
            memory = psutil.virtual_memory()
            available_gb = memory.available / (1024**3)
            
            logger.info(f"💾 Memory: {memory.percent:.1f}% used, {available_gb:.1f}GB available")
            
            if memory.percent > 80 or file_size_mb > 50:
                logger.warning(f"⚠️ Memory pressure detected, forcing streaming mode")
                return "streaming"
            return "auto"
        except Exception as e:
            logger.warning(f"Memory check failed: {e}")
            return "auto"

    async def connect(self):
        """Verbinde mit der PostgreSQL-Datenbank"""
        try:
            self.db_pool = await asyncpg.create_pool(self.db_url)
            logger.info("✅ Verbindung zur PostgreSQL-Datenbank hergestellt")
        except Exception as e:
            logger.error(f"❌ Datenbankverbindung fehlgeschlagen: {e}")
            raise

    async def close(self):
        """Schließe die Verbindung zur Datenbank"""
        if self.db_pool:
            await self.db_pool.close()
        logger.info("🛑 OCR Agent beendet")

    async def process_queue(self):
        """Verarbeite die Warteschlange kontinuierlich"""
        while self.running:
            try:
                async with self.db_pool.acquire() as conn:
                    task = await conn.fetchrow(
                        """
                        UPDATE processing_queue 
                        SET status = 'processing', started_at = NOW()
                        WHERE id = (
                            SELECT id FROM processing_queue 
                            WHERE status = 'pending' 
                            AND task_type IN ('ocr', 'extract')
                            AND (NOT deferred OR EXTRACT(hour FROM NOW()) BETWEEN 22 AND 6)
                            ORDER BY priority DESC, created_at ASC
                            LIMIT 1
                            FOR UPDATE SKIP LOCKED
                        )
                        RETURNING *
                        """
                    )
                    if task:
                        await self.process_document(task)
                    else:
                        await asyncio.sleep(5)
            except Exception as e:
                logger.error(f"⚠️ Fehler bei Warteschlange: {e}")
                await asyncio.sleep(10)

    async def process_document(self, task: Dict):
        """Verarbeite ein Dokument"""
        async with self.db_pool.acquire() as conn:
            try:
                doc = await conn.fetchrow(
                    "SELECT * FROM documents WHERE id = $1",
                    task["document_id"]
                )
                if not doc:
                    raise ValueError(f"🚫 Dokument {task['document_id']} nicht gefunden")

                logger.info(f"📄 Verarbeite: {doc['original_name']} (ID: {doc['id']})")

                # Extrahiere Inhalt
                chunks = await self.extract_content(doc)
                
                if not chunks:
                    logger.warning(f"⚠️ Keine Chunks erstellt für {doc['original_name']}")
                    chunks = [self._create_error_chunk("Keine Inhalte extrahiert")]

                # Validiere und speichere Chunks
                chunks = self._validate_chunks(chunks)
                await self._save_chunks_to_db(conn, doc["id"], chunks)

                # Update document status
                await conn.execute(
                    """UPDATE documents 
                       SET status = 'completed', processed_at = NOW()
                       WHERE id = $1""",
                    doc["id"]
                )

                # Complete task
                await conn.execute(
                    """UPDATE processing_queue 
                       SET status = 'completed', completed_at = NOW()
                       WHERE id = $1""",
                    task["id"]
                )

                # Create enhancement task
                await conn.execute(
                    """INSERT INTO processing_queue (document_id, task_type, priority)
                       VALUES ($1, 'enhance', $2)""",
                    doc["id"], task["priority"]
                )

                logger.info(f"✅ Dokument verarbeitet: {doc['original_name']} ({len(chunks)} Chunks)")

            except Exception as e:
                logger.error(f"❌ Fehler bei Dokumentenverarbeitung: {e}")
                await self._mark_task_failed(conn, task["id"], task["document_id"], str(e))

    async def extract_content(self, doc: Dict) -> List[Dict]:
        """Hauptmethode für Content-Extraktion"""
        try:
            file_type = doc["file_type"].lower()
            
            # Map MIME types to simple extensions
            file_type_mapping = {
                "application/pdf": "pdf",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
                "application/msword": "doc",
                "text/plain": "txt",
                "text/html": "html",
                "image/png": "png",
                "image/jpeg": "jpg",
            }
            file_type = file_type_mapping.get(file_type, file_type.split('/')[-1])
            
            metadata = self._parse_metadata_robust(doc.get("metadata", {}))
            ocr_config = metadata.get_ocr_config()
            
            # Log configuration
            logger.debug(f"📋 OCR Config: engine={ocr_config.engine}, "
                         f"languages={ocr_config.languages}, "
                         f"extraction_mode={ocr_config.extraction_mode}")
            
            # Determine processing mode
            processing_mode = metadata.determine_processing_mode()
            logger.info(f"🔍 Processing mode: {processing_mode}, file_type: {file_type}")
            
            # Get file source
            file_path = metadata.get_file_path()
            file_content = metadata.get_file_content()
            
            # Streaming mode
            if processing_mode == "streaming" and file_path:
                return await self._extract_content_streaming(file_path, file_type, doc["original_name"], ocr_config)
            
            # In-memory mode
            file_bytes = await self._get_file_bytes(file_content, file_path)
            if not file_bytes:
                raise ValueError("Keine Dateidaten gefunden")
                
            return await self._dispatch_extraction(file_type, file_bytes, doc["original_name"], ocr_config)
            
        except Exception as e:
            logger.error(f"❌ Content-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Content-Extraktion fehlgeschlagen: {e}")]

    def _parse_metadata_robust(self, raw_metadata: Union[str, dict]) -> DocumentMetadata:
        """Parse metadata to DocumentMetadata object"""
        try:
            if isinstance(raw_metadata, str):
                try:
                    metadata_dict = json.loads(raw_metadata)
                except json.JSONDecodeError:
                    logger.warning(f"⚠️ JSON-Parsing fehlgeschlagen")
                    metadata_dict = {}
            elif isinstance(raw_metadata, dict):
                metadata_dict = raw_metadata
            else:
                metadata_dict = {}
            
            logger.debug(f"📋 Metadata keys: {list(metadata_dict.keys())}")
            
            try:
                return DocumentMetadata(**metadata_dict)
            except Exception as e:
                logger.warning(f"⚠️ DocumentMetadata validation failed: {e}")
                return DocumentMetadata(
                    file_content=metadata_dict.get('file_content') or metadata_dict.get('fileContent'),
                    file_path=metadata_dict.get('file_path') or metadata_dict.get('filePath'),
                    original_filename=metadata_dict.get('original_filename') or metadata_dict.get('fileName'),
                    source=metadata_dict.get('source', 'unknown')
                )
                
        except Exception as e:
            logger.error(f"❌ Metadata-Parsing komplett fehlgeschlagen: {e}")
            return DocumentMetadata()

    async def _get_file_bytes(self, file_content: Optional[str], file_path: Optional[str]) -> Optional[bytes]:
        """Get file bytes from content or path"""
        logger.debug(f"🔍 Suche Dateidaten - file_content: {bool(file_content)}, file_path: {file_path}")
        
        if file_content:
            if self._validate_base64(file_content):
                try:
                    cleaned = self._clean_base64(file_content)
                    return base64.b64decode(cleaned)
                except Exception as e:
                    logger.warning(f"⚠️ Base64-Dekodierung fehlgeschlagen: {e}")
        
        if file_path and os.path.exists(file_path):
            try:
                with open(file_path, "rb") as f:
                    return f.read()
            except Exception as e:
                logger.warning(f"⚠️ Datei-Lesen fehlgeschlagen: {e}")
        
        logger.error(f"❌ Keine Dateidaten gefunden")
        return None

    def _validate_base64(self, content: str) -> bool:
        """Validate base64 content"""
        if not content or not isinstance(content, str):
            return False
            
        if content.startswith("data:"):
            content = content.split(",", 1)[-1] if "," in content else content
            
        cleaned = content.strip().replace('\n', '').replace(' ', '')
        
        if not re.match(r'^[A-Za-z0-9+/]*={0,2}$', cleaned):
            return False
            
        try:
            test = cleaned[:1000] if len(cleaned) > 1000 else cleaned
            base64.b64decode(test + '=' * (4 - len(test) % 4))
            return True
        except:
            return False

    def _clean_base64(self, content: str) -> str:
        """Clean base64 content"""
        if content.startswith("data:") and "," in content:
            content = content.split(",", 1)[1]
        
        return content.strip().replace('\n', '').replace(' ', '').replace('\r', '')

    async def _dispatch_extraction(self, file_type: str, file_bytes: bytes, filename: str, ocr_config: OCRConfig) -> List[Dict]:
        """Dispatch to specific extraction method"""
        try:
            if file_type == "pdf":
                return await self.extract_pdf_intelligent(file_bytes, filename, ocr_config)
            elif file_type in ["png", "jpg", "jpeg", "tiff", "tif"]:
                return await self.extract_image(file_bytes, ocr_config)
            elif file_type in ["txt", "md", "log"]:
                content = file_bytes.decode("utf-8", errors="ignore")
                return self.extract_text(content, file_type)
            elif file_type in ["docx", "doc"]:
                return await self.extract_docx(file_bytes)
            elif file_type in ["xlsx", "xls", "csv"]:
                return await self.extract_spreadsheet(file_bytes, file_type)
            elif file_type == "eml":
                return await self.extract_email(file_bytes)
            elif file_type == "html":
                content = file_bytes.decode("utf-8", errors="ignore")
                return self.extract_html(content)
            elif file_type == "zip":
                return await self.extract_archive(file_bytes)
            else:
                content = file_bytes.decode("utf-8", errors="ignore")
                return self.extract_text(content, file_type)
                
        except Exception as e:
            logger.error(f"❌ Extraktion für {file_type} fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Extraktion fehlgeschlagen: {e}")]

    # === INTELLIGENT PDF EXTRACTION ===

    async def extract_pdf_intelligent(self, pdf_bytes: bytes, filename: str, ocr_config: OCRConfig) -> List[Dict]:
        """Intelligente PDF-Extraktion mit PyMuPDF"""
        chunks = []
        extraction_mode = ocr_config.extraction_mode
        
        try:
            # Open PDF with PyMuPDF
            pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            total_pages = len(pdf_doc)
            
            logger.info(f"📄 PDF: {filename} ({total_pages} Seiten)")
            logger.info(f"🔧 Extraction mode: {extraction_mode}")
            
            for page_num, page in enumerate(pdf_doc, 1):
                if page_num % 10 == 0:
                    logger.info(f"📄 Progress: {page_num}/{total_pages}")
                
                # Determine extraction strategy
                if extraction_mode == ExtractionMode.AUTO:
                    # Check if page has text
                    text = page.get_text().strip()
                    if text:
                        extraction_mode = ExtractionMode.HYBRID
                    else:
                        extraction_mode = ExtractionMode.OCR
                
                if extraction_mode in [ExtractionMode.NATIVE, ExtractionMode.HYBRID]:
                    # Extract native text
                    text = page.get_text()
                    if text.strip():
                        # Extract with layout preservation
                        text_dict = page.get_text("dict")
                        page_chunks = self._process_pdf_text_dict(text_dict, page_num, "pymupdf_native")
                        chunks.extend(page_chunks)
                        
                        if extraction_mode == ExtractionMode.NATIVE:
                            continue
                
                if extraction_mode in [ExtractionMode.OCR, ExtractionMode.HYBRID]:
                    # OCR processing
                    # Check for images in page
                    image_list = page.get_images()
                    
                    if image_list or extraction_mode == ExtractionMode.OCR:
                        # Render page as image
                        mat = fitz.Matrix(ocr_config.dpi/72.0, ocr_config.dpi/72.0)
                        pix = page.get_pixmap(matrix=mat)
                        img_data = pix.tobytes("png")
                        
                        # Run OCR
                        ocr_results = await self._run_ocr(img_data, ocr_config)
                        
                        for ocr_result in ocr_results:
                            chunk = {
                                "content": ocr_result.text,
                                "type": "text",
                                "page": page_num,
                                "metadata": {
                                    "extraction_method": f"ocr_{ocr_config.engine}",
                                    "page_total": total_pages
                                }
                            }
                            
                            if ocr_result.confidence is not None:
                                chunk["metadata"]["confidence"] = ocr_result.confidence
                            if ocr_result.bbox is not None:
                                chunk["metadata"]["bbox"] = ocr_result.bbox
                                
                            chunks.append(chunk)
                        
                        del pix
                
                # Memory cleanup
                if page_num % 5 == 0:
                    gc.collect()
            
            pdf_doc.close()
            
        except Exception as e:
            logger.error(f"❌ PDF-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"PDF-Fehler: {e}")]
            
        return chunks or [self._create_error_chunk("Kein Text aus PDF extrahiert")]

    def _process_pdf_text_dict(self, text_dict: dict, page_num: int, method: str) -> List[Dict]:
        """Process PyMuPDF text dictionary to preserve structure"""
        chunks = []
        current_block = []
        
        for block in text_dict.get("blocks", []):
            if block["type"] == 0:  # Text block
                block_text = ""
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        block_text += span.get("text", "")
                    block_text += "\n"
                
                if block_text.strip():
                    current_block.append(block_text)
                    
                    if len("\n".join(current_block)) > self.chunk_size:
                        chunks.append({
                            "content": "\n".join(current_block),
                            "type": "text",
                            "page": page_num,
                            "metadata": {
                                "extraction_method": method,
                                "block_bbox": block["bbox"]
                            }
                        })
                        current_block = []
        
        # Last block
        if current_block:
            chunks.append({
                "content": "\n".join(current_block),
                "type": "text",
                "page": page_num,
                "metadata": {"extraction_method": method}
            })
            
        return chunks

    async def _run_ocr(self, image_bytes: bytes, ocr_config: OCRConfig) -> List[OCRResult]:
        """Run OCR with configured engine"""
        try:
            engine = self.ocr_manager.get_engine(ocr_config.engine, ocr_config)
            results = await engine.process(image_bytes)
            
            # Combine results into text chunks if coordinates not needed
            if not ocr_config.enable_coordinates and not ocr_config.enable_confidence:
                combined_text = " ".join([r.text for r in results])
                return [OCRResult(text=combined_text)]
                
            return results
            
        except Exception as e:
            logger.error(f"OCR failed: {e}")
            return []

    # === OTHER EXTRACTION METHODS (Updated) ===

    async def extract_image(self, image_bytes: bytes, ocr_config: OCRConfig) -> List[Dict]:
        """Extract text from image with configured OCR engine"""
        try:
            ocr_results = await self._run_ocr(image_bytes, ocr_config)
            
            if not ocr_results:
                return [{"content": "[Bild ohne Text]", "type": "image", "metadata": {}}]
            
            chunks = []
            for ocr_result in ocr_results:
                text_chunks = self.split_text(ocr_result.text)
                for chunk in text_chunks:
                    chunk_data = {
                        "content": chunk,
                        "type": "text",
                        "page": 1,
                        "metadata": {
                            "extraction_method": f"image_ocr_{ocr_config.engine}"
                        }
                    }
                    
                    if ocr_result.confidence is not None:
                        chunk_data["metadata"]["confidence"] = ocr_result.confidence
                    if ocr_result.bbox is not None:
                        chunk_data["metadata"]["bbox"] = ocr_result.bbox
                        
                    chunks.append(chunk_data)
                    
            return chunks
            
        except Exception as e:
            logger.error(f"❌ Bild-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Bild-Fehler: {e}")]

    def extract_text(self, text: str, file_type: str = "txt") -> List[Dict]:
        """Extract text with chunking"""
        if not text.strip():
            return [{"content": "[Leerer Text]", "type": "text", "metadata": {}}]
            
        chunks = self.split_text(text)
        return [{
            "content": chunk,
            "type": "text",
            "metadata": {
                "extraction_method": "text_direct",
                "file_type": file_type
            }
        } for chunk in chunks]

    async def extract_docx(self, docx_bytes: bytes) -> List[Dict]:
        """Extract text from DOCX"""
        if not HAS_DOCX:
            return [self._create_error_chunk("DOCX-Support nicht verfügbar")]
            
        try:
            doc = DocxDocument(io.BytesIO(docx_bytes))
            text = "\n".join([para.text for para in doc.paragraphs])
            
            result = self.extract_text(text, "docx")
            
            # Extract tables
            for idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                    
                if table_data:
                    result.append({
                        "content": json.dumps(table_data),
                        "type": "table",
                        "metadata": {
                            "extraction_method": "docx_table",
                            "table_index": idx
                        }
                    })
                    
            return result
            
        except Exception as e:
            logger.error(f"❌ DOCX-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"DOCX-Fehler: {e}")]

    async def extract_spreadsheet(self, file_bytes: bytes, file_type: str) -> List[Dict]:
        """Extract data from spreadsheet"""
        if not HAS_PANDAS:
            return [self._create_error_chunk("Spreadsheet-Support nicht verfügbar")]
            
        try:
            if file_type == "csv":
                df = pd.read_csv(io.BytesIO(file_bytes))
            else:
                df = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
                
            chunks = []
            
            if isinstance(df, dict):  # Multiple sheets
                for sheet_name, sheet_df in df.items():
                    chunks.extend(self._process_dataframe(sheet_df, {"sheet": sheet_name}))
            else:
                chunks.extend(self._process_dataframe(df))
                
            return chunks or [self._create_error_chunk("Keine Daten extrahiert")]
            
        except Exception as e:
            logger.error(f"❌ Spreadsheet-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Spreadsheet-Fehler: {e}")]

    def _process_dataframe(self, df: "pd.DataFrame", metadata: dict = None) -> List[Dict]:
        """Process DataFrame into chunks"""
        chunks = []
        metadata = metadata or {}
        
        try:
            # Split into chunks of 50 rows
            for i in range(0, len(df), 50):
                chunk_df = df.iloc[i:i + 50]
                chunks.append({
                    "content": chunk_df.to_json(orient="records"),
                    "type": "table",
                    "metadata": {
                        **metadata,
                        "row_start": i,
                        "row_count": len(chunk_df),
                        "extraction_method": "pandas"
                    }
                })
        except Exception as e:
            logger.error(f"DataFrame-Verarbeitung fehlgeschlagen: {e}")
            
        return chunks

    async def extract_email(self, eml_bytes: bytes) -> List[Dict]:
        """Extract email content"""
        try:
            msg = BytesParser(policy=policy.default).parsebytes(eml_bytes)
            
            # Email metadata
            chunks = [{
                "content": json.dumps({
                    "from": msg["From"],
                    "to": msg["To"],
                    "subject": msg["Subject"],
                    "date": msg["Date"],
                }),
                "type": "metadata",
                "metadata": {"extraction_method": "email_header"}
            }]
            
            # Email body
            body = msg.get_body(preferencelist=("plain", "html"))
            if body:
                content = body.get_content()
                if content and content.strip():
                    text_chunks = self.split_text(content)
                    chunks.extend([{
                        "content": chunk,
                        "type": "text",
                        "metadata": {"extraction_method": "email_body"}
                    } for chunk in text_chunks])
                    
            return chunks
            
        except Exception as e:
            logger.error(f"❌ Email-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Email-Fehler: {e}")]

    def extract_html(self, html: str) -> List[Dict]:
        """Extract text from HTML"""
        try:
            if HAS_BS4:
                soup = BeautifulSoup(html, "html.parser")
                text = soup.get_text(separator="\n", strip=True)
            else:
                # Simple regex fallback
                text = re.sub('<[^<]+?>', '', html)
                
            return self.extract_text(text, "html")
            
        except Exception as e:
            logger.error(f"❌ HTML-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"HTML-Fehler: {e}")]

    async def extract_archive(self, zip_bytes: bytes) -> List[Dict]:
        """Extract file list from archive"""
        try:
            with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
                files = zf.namelist()
                return [{
                    "content": f"Archive enthält: {name}",
                    "type": "metadata",
                    "metadata": {
                        "archived_file": name,
                        "extraction_method": "zip_listing"
                    }
                } for name in files if not name.endswith("/")]
                
        except Exception as e:
            logger.error(f"❌ Archiv-Extraktion fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Archiv-Fehler: {e}")]

    # === STREAMING METHODS ===

    async def _extract_content_streaming(self, file_path: str, file_type: str, filename: str, ocr_config: OCRConfig) -> List[Dict]:
        """Streaming extraction for large files"""
        try:
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            logger.info(f"🚀 Streaming für {filename}: {file_size_mb:.1f}MB")
            
            if file_type == "pdf":
                return await self._stream_large_pdf(file_path, ocr_config)
            elif file_type in ["txt", "md", "log"]:
                return await self._stream_large_text(file_path)
            elif file_type in ["docx", "doc"] and HAS_DOCX:
                return await self._stream_large_docx(file_path)
            else:
                # Fallback to normal loading
                with open(file_path, "rb") as f:
                    file_bytes = f.read()
                return await self._dispatch_extraction(file_type, file_bytes, filename, ocr_config)
                
        except Exception as e:
            logger.error(f"❌ Streaming fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Streaming-Fehler: {e}")]

    async def _stream_large_pdf(self, file_path: str, ocr_config: OCRConfig) -> List[Dict]:
        """Stream large PDF with intelligent processing"""
        chunks = []
        
        try:
            pdf_doc = fitz.open(file_path)
            total_pages = len(pdf_doc)
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            
            # Adjust quality for large files
            if file_size_mb > 50:
                ocr_config.dpi = 150
            
            for page_num, page in enumerate(pdf_doc, 1):
                if page_num % 10 == 0:
                    logger.info(f"📄 Progress: {page_num}/{total_pages}")
                
                # Try native extraction first
                text = page.get_text()
                
                if text.strip():
                    # Native text available
                    page_chunks = self.split_text(text)
                    for chunk in page_chunks:
                        chunks.append({
                            "content": chunk,
                            "type": "text",
                            "page": page_num,
                            "metadata": {
                                "extraction_method": "pdf_streaming_native",
                                "page_total": total_pages
                            }
                        })
                else:
                    # Need OCR
                    mat = fitz.Matrix(ocr_config.dpi/72.0, ocr_config.dpi/72.0)
                    pix = page.get_pixmap(matrix=mat)
                    img_data = pix.tobytes("png")
                    
                    ocr_results = await self._run_ocr(img_data, ocr_config)
                    
                    for result in ocr_results:
                        if result.text.strip():
                            chunks.append({
                                "content": result.text,
                                "type": "text",
                                "page": page_num,
                                "metadata": {
                                    "extraction_method": f"pdf_streaming_ocr_{ocr_config.engine}",
                                    "dpi": ocr_config.dpi
                                }
                            })
                    
                    del pix
                
                # Memory cleanup
                if page_num % 5 == 0:
                    gc.collect()
                    
            pdf_doc.close()
                    
        except Exception as e:
            logger.error(f"❌ PDF-Streaming fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"PDF-Streaming-Fehler: {e}")]
            
        return chunks

    async def _stream_large_text(self, file_path: str) -> List[Dict]:
        """Stream large text file"""
        chunks = []
        current_chunk = []
        current_size = 0
        
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                for line in f:
                    current_chunk.append(line)
                    current_size += len(line)
                    
                    if current_size >= self.chunk_size:
                        chunks.append({
                            "content": "".join(current_chunk),
                            "type": "text",
                            "metadata": {"extraction_method": "text_streaming"}
                        })
                        current_chunk = []
                        current_size = 0
                        
                # Last chunk
                if current_chunk:
                    chunks.append({
                        "content": "".join(current_chunk),
                        "type": "text",
                        "metadata": {"extraction_method": "text_streaming"}
                    })
                    
        except Exception as e:
            logger.error(f"❌ Text-Streaming fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"Text-Streaming-Fehler: {e}")]
            
        return chunks

    async def _stream_large_docx(self, file_path: str) -> List[Dict]:
        """Stream large DOCX file"""
        if not HAS_DOCX:
            return [self._create_error_chunk("DOCX-Support nicht verfügbar")]
            
        try:
            doc = DocxDocument(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            chunks = self.split_text(text)
            
            result = [{
                "content": chunk,
                "type": "text",
                "metadata": {"extraction_method": "docx_streaming"}
            } for chunk in chunks]
            
            # Extract tables
            for idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                if table_data:
                    result.append({
                        "content": json.dumps(table_data),
                        "type": "table",
                        "metadata": {"table_index": idx}
                    })
                    
            return result
            
        except Exception as e:
            logger.error(f"❌ DOCX-Streaming fehlgeschlagen: {e}")
            return [self._create_error_chunk(f"DOCX-Streaming-Fehler: {e}")]

    # === UTILITY METHODS ===

    def split_text(self, text: str, chunk_size: int = None) -> List[str]:
        """Split text into chunks"""
        chunk_size = chunk_size or self.chunk_size
        words = text.split()
        chunks = []
        current_chunk = []
        current_size = 0

        for word in words:
            current_chunk.append(word)
            current_size += len(word) + 1
            
            if current_size >= chunk_size:
                chunks.append(" ".join(current_chunk))
                current_chunk = []
                current_size = 0

        if current_chunk:
            chunks.append(" ".join(current_chunk))

        return chunks or [""]

    def _create_error_chunk(self, error_message: str) -> Dict:
        """Create error chunk"""
        return {
            "content": f"[FEHLER: {error_message}]",
            "type": "error",
            "metadata": {
                "extraction_method": "error",
                "error_message": error_message,
                "timestamp": datetime.now().isoformat()
            }
        }

    def _validate_chunks(self, chunks: List[Dict]) -> List[Dict]:
        """Validate chunk format"""
        validated = []
        for chunk in chunks:
            if not isinstance(chunk, dict):
                chunk = {"content": str(chunk), "type": "text", "metadata": {}}
                
            # Ensure required fields
            chunk.setdefault("content", "[Leer]")
            chunk.setdefault("type", "text")
            chunk.setdefault("metadata", {})
            
            validated.append(chunk)
            
        return validated

    async def _save_chunks_to_db(self, conn, document_id: str, chunks: List[Dict]):
        """Save chunks to database"""
        for idx, chunk in enumerate(chunks):
            await conn.execute(
                """
                INSERT INTO chunks (
                    document_id, chunk_index, content, content_type,
                    page_number, position, metadata
                ) VALUES ($1, $2, $3, $4, $5, $6, $7)
                ON CONFLICT (document_id, chunk_index) DO UPDATE
                SET content = EXCLUDED.content,
                    content_type = EXCLUDED.content_type,
                    metadata = EXCLUDED.metadata
                """,
                document_id,
                idx,
                chunk["content"],
                chunk.get("type", "text"),
                chunk.get("page"),
                chunk.get("position"),
                json.dumps(chunk.get("metadata", {}))
            )

    async def _mark_task_failed(self, conn, task_id: str, doc_id: str, error_message: str):
        """Mark task as failed"""
        await conn.execute(
            """UPDATE processing_queue 
               SET status = 'failed', error_message = $2, completed_at = NOW()
               WHERE id = $1""",
            task_id, error_message
        )
        await conn.execute(
            """UPDATE documents SET status = 'failed' WHERE id = $1""",
            doc_id
        )

    async def run(self):
        """Main run loop"""
        await self.connect()
        logger.info("🧠 OCR Agent V2 gestartet")
        
        try:
            await self.process_queue()
        except KeyboardInterrupt:
            logger.info("🛑 OCR Agent wird beendet")
        except Exception as e:
            logger.error(f"💥 Unerwarteter Fehler: {e}")
            raise
        finally:
            await self.close()


# Import numpy for PaddleOCR
try:
    import numpy as np
except ImportError:
    logger.warning("numpy not available - required for PaddleOCR")


def main():
    """Main entry point"""
    db_url = os.getenv(
        "DATABASE_URL",
        "postgresql://semanticuser:semantic2024@postgres:5432/semantic_doc_finder"
    )
    
    agent = OCRAgentOptimizedV2(db_url)
    asyncio.run(agent.run())


if __name__ == "__main__":
    main()
