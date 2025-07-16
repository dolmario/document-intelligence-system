import asyncio
import hashlib
import io
import json
import logging
import os
import sys
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional
from uuid import uuid4

import asyncpg
import pytesseract
from PIL import Image
import pdf2image

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger('ocr_agent')

# Optional imports with fallbacks
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logging.warning("python-docx not available - DOCX support disabled")

try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False
    logging.warning("pandas not available - Excel/CSV support limited")

import email
from email import policy
from email.parser import BytesParser
import zipfile

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger('ocr_agent')


class OCRAgentV2:
    def __init__(self, db_url: str):
        self.db_url = db_url
        self.db_pool = None
        self.chunk_size = int(os.getenv('CHUNK_SIZE', '1000'))
        self.running = True
        
    async def connect(self):
        self.db_pool = await asyncpg.create_pool(self.db_url)
        logger.info("Connected to PostgreSQL")
    
    async def close(self):
        if self.db_pool:
            await self.db_pool.close()
    
    async def process_queue(self):
        while self.running:
            try:
                async with self.db_pool.acquire() as conn:
                    # Get next task
                    task = await conn.fetchrow("""
                        UPDATE processing_queue 
                        SET status = 'processing', started_at = NOW()
                        WHERE id = (
                            SELECT id FROM processing_queue 
                            WHERE status = 'pending' 
                            AND task_type IN ('ocr', 'extract')
                            AND (NOT deferred OR EXTRACT(hour FROM NOW()) BETWEEN 22 AND 6)
                            ORDER BY priority DESC, created_at ASC
                            LIMIT 1
                            FOR UPDATE SKIP LOCKED
                        )
                        RETURNING *
                    """)
                    
                    if task:
                        await self.process_document(task)
                    else:
                        await asyncio.sleep(5)
                        
            except Exception as e:
                logger.error(f"Queue processing error: {e}")
                await asyncio.sleep(10)
    
    async def process_document(self, task: Dict):
        async with self.db_pool.acquire() as conn:
            try:
                # Get document info
                doc = await conn.fetchrow(
                    "SELECT * FROM documents WHERE id = $1",
                    task['document_id']
                )
                
                if not doc:
                    raise ValueError(f"Document {task['document_id']} not found")
                
                # Extract based on file type
                chunks = await self.extract_content(doc)
                
                # Save chunks
                for idx, chunk in enumerate(chunks):
                    await conn.execute("""
                        INSERT INTO chunks (
                            document_id, chunk_index, content, content_type,
                            page_number, position, metadata
                        ) VALUES ($1, $2, $3, $4, $5, $6, $7)
                        ON CONFLICT (document_id, chunk_index) DO UPDATE
                        SET content = EXCLUDED.content,
                            content_type = EXCLUDED.content_type
                    """, doc['id'], idx, chunk['content'], chunk['type'],
                        chunk.get('page'), chunk.get('position'), 
                        json.dumps(chunk.get('metadata', {})))
                
                # Update document status
                await conn.execute("""
                    UPDATE documents 
                    SET status = 'completed', processed_at = NOW()
                    WHERE id = $1
                """, doc['id'])
                
                # Complete task
                await conn.execute("""
                    UPDATE processing_queue 
                    SET status = 'completed', completed_at = NOW()
                    WHERE id = $1
                """, task['id'])
                
                # Create enhancement tasks
                await conn.execute("""
                    INSERT INTO processing_queue (document_id, task_type, priority)
                    VALUES ($1, 'enhance', $2)
                """, doc['id'], task['priority'])
                
                logger.info(f"Processed document {doc['original_name']}: {len(chunks)} chunks")
                
            except Exception as e:
                logger.error(f"Document processing error: {e}")
                await conn.execute("""
                    UPDATE processing_queue 
                    SET status = 'failed', error_message = $2, completed_at = NOW()
                    WHERE id = $1
                """, task['id'], str(e))
                
                await conn.execute("""
                    UPDATE documents SET status = 'failed' WHERE id = $1
                """, task['document_id'])
    
    async def extract_content(self, doc: Dict) -> List[Dict]:
        file_type = doc['file_type'].lower()
        file_path = doc['metadata'].get('file_path')
        file_content = doc['metadata'].get('file_content')
        
        if file_content:
            # Base64 content from n8n
            import base64
            file_bytes = base64.b64decode(file_content)
        elif file_path and os.path.exists(file_path):
            with open(file_path, 'rb') as f:
                file_bytes = f.read()
        else:
            raise ValueError("No file content available")
        
        # Extract based on type
        if file_type == 'pdf':
            return await self.extract_pdf(file_bytes)
        elif file_type in ['png', 'jpg', 'jpeg', 'tiff', 'tif']:
            return await self.extract_image(file_bytes)
        elif file_type in ['txt', 'md', 'log']:
            return self.extract_text(file_bytes.decode('utf-8', errors='ignore'))
        elif file_type in ['docx', 'doc']:
                            if HAS_DOCX:
                    return await self.extract_docx(file_bytes)
                else:
                    return self.extract_text(file_bytes.decode('utf-8', errors='ignore'))
        elif file_type in ['xlsx', 'xls', 'csv']:
            if HAS_PANDAS:
                return await self.extract_spreadsheet(file_bytes, file_type)
            else:
                return [{'content': f'[Spreadsheet file: {doc["original_name"]}]', 'type': 'metadata'}]
        elif file_type == 'eml':
            return await self.extract_email(file_bytes)
        elif file_type == 'html':
            return self.extract_html(file_bytes.decode('utf-8', errors='ignore'))
        elif file_type == 'zip':
            return await self.extract_archive(file_bytes)
        else:
            # Fallback to text
            return self.extract_text(file_bytes.decode('utf-8', errors='ignore'))
    
    async def extract_pdf(self, pdf_bytes: bytes) -> List[Dict]:
        chunks = []
        images = pdf2image.convert_from_bytes(pdf_bytes, dpi=300)
        
        for page_num, image in enumerate(images, 1):
            text = pytesseract.image_to_string(
                image, 
                lang=os.getenv('TESSERACT_LANG', 'deu+eng')
            )
            
            if text.strip():
                page_chunks = self.split_text(text)
                for chunk in page_chunks:
                    chunks.append({
                        'content': chunk,
                        'type': 'text',
                        'page': page_num,
                        'position': None
                    })
        
        return chunks
    
    async def extract_image(self, image_bytes: bytes) -> List[Dict]:
        image = Image.open(io.BytesIO(image_bytes))
        text = pytesseract.image_to_string(
            image,
            lang=os.getenv('TESSERACT_LANG', 'deu+eng')
        )
        
        chunks = self.split_text(text) if text.strip() else ["[Bild ohne erkannten Text]"]
        return [{'content': chunk, 'type': 'text', 'page': 1} for chunk in chunks]
    
    async def extract_docx(self, docx_bytes: bytes) -> List[Dict]:
        if not HAS_DOCX:
            return [{'content': '[DOCX file - python-docx not installed]', 'type': 'text'}]
            
        chunks = []
        doc = DocxDocument(io.BytesIO(docx_bytes))
        
        full_text = '\n'.join([para.text for para in doc.paragraphs])
        text_chunks = self.split_text(full_text)
        
        for chunk in text_chunks:
            chunks.append({'content': chunk, 'type': 'text'})
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                table_data.append([cell.text for cell in row.cells])
            
            if table_data:
                chunks.append({
                    'content': json.dumps(table_data),
                    'type': 'table',
                    'metadata': {'table_index': table_idx}
                })
        
        return chunks
    
    async def extract_spreadsheet(self, file_bytes: bytes, file_type: str) -> List[Dict]:
        if not HAS_PANDAS:
            return [{'content': '[Spreadsheet file - pandas not installed]', 'type': 'text'}]
            
        chunks = []
        
        if file_type == 'csv':
            df = pd.read_csv(io.BytesIO(file_bytes))
        else:
            df = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
        
        if isinstance(df, dict):  # Multiple sheets
            for sheet_name, sheet_df in df.items():
                chunks.extend(self._process_dataframe(sheet_df, {'sheet': sheet_name}))
        else:
            chunks.extend(self._process_dataframe(df))
        
        return chunks
    
    def _process_dataframe(self, df: pd.DataFrame, metadata: dict = None) -> List[Dict]:
        chunks = []
        
        # Split large tables
        for i in range(0, len(df), 50):
            chunk_df = df.iloc[i:i+50]
            chunks.append({
                'content': chunk_df.to_json(orient='records'),
                'type': 'table',
                'metadata': {**(metadata or {}), 'row_start': i}
            })
        
        return chunks
    
    async def extract_email(self, eml_bytes: bytes) -> List[Dict]:
        chunks = []
        msg = BytesParser(policy=policy.default).parsebytes(eml_bytes)
        
        # Email metadata
        chunks.append({
            'content': json.dumps({
                'from': msg['From'],
                'to': msg['To'],
                'subject': msg['Subject'],
                'date': msg['Date']
            }),
            'type': 'metadata'
        })
        
        # Email body
        body = msg.get_body(preferencelist=('plain', 'html'))
        if body:
            content = body.get_content()
            chunks.extend([
                {'content': chunk, 'type': 'text'} 
                for chunk in self.split_text(content)
            ])
        
        return chunks
    
    def extract_text(self, text: str) -> List[Dict]:
        chunks = self.split_text(text)
        return [{'content': chunk, 'type': 'text'} for chunk in chunks]
    
    def extract_html(self, html: str) -> List[Dict]:
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html, 'html.parser')
            text = soup.get_text(separator='\n', strip=True)
            return self.extract_text(text)
        except ImportError:
            # Fallback if BeautifulSoup not available
            import re
            text = re.sub('<[^<]+?>', '', html)
            return self.extract_text(text)
    
    async def extract_archive(self, zip_bytes: bytes) -> List[Dict]:
        chunks = []
        
        with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
            for name in zf.namelist():
                if not name.endswith('/'):
                    chunks.append({
                        'content': f"Archive contains: {name}",
                        'type': 'metadata',
                        'metadata': {'archived_file': name}
                    })
        
        return chunks
    
    def split_text(self, text: str) -> List[str]:
        words = text.split()
        chunks = []
        current_chunk = []
        current_size = 0
        
        for word in words:
            current_chunk.append(word)
            current_size += len(word) + 1
            
            if current_size >= self.chunk_size:
                chunks.append(' '.join(current_chunk))
                current_chunk = []
                current_size = 0
        
        if current_chunk:
            chunks.append(' '.join(current_chunk))
        
        return chunks or [""]
    
    async def run(self):
        await self.connect()
        logger.info("OCR Agent V2 started")
        
        try:
            await self.process_queue()
        except KeyboardInterrupt:
            logger.info("Shutting down")
        finally:
            await self.close()


def main():
    db_url = os.getenv('DATABASE_URL', 'postgresql://docintell:docintell123@postgres:5432/document_intelligence')
    agent = OCRAgentV2(db_url)
    asyncio.run(agent.run())


if __name__ == "__main__":
    main()
